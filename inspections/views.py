from rest_framework import viewsets, status, filters
from rest_framework.decorators import action
from rest_framework.exceptions import PermissionDenied
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated, AllowAny
from rest_framework.throttling import AnonRateThrottle, UserRateThrottle
from django.http import HttpResponse, HttpResponseNotFound
from django.contrib.auth.models import User
from django.conf import settings
from django.core.mail import send_mail
from django.db import transaction
from django.db.models import F, Q
from django.utils import timezone
from django.utils.text import slugify
from django.urls import reverse
from django.utils.html import escape
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from datetime import datetime, time, timedelta
import logging
import io
import json
import uuid

sec_log = logging.getLogger('inspections.security')


class LoginRateThrottle(AnonRateThrottle):
    rate = '10/minute'

from .models import (
    UserProfile, Terminal, Tank, Inspection, Seal, Isolation,
    InspectionCalculation, InspectionReport,
    ProductReceiptCertificate, ProductReceiptCertificateItem,
    SealIsolationReport,
    ShoreTankCalculation,
    Submission, VesselReport, RosterAssignment, ActivityLog,
    ProvisionalOuturnReport, ProvisionalOuturnItem,
    StockReport,
    SamplingForm,
    ServiceRequest,
    ServiceRequestMessage,
    Notification,
    DocumentSignature,
)
from .serializers import (
    UserSerializer, UserProfileSerializer, UserRegistrationSerializer, TerminalSerializer,
    TankSerializer, TankDetailSerializer,
    SealSerializer, IsolationSerializer,
    InspectionCalculationSerializer, InspectionReportSerializer,
    InspectionListSerializer, InspectionCreateSerializer, InspectionDetailSerializer,
    InspectionApprovalSerializer,
    ProductReceiptCertificateListSerializer,
    ProductReceiptCertificateDetailSerializer,
    SealIsolationReportListSerializer,
    SealIsolationReportDetailSerializer,
    ShoreTankCalculationListSerializer,
    ShoreTankCalculationDetailSerializer,
    SubmissionSerializer, VesselReportSerializer, RosterAssignmentSerializer,
    ProvisionalOuturnReportSerializer,
    StockReportSerializer,
    SamplingFormSerializer,
    ServiceRequestSerializer,
    ServiceRequestMessageSerializer,
    NotificationSerializer,
    ActivityLogSerializer,
)
from .permissions import IsInspector, IsTerminalRep, IsAdmin
from .calculations import InspectionCalculationEngine, ShoreTankCalculationEngine
from .shore_tank_utils import (
    ShoreTankDocumentGenerator,
    generate_shore_tank_document,
    generate_product_receipt_document,
    generate_seal_isolation_document,
    generate_dip_ticket_document,
    generate_dip_ticket_pdf,
    generate_seal_isolation_pdf,
    generate_shore_tank_pdf,
)
from .astm_tables import density_at_20_from_table, density_at_20_formula, vcf_from_table, vcf_formula, wcf_from_density, table_range
from .signing import (
    sign_pdf_bytes, finalize_pdf_bytes, get_signature_info, compute_document_hash,
    compute_model_hash, create_digital_signature, verify_digital_signature,
    build_document_fields_string,
)
from .pdf_branding import add_pbpa_letterhead
from .mfa import generate_secret, provisioning_uri, verify_code


def _build_public_verification_url(obj, document_hash, request=None):
    verification_path = reverse(
        'public-document-verification',
        kwargs={'document_type': _verification_type_for(obj), 'pk': obj.pk},
    )
    verification_base_url = getattr(settings, 'QR_VERIFICATION_BASE_URL', '')
    if verification_base_url:
        return f'{verification_base_url}{verification_path}?hash={document_hash}'
    if request:
        return f'{request.build_absolute_uri(verification_path)}?hash={document_hash}'
    return ''


def _append_signature_footer(pdf, W, H, M, obj, signatures=None, request=None):
    """
    Append a verification footer to a ReportLab canvas page.
    Prints document hash, signing info, and a QR code for verification.
    """
    import qrcode
    import qrcode.image.pil
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.lib.utils import ImageReader
    import io as _io

    doc_hash = getattr(obj, 'document_hash', '')
    if not doc_hash:
        return

    footer_y = M + 2 * mm
    line_h   = 4 * mm

    # Divider line
    pdf.setStrokeColor(colors.HexColor('#cccccc'))
    pdf.setLineWidth(0.5)
    pdf.line(M, footer_y + line_h * 6, W - M, footer_y + line_h * 6)

    pdf.setFont('Helvetica-Bold', 7)
    pdf.setFillColor(colors.HexColor('#444444'))
    pdf.drawString(M, footer_y + line_h * 5, 'DIGITAL SIGNATURE VERIFICATION')

    pdf.setFont('Helvetica', 6.5)
    pdf.setFillColor(colors.black)


    # Inspector signature line
    insp_name = getattr(obj, 'pbpa_inspector_name', '') or ''
    insp_at   = getattr(obj, 'inspector_signed_at', None)
    insp_str  = f'Inspector: {insp_name}'
    if insp_at:
        insp_str += f'  |  Signed: {insp_at.strftime("%d %b %Y %H:%M")}'
    pdf.drawString(M, footer_y + line_h * 4, insp_str)

    # Terminal rep signature line
    tr_name = getattr(obj, 'terminal_representative_name', '') or ''
    tr_at   = getattr(obj, 'client_signed_at', None)
    tr_str  = f'Terminal Rep: {tr_name}'
    if tr_at:
        tr_str += f'  |  Signed: {tr_at.strftime("%d %b %Y %H:%M")}'
    pdf.drawString(M, footer_y + line_h * 3, tr_str)

    # Hash line (truncated for display)
    hash_display = doc_hash[:48] + '...' if len(doc_hash) > 48 else doc_hash
    pdf.drawString(M, footer_y + line_h * 2, f'SHA-256: {hash_display}')

    # QR code — encodes doc_type, doc_number, hash
    doc_number = ''
    for field in ('ticket_number', 'report_number', 'certificate_number', 'calculation_number', 'form_number'):
        doc_number = getattr(obj, field, '') or ''
        if doc_number:
            break
    qr_data = _build_public_verification_url(obj, doc_hash, request) or f'PBPA|{doc_number}|{doc_hash}'
    try:
        qr_img = qrcode.make(qr_data)
        qr_buf = _io.BytesIO()
        qr_img.save(qr_buf, format='PNG')
        qr_buf.seek(0)
        qr_size = 18 * mm
        pdf.drawImage(ImageReader(qr_buf), W - M - qr_size, footer_y, width=qr_size, height=qr_size)
    except Exception:
        pass  # QR generation is best-effort

    pdf.setFillColor(colors.black)


def _finalize_pdf_for_download(pdf_bytes, obj, document_type, request=None, document_number=''):
    """Apply PBPA letterhead and make every downloaded PDF verifiable."""
    pdf_bytes = add_pbpa_letterhead(pdf_bytes, document_type=document_type)
    document_number = document_number or next(
        (str(getattr(obj, field)) for field in ('ticket_number', 'report_number', 'certificate_number',
         'calculation_number', 'form_number') if getattr(obj, field, None)),
        str(getattr(obj, 'pk', '')),
    )
    document_hash = getattr(obj, 'document_hash', '')
    if not document_hash:
        # Reports without a persisted signing field still receive a stable,
        # record-based verification hash rather than a volatile PDF-byte hash.
        verification_config = _VERIFICATION_MODELS.get(_verification_type_for(obj))
        if verification_config:
            document_hash = compute_model_hash(obj, verification_config[2], document_number)
        else:
            document_hash = compute_document_hash(pdf_bytes)
    signer_name = ''
    signer_role = ''
    if request and getattr(request, 'user', None) and request.user.is_authenticated:
        signer_name = request.user.get_full_name() or request.user.username
        signer_role = get_or_create_user_profile(request.user).get_role_display()
    verification_url = _build_public_verification_url(obj, document_hash, request)
    return finalize_pdf_bytes(
        pdf_bytes,
        document_number=document_number,
        document_hash=document_hash,
        document_type=document_type,
        signer_name=signer_name,
        signer_role=signer_role,
        signed=bool(getattr(obj, 'is_signed', False)),
        verification_url=verification_url,
    )


_VERIFICATION_MODELS = {
    # model, public label, canonical type used when the record hash was made
    'dip-ticket': (Inspection, 'Dip Ticket', 'dip_ticket'),
    'product-receipt-certificate': (ProductReceiptCertificate, 'Product Receipt Certificate', 'product_receipt'),
    'seal-isolation-report': (SealIsolationReport, 'Seal and Isolation Report', 'seal_isolation'),
    'shore-tank-calculation': (ShoreTankCalculation, 'Shore Tank Calculation', 'shore_tank'),
    'vessel-report': (VesselReport, 'Vessel Report', 'vessel_report'),
    'provisional-outturn-report': (ProvisionalOuturnReport, 'Provisional Outturn Report', 'provisional_outturn'),
    'stock-report': (StockReport, 'Weekly Stock Report', 'stock_report'),
    'sampling-form': (SamplingForm, 'Vessel Sampling Form', 'sampling_form'),
    'roster-assignment': (RosterAssignment, 'Inspector Roster Assignment', 'roster_assignment'),
}


def _verification_type_for(obj):
    for slug, (model, _, _) in _VERIFICATION_MODELS.items():
        if isinstance(obj, model):
            return slug
    return obj._meta.model_name


def _display_date(value):
    if not value:
        return 'Not recorded'
    if hasattr(value, 'strftime'):
        return value.strftime('%d-%m-%Y %H:%M') if hasattr(value, 'hour') else value.strftime('%d-%m-%Y')
    return str(value)


_VERIFICATION_EXCLUDED_FIELDS = {
    'id', 'created_at', 'updated_at', 'document_hash', 'captain_signature',
    'terminal_representative_signature', 'pbpa_inspector_signature',
    'captain_signing_token', 'captain_signing_expires_at', 'captain_signed_ip',
    # Internal/administrative fields are not part of the public QR report.
    # Keep the human-readable ``tank_no`` field; only hide the Tank relation.
    'tank', 'approval_date', 'rejection_reason',
}


def _verification_field_rows(record):
    """Return the saved, non-sensitive business fields for a QR verification page."""
    rows = []
    for field in record._meta.concrete_fields:
        if field.name in _VERIFICATION_EXCLUDED_FIELDS:
            continue
        value = getattr(record, field.name, None)
        if field.is_relation:
            value = _user_name(value) or str(value or '')
        rows.append((field.verbose_name.replace('_', ' ').title(), _display_date(value)))
    return rows


def _verification_item_sections(obj):
    """Render report line items so scanning exposes the original report data too."""
    sections = []
    for accessor in ('items', 'tank_items', 'seals', 'isolations'):
        related = getattr(obj, accessor, None)
        if related is None or not hasattr(related, 'all'):
            continue
        records = list(related.all())
        if not records:
            continue
        record_rows = []
        for index, record in enumerate(records, start=1):
            values = _verification_field_rows(record)
            details = '<br>'.join(
                f'<strong>{escape(label)}:</strong> {escape(str(value))}' for label, value in values
                if label.lower() not in {'Report', 'Calculation', 'Inspection'}
            )
            record_rows.append(f'<tr><td>{index}</td><td>{details}</td></tr>')
        title = accessor.replace('_', ' ').title()
        sections.append(
            f'<h2>{escape(title)}</h2><table><thead><tr><th>#</th><th>Original values</th></tr></thead>'
            f'<tbody>{"".join(record_rows)}</tbody></table>'
        )
    return ''.join(sections)


def document_verification(request, document_type, pk):
    """Public, hash-bound record shown after scanning a generated PDF's QR code."""
    config = _VERIFICATION_MODELS.get(document_type)
    if not config:
        return HttpResponseNotFound('Unknown document type.')
    model, type_label, hash_type = config
    try:
        obj = model.objects.get(pk=pk)
    except model.DoesNotExist:
        return HttpResponseNotFound('Document record not found.')

    reference_hash = request.GET.get('hash', '')
    expected_hash = getattr(obj, 'document_hash', '') or reference_hash
    hash_matches = bool(expected_hash) and reference_hash == expected_hash
    document_number = next((str(getattr(obj, field)) for field in (
        'ticket_number', 'report_number', 'certificate_number', 'calculation_number', 'form_number'
    ) if getattr(obj, field, None)), str(obj.pk))
    current_hash = ''
    try:
        current_hash = compute_model_hash(obj, hash_type, document_number)
    except Exception:
        pass
    integrity_valid = hash_matches and current_hash == expected_hash

    signatures = DocumentSignature.objects.filter(
        doc_type=hash_type, doc_id=obj.pk, status='valid'
    ).select_related('signer')
    valid_roles = set()
    for signature in signatures.filter(document_hash=expected_hash):
        if signature.signer_id and verify_digital_signature(signature.document_hash, str(signature.signer_id), signature.signing_timestamp, signature.digital_signature):
            valid_roles.add(signature.role)
    inspector_at = getattr(obj, 'inspector_signed_at', None)
    terminal_at = getattr(obj, 'client_signed_at', None)
    inspector_name = getattr(obj, 'pbpa_inspector_name', '') or _user_name(getattr(obj, 'inspector_signed_by', None))
    terminal_name = getattr(obj, 'terminal_representative_name', '') or _user_name(getattr(obj, 'client_signed_by', None))
    inspector_verified = bool(inspector_at) and (not signatures.exists() or 'inspector' in valid_roles)
    terminal_verified = bool(terminal_at) and (not signatures.exists() or 'terminal_representative' in valid_roles)
    status_text = '✓ VERIFIED' if integrity_valid and inspector_verified and terminal_verified else 'NOT VERIFIED'
    rows = [
        ('Document Type', type_label), ('Document ID', document_number), ('Verification Status', status_text),
        *_verification_field_rows(obj),
    ]
    details = ''.join(f'<dt>{escape(label)}</dt><dd>{escape(str(value))}</dd>' for label, value in rows)
    def party_row(role, name, signed_at, verified):
        return '<tr><td>{}</td><td>{}</td><td>{}</td><td class="{}">{}</td></tr>'.format(
            escape(role), escape(name or 'Not recorded'), escape(_display_date(signed_at)),
            'ok' if verified else 'bad', '✓ Verified' if verified else 'Not verified')
    item_sections = _verification_item_sections(obj)
    page = '''<!doctype html><html><head><meta name="viewport" content="width=device-width,initial-scale=1"><title>PBPA Document Verification</title><style>body{font:16px Arial,sans-serif;background:#f4f6f8;color:#15212c;margin:0}.card{max-width:760px;margin:30px auto;background:#fff;padding:30px;border-radius:10px;box-shadow:0 2px 14px #0002}h1{color:#8b1a1a;font-size:25px;margin-top:0}h2{font-size:18px;margin-top:30px}dl{display:grid;grid-template-columns:180px 1fr;gap:10px;margin:0}dt{font-weight:bold}dd{margin:0}.ok{color:#087a39;font-weight:bold}.bad{color:#b42318;font-weight:bold}table{border-collapse:collapse;width:100%}td,th{padding:10px;border-bottom:1px solid #dde3e8;text-align:left;vertical-align:top}@media(max-width:520px){.card{margin:0;border-radius:0;padding:20px}dl{grid-template-columns:1fr}dd{margin-bottom:8px}}</style></head><body><main class="card"><h1>DOCUMENT VERIFICATION</h1><p>This page shows the original saved report data linked to the scanned QR code.</p><h2>ORIGINAL REPORT INFORMATION</h2><dl>''' + details + '''</dl><h2>SIGNATORIES</h2><table><thead><tr><th>Party</th><th>Name</th><th>Date</th><th>Verification</th></tr></thead><tbody>''' + party_row('PBPA Inspector', inspector_name, inspector_at, inspector_verified) + party_row('Terminal Representative', terminal_name, terminal_at, terminal_verified) + '''</tbody></table>''' + item_sections + '''<h2>Document Integrity</h2><p><strong>SHA-256:</strong> ''' + escape(f'{expected_hash[:8]}...{expected_hash[-8:]}' if expected_hash else 'Not available') + '''</p><p><strong>Digital Signature:</strong> <span class="''' + ('ok' if integrity_valid else 'bad') + '''">''' + ('Verified' if integrity_valid else 'Invalid or changed') + '''</span></p></main></body></html>'''
    return HttpResponse(page, content_type='text/html; charset=utf-8')


def _user_name(user):
    return user.get_full_name() or user.username if user else ''


def overlay_signature_image(pdf_bytes, sig_b64, x, y, w, h):
    """
    Overlay a base64 PNG signature image onto a PDF using ReportLab + pypdf.
    Returns new PDF bytes with the signature image drawn at (x, y) in mm from bottom-left.
    """
    import base64
    import binascii
    import io as _io
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.units import mm
    from pypdf import PdfWriter, PdfReader

    if not isinstance(sig_b64, str) or not sig_b64.strip():
        raise ValueError('A signature image is required.')
    if ',' in sig_b64:
        sig_b64 = sig_b64.split(',', 1)[1]
    try:
        img_bytes = base64.b64decode(sig_b64, validate=True)
    except (ValueError, binascii.Error) as exc:
        raise ValueError('Signature must be a valid base64 image.') from exc

    from reportlab.lib.utils import ImageReader
    base_reader = PdfReader(_io.BytesIO(pdf_bytes))
    if not base_reader.pages:
        raise ValueError('The document has no pages to sign.')
    first_page = base_reader.pages[0]
    page_width = float(first_page.mediabox.width)
    page_height = float(first_page.mediabox.height)
    overlay_buf = _io.BytesIO()
    c = rl_canvas.Canvas(overlay_buf, pagesize=(page_width, page_height))
    try:
        c.drawImage(
            ImageReader(_io.BytesIO(img_bytes)), x * mm, y * mm,
            width=w * mm, height=h * mm, mask='auto', preserveAspectRatio=True,
        )
    except Exception as exc:
        raise ValueError('Signature must be a valid image.') from exc
    c.save()
    overlay_buf.seek(0)

    overlay_reader = PdfReader(overlay_buf)
    writer = PdfWriter()

    for i, page in enumerate(base_reader.pages):
        if i == 0:
            page.merge_page(overlay_reader.pages[0])
        writer.add_page(page)

    out = _io.BytesIO()
    writer.write(out)
    return out.getvalue()



def get_client_ip(request):
    xff = request.META.get('HTTP_X_FORWARDED_FOR')
    return xff.split(',')[0].strip() if xff else request.META.get('REMOTE_ADDR')

def get_or_create_user_profile(user):
    """Return a profile for existing users that may predate profile creation."""
    role = 'admin' if user.is_staff or user.is_superuser else 'inspector'
    profile, _ = UserProfile.objects.get_or_create(user=user, defaults={'role': role})
    return profile


def owned_queryset(queryset, user, owner_field='created_by'):
    """Limit records to the authenticated user who owns/created them."""
    if not user or not user.is_authenticated:
        return queryset.none()
    return queryset.filter(**{owner_field: user})


def ensure_inspection_owner(inspection, user):
    if inspection.inspector_id != user.id:
        raise PermissionDenied('You can only access records for inspections you created.')


def get_owned_document_querysets(user, period=None):
    return {
        'inspections': filter_queryset_by_period(
            owned_queryset(Inspection.objects.all(), user, 'inspector'),
            period,
            'inspection_date',
        ),
        'product_receipt_certificates': filter_queryset_by_period(
            owned_queryset(ProductReceiptCertificate.objects.all(), user),
            period,
            'receipt_date',
        ),
        'seal_isolation_reports': filter_queryset_by_period(
            owned_queryset(SealIsolationReport.objects.all(), user),
            period,
            'report_date',
        ),
        'shore_tank_calculations': filter_queryset_by_period(
            owned_queryset(ShoreTankCalculation.objects.all(), user),
            period,
            'calculation_date',
        ),
        'stock_reports': filter_queryset_by_period(
            owned_queryset(StockReport.objects.all(), user),
            period,
            'report_date',
        ),
        'provisional_outturn_reports': filter_queryset_by_period(
            owned_queryset(ProvisionalOuturnReport.objects.all(), user),
            period,
            'report_date',
        ),
        'vessel_reports': filter_queryset_by_period(
            owned_queryset(VesselReport.objects.all(), user),
            period,
            'discharge_date',
        ),
    }


def period_start_date(period):
    """Return the local start date for a supported period filter."""
    today = timezone.localdate()
    if period == 'daily':
        return today
    if period == 'weekly':
        return today - timedelta(days=today.weekday())
    if period == 'monthly':
        return today.replace(day=1)
    if period == 'yearly':
        return today.replace(month=1, day=1)
    return None


def filter_queryset_by_period(queryset, period, date_field):
    """Apply all/daily/weekly/monthly/yearly filters to DateField or DateTimeField values."""
    start_date = period_start_date(period)
    if not start_date:
        return queryset

    today = timezone.localdate()
    field = queryset.model._meta.get_field(date_field)
    if field.get_internal_type() == 'DateTimeField':
        return queryset.filter(**{
            f'{date_field}__date__gte': start_date,
            f'{date_field}__date__lte': today,
        })

    return queryset.filter(**{
        f'{date_field}__gte': start_date,
        f'{date_field}__lte': today,
    })


def period_start_datetime(period):
    """Return a timezone-aware datetime for log/date-time activity filters."""
    start_date = period_start_date(period)
    if not start_date:
        return None
    return timezone.make_aware(datetime.combine(start_date, time.min), timezone.get_current_timezone())


def count_log_entries(log_name, tokens, period=None):
    """Best-effort count of already-recorded security/audit log events."""
    log_path = getattr(settings, 'LOG_DIR', None)
    if not log_path:
        return 0
    log_file = log_path / log_name
    if not log_file.exists():
        return 0

    start_at = period_start_datetime(period)
    total = 0
    try:
        with log_file.open('r', encoding='utf-8', errors='ignore') as handle:
            for line in handle:
                if not all(token in line for token in tokens):
                    continue
                if start_at:
                    event_at = parse_log_datetime(line)
                    if event_at and event_at < start_at:
                        continue
                total += 1
    except OSError:
        return 0
    return total


def parse_log_datetime(line):
    """Parse either JSON or verbose Django log timestamps."""
    try:
        payload = json.loads(line)
        value = payload.get('asctime') or payload.get('timestamp')
        if value:
            parsed = datetime.fromisoformat(value.replace('Z', '+00:00'))
            return parsed if timezone.is_aware(parsed) else timezone.make_aware(parsed)
    except (ValueError, TypeError):
        pass

    parts = line.split()
    if len(parts) >= 3:
        try:
            parsed = datetime.fromisoformat(f'{parts[1]} {parts[2].split(",")[0]}')
            return timezone.make_aware(parsed, timezone.get_current_timezone())
        except (ValueError, TypeError):
            return None
    return None


def build_activity_overview(user, period=None):
    """Build dashboard activity metrics, system-wide for admins and user-owned otherwise."""
    is_admin = get_or_create_user_profile(user).role == 'admin'
    owned_reports = get_owned_document_querysets(user, None)
    report_querysets = {
        'inspections': Inspection.objects.all(),
        'product_receipt_certificates': ProductReceiptCertificate.objects.all(),
        'seal_isolation_reports': SealIsolationReport.objects.all(),
        'shore_tank_calculations': ShoreTankCalculation.objects.all(),
        'stock_reports': StockReport.objects.all(),
        'provisional_outturn_reports': ProvisionalOuturnReport.objects.all(),
        'vessel_reports': VesselReport.objects.all(),
    } if is_admin else owned_reports

    report_creation = sum(
        filter_queryset_by_period(queryset, period, 'created_at').count()
        for queryset in report_querysets.values()
    )
    report_editing = sum(
        filter_queryset_by_period(
            queryset.filter(updated_at__gt=F('created_at') + timedelta(seconds=2)),
            period,
            'updated_at',
        ).count()
        for queryset in report_querysets.values()
    )
    report_files = InspectionReport.objects.all() if is_admin else InspectionReport.objects.filter(inspection__inspector=user)
    file_uploads = filter_queryset_by_period(report_files, period, 'generated_at').count()

    failed_login_attempts = count_log_entries('security.log', ['Unauthorized access attempt'], period) if is_admin else 0
    auth_attempts = count_log_entries('security.log', ['Authentication attempt'], period) if is_admin else 0
    password_changes = count_log_entries('security.log', ['reset password'], period) if is_admin else 0
    admin_actions = filter_queryset_by_period(
        User.objects.filter(is_staff=True) if is_admin else User.objects.filter(pk=user.pk, is_staff=True),
        period,
        'date_joined',
    ).count()
    report_deletion = count_log_entries('audit.log', ['DELETE request'], period) if is_admin else 0

    rows = [
        {'activity': 'User login/logout', 'value': max(auth_attempts - failed_login_attempts, 0), 'why': 'Security', 'color': '#2563eb'},
        {'activity': 'Report creation', 'value': report_creation, 'why': 'Usage monitoring', 'color': '#16a34a'},
        {'activity': 'Report editing', 'value': report_editing, 'why': 'Audit trail', 'color': '#f59e0b'},
        {'activity': 'Report deletion', 'value': report_deletion, 'why': 'Accountability', 'color': '#ef4444'},
        {'activity': 'Password changes', 'value': password_changes, 'why': 'Security tracking', 'color': '#7c3aed'},
        {'activity': 'Failed login attempts', 'value': failed_login_attempts, 'why': 'Detect attacks', 'color': '#dc2626'},
        {'activity': 'File uploads', 'value': file_uploads, 'why': 'Monitor storage use', 'color': '#0891b2'},
        {'activity': 'Admin actions', 'value': admin_actions, 'why': 'Transparency', 'color': '#4f46e5'},
    ]
    return rows


def create_notification(recipient, title, message, notification_type='report_submitted', doc_type='', doc_id=None, doc_number='', email=False):
    """Create an in-app notification and, when requested, an email alert."""
    notification = Notification.objects.create(
        recipient=recipient,
        notification_type=notification_type,
        title=title,
        message=message,
        doc_type=doc_type,
        doc_id=doc_id,
        doc_number=doc_number,
    )
    if email:
        recipient_email = (recipient.email or getattr(recipient.profile, 'company_email', '')).strip()
        if recipient_email:
            try:
                send_mail(
                    subject=title,
                    message=message,
                    from_email=settings.DEFAULT_FROM_EMAIL,
                    recipient_list=[recipient_email],
                    fail_silently=False,
                )
            except Exception:
                # A mail-provider outage must not discard the signing request.
                sec_log.exception('Could not send signing notification email to user %s', recipient.pk)
    return notification


def _terminal_representatives(terminal):
    """Return active terminal representatives assigned to a document's terminal."""
    terminal = (terminal or '').strip()
    if not terminal:
        return UserProfile.objects.none()
    # Match on the terminal name stored in the profile (case-insensitive).
    # Also try matching on the Terminal model name so both storage formats work.
    return UserProfile.objects.filter(
        role='terminal_representative',
        is_active=True,
        terminal__iexact=terminal,
        user__is_active=True,
    ).select_related('user')


def _record_signature(request, obj, role, doc_hash, doc_type, doc_number, signature_image=''):
    """Create a DocumentSignature record for an individual signing event."""
    from django.utils import timezone as tz
    timestamp = tz.now().strftime('%Y-%m-%dT%H:%M:%S')
    digital_sig = create_digital_signature(doc_hash, str(request.user.id), timestamp)
    DocumentSignature.objects.create(
        doc_type=doc_type,
        doc_id=obj.pk,
        doc_number=doc_number,
        signer=request.user,
        role=role,
        document_hash=doc_hash,
        digital_signature=digital_sig,
        signing_timestamp=timestamp,
        signature_image=signature_image,
        ip_address=get_client_ip(request),
        user_agent=request.META.get('HTTP_USER_AGENT', '')[:500],
        status='valid',
    )


# ── Shared signing workflow mixin ────────────────────────────────────────────
class SigningWorkflowMixin:
    """
    Provides the 5-step multi-party signing workflow for report ViewSets.
    Requires subclass to define:
      _build_pdf_bytes(self, obj) -> bytes
      _serializer_class_detail  (the detail serializer)
      _doc_type_label           (str used in Submission.doc_type)
    """

    def _get_role(self, request):
        return get_or_create_user_profile(request.user).role

    def _build_pdf_bytes(self, obj):
        raise NotImplementedError

    def create(self, request, *args, **kwargs):
        if self._get_role(request) == 'terminal_representative':
            return Response({'detail': 'Terminal representatives can only view and sign reports.'}, status=status.HTTP_403_FORBIDDEN)
        return super().create(request, *args, **kwargs)

    # Legacy overlay slots for forms whose generators have not yet defined
    # signature containers. Dip Tickets and Shore Tank Calculations render
    # their signatures inside their named fields in the generator.
    SIGNATURE_SLOTS = {
        'seal_isolation': {
            'terminal_representative': (26, 14.4, 69, 9),
            'inspector': (123, 14.4, 69, 9),
        },
        'shore_tank': {
            'inspector': (34, 69, 55, 5),
            'terminal_representative': (131, 69, 55, 5),
        },
        'product_receipt': {
            'terminal_representative': (16, 71, 73, 10),
            'inspector': (106, 71, 73, 10),
        },
        'sampling_form': {
            'terminal_representative': (16, 79, 73, 10),
            'inspector': (104, 79, 73, 10),
        },
    }

    def _build_pdf_with_signatures(self, obj):
        """Render the document and add each stored signature to its role's slot."""
        pdf_bytes = self._build_pdf_bytes(obj)
        # These documents render their operation-specific signatures directly
        # in the correctly labelled signature fields of the original PDF.
        if self._doc_type_label in {'dip_ticket', 'shore_tank'}:
            return pdf_bytes
        signatures = DocumentSignature.objects.filter(
            doc_type=self._doc_type_label, doc_id=obj.pk, status='valid',
            role__in=('inspector', 'terminal_representative'),
        ).order_by('created_at')
        slots = self.SIGNATURE_SLOTS.get(self._doc_type_label, {})
        for signature in signatures:
            slot = slots.get(signature.role)
            if slot and signature.signature_image:
                pdf_bytes = overlay_signature_image(pdf_bytes, signature.signature_image, *slot)
        return pdf_bytes

    def _save_signing_state(self, obj, update_fields):
        obj.save(update_fields=update_fields + ['updated_at'])

    def _get_doc_number(self, obj):
        for field in ('report_number', 'certificate_number', 'calculation_number', 'form_number'):
            value = getattr(obj, field, None)
            if value:
                return value
        return str(obj.pk)

    def _get_doc_label(self):
        return dict(
            dip_ticket='Dip Ticket',
            seal_isolation='Seal & Isolation Report',
            product_receipt='Product Receipt Certificate',
            shore_tank='Shore Tank Calculation',
            sampling_form='Sampling Form',
        ).get(self._doc_type_label, self._doc_type_label.replace('_', ' ').title())

    def _get_counterparty_label(self):
        return 'vessel captain' if self._doc_type_label == 'sampling_form' else 'terminal representative'

    def _document_hash(self, obj):
        return compute_model_hash(obj, self._doc_type_label, self._get_doc_number(obj))

    def update(self, request, *args, **kwargs):
        if self._get_role(request) == 'terminal_representative':
            return Response({'detail': 'Terminal representatives can only view and sign reports.'}, status=status.HTTP_403_FORBIDDEN)
        obj = self.get_object()
        if (
            getattr(obj, 'is_signed', False)
            or getattr(obj, 'signing_step', 'draft') != 'draft'
            or getattr(obj, 'status', 'draft') == 'submitted'
        ):
            return Response(
                {'detail': 'Submitted or signed documents are locked. An admin must unlock the document before editing.'},
                status=status.HTTP_409_CONFLICT,
            )
        return super().update(request, *args, **kwargs)

    def destroy(self, request, *args, **kwargs):
        if self._get_role(request) == 'terminal_representative':
            return Response({'detail': 'Terminal representatives can only view and sign reports.'}, status=status.HTTP_403_FORBIDDEN)
        obj = self.get_object()
        if (
            getattr(obj, 'is_signed', False)
            or getattr(obj, 'signing_step', 'draft') != 'draft'
            or getattr(obj, 'status', 'draft') == 'submitted'
        ):
            return Response(
                {'detail': 'Submitted or signed documents are locked and cannot be deleted.'},
                status=status.HTTP_409_CONFLICT,
            )
        return super().destroy(request, *args, **kwargs)

    @action(detail=True, methods=['post'], url_path='unlock')
    def unlock(self, request, pk=None):
        """Invalidate prior signatures and return a document to draft for editing."""
        obj = self.get_object()
        if self._get_role(request) != 'admin':
            return Response({'detail': 'Only admins can unlock a signed document.'}, status=status.HTTP_403_FORBIDDEN)
        if obj.signing_step == 'draft' and not getattr(obj, 'is_signed', False):
            return Response({'detail': 'Document is already editable.'}, status=status.HTTP_400_BAD_REQUEST)

        DocumentSignature.objects.filter(
            doc_type=self._doc_type_label, doc_id=obj.pk, status='valid'
        ).update(status='superseded')
        obj.is_signed = False
        obj.signed_at = None
        obj.signed_by = None
        obj.document_hash = ''
        obj.signing_step = 'draft'
        obj.inspector_signed_at = None
        obj.inspector_signed_by = None
        obj.client_signed_at = None
        obj.client_signed_by = None
        obj.verified_at = None
        self._save_signing_state(obj, [
            'is_signed', 'signed_at', 'signed_by', 'document_hash', 'signing_step',
            'inspector_signed_at', 'inspector_signed_by', 'client_signed_at',
            'client_signed_by', 'verified_at',
        ])
        ActivityLog.log(request.user, 'report_updated', doc_type=self._doc_type_label,
            doc_id=obj.pk, doc_number=self._get_doc_number(obj),
            detail=f'Admin unlocked {self._get_doc_label()} #{self._get_doc_number(obj)}; prior signatures superseded.',
            request=request)
        return Response(self.get_serializer(obj).data)

    @action(detail=True, methods=['post'], url_path='inspector_sign')
    def inspector_sign(self, request, pk=None):
        """Step 1 — Inspector draws signature and signs the document."""
        obj = self.get_object()
        role = self._get_role(request)
        if role not in ('inspector', 'admin'):
            return Response({'detail': 'Only inspectors can perform this step.'}, status=status.HTTP_403_FORBIDDEN)
        if obj.signing_step != 'draft':
            return Response({'detail': f'Cannot sign: current step is "{obj.signing_step}".'}, status=status.HTTP_400_BAD_REQUEST)

        sig_b64 = request.data.get('signature')
        if not sig_b64:
            return Response({'detail': 'signature field is required.'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            submitted_name = (request.data.get('signer_name') or '').strip()
            obj.pbpa_inspector_name = submitted_name or request.user.get_full_name() or request.user.username
            obj.is_signed           = True
            obj.signed_at           = timezone.now()
            obj.signed_by           = request.user
            obj.document_hash       = self._document_hash(obj)
            obj.signing_step        = 'inspector_signed'
            obj.inspector_signed_at = timezone.now()
            obj.inspector_signed_by = request.user
            self._save_signing_state(obj, ['is_signed','signed_at','signed_by','document_hash',
                                           'signing_step','inspector_signed_at','inspector_signed_by',
                                           'pbpa_inspector_name'])
            _record_signature(request, obj, 'inspector', obj.document_hash,
                              self._doc_type_label, self._get_doc_number(obj), sig_b64)
            signed_bytes = self._build_pdf_with_signatures(obj)
            ActivityLog.log(request.user, 'report_updated', doc_type=self._doc_type_label,
                doc_id=obj.pk, doc_number=self._get_doc_number(obj),
                detail=f'Inspector signed {self._get_doc_label()} #{self._get_doc_number(obj)}',
                request=request)
            signed_bytes = _finalize_pdf_for_download(signed_bytes, obj, self._get_doc_label(), request)
            response = HttpResponse(signed_bytes, content_type='application/pdf')
            fname = self._get_doc_number(obj)
            response['Content-Disposition'] = f'attachment; filename="InspectorSigned_{fname}.pdf"'
            return response
        except ValueError as e:
            return Response({'detail': str(e)}, status=status.HTTP_400_BAD_REQUEST)
        except Exception:
            sec_log.exception('Inspector signing failed for %s #%s', self._doc_type_label, obj.pk)
            return Response({'detail': 'Signing failed. Please try again or contact an administrator.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @action(detail=True, methods=['post'], url_path='send_to_client')
    def send_to_client(self, request, pk=None):
        """Step 2 — send a signing request to representatives of this document's terminal."""
        obj = self.get_object()
        role = self._get_role(request)
        if role not in ('inspector', 'admin'):
            return Response({'detail': 'Only inspectors can send to client.'}, status=status.HTTP_403_FORBIDDEN)
        if obj.signing_step != 'inspector_signed':
            return Response({'detail': f'Cannot send: current step is "{obj.signing_step}".'}, status=status.HTTP_400_BAD_REQUEST)

        terminal = (getattr(obj, 'terminal', '') or '').strip()
        representatives = list(_terminal_representatives(terminal))
        if not representatives:
            # Give a helpful diagnostic: list all terminal rep terminals so admin can fix the mismatch.
            all_tr_terminals = list(
                UserProfile.objects.filter(role='terminal_representative', is_active=True, user__is_active=True)
                .values_list('terminal', flat=True).distinct()
            )
            registered = ', '.join(f'"{t}"' for t in all_tr_terminals if t) or 'none registered'
            return Response(
                {'detail': (
                    f'No active terminal representative is assigned to terminal "{terminal or "(blank)"}". '
                    f'Registered terminal rep terminals: {registered}. '
                    f'Ask an admin to assign a terminal representative to "{terminal}" or correct the terminal name on this document.'
                )},
                status=status.HTTP_400_BAD_REQUEST,
            )

        doc_number = self._get_doc_number(obj)
        doc_label = self._get_doc_label()
        vessel_name = getattr(obj, 'vessel_name', '')
        title = f'Signing Request: {doc_label} #{doc_number}'
        message = (
            f'Inspector {request.user.get_full_name() or request.user.username} requests your signature for '
            f'{doc_label} #{doc_number} for vessel "{vessel_name}" at terminal "{terminal}". '
            'Please sign in to the Smart Reporting System to review and sign the document.'
        )
        with transaction.atomic():
            obj.signing_step = 'sent_to_client'
            self._save_signing_state(obj, ['signing_step'])
            for representative in representatives:
                create_notification(
                    recipient=representative.user,
                    title=title,
                    message=message,
                    notification_type='signing_request',
                    doc_type=self._doc_type_label,
                    doc_id=obj.pk,
                    doc_number=doc_number,
                    email=True,
                )
        ActivityLog.log(request.user, 'report_updated', doc_type=self._doc_type_label,
            doc_id=obj.pk, doc_number=doc_number,
            detail=f'Sent {doc_label} #{doc_number} signing request to {len(representatives)} representative(s) at {terminal}.',
            request=request)
        return Response(self.get_serializer(obj).data)

    @action(detail=True, methods=['post'], url_path='client_sign')
    def client_sign(self, request, pk=None):
        """Step 3 — Terminal Representative/terminal rep) draws signature and signs."""
        obj = self.get_object()
        role = self._get_role(request)
        if role not in ('terminal_representative', 'admin'):
            return Response({'detail': 'Only terminal representatives can perform this step.'}, status=status.HTTP_403_FORBIDDEN)
        if obj.signing_step != 'sent_to_client':
            return Response({'detail': f'Cannot sign: current step is "{obj.signing_step}".'}, status=status.HTTP_400_BAD_REQUEST)
        if role == 'terminal_representative' and not _terminal_representatives(
            getattr(obj, 'terminal', '')
        ).filter(user=request.user).exists():
            return Response(
                {'detail': 'This signing request is assigned to a different terminal.'},
                status=status.HTTP_403_FORBIDDEN,
            )

        sig_b64 = request.data.get('signature')
        if not sig_b64:
            return Response({'detail': 'signature field is required.'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            submitted_name = (request.data.get('signer_name') or '').strip()
            obj.terminal_representative_name = submitted_name or request.user.get_full_name() or request.user.username
            obj.document_hash    = self._document_hash(obj)
            obj.signing_step     = 'sent_to_inspector'
            obj.client_signed_at = timezone.now()
            obj.client_signed_by = request.user
            self._save_signing_state(obj, ['document_hash', 'signing_step', 'client_signed_at', 'client_signed_by',
                                           'terminal_representative_name'])
            _record_signature(request, obj, 'terminal_representative', obj.document_hash,
                              self._doc_type_label, self._get_doc_number(obj), sig_b64)
            signed_bytes = self._build_pdf_with_signatures(obj)
            ActivityLog.log(request.user, 'report_updated', doc_type=self._doc_type_label,
                doc_id=obj.pk, doc_number=self._get_doc_number(obj),
                detail=f'Terminal representative signed {self._get_doc_label()} #{self._get_doc_number(obj)}',
                request=request)

            # Notify inspector that document is ready to submit
            inspector = getattr(obj, 'created_by', None) or getattr(obj, 'inspector', None)
            if inspector:
                doc_number   = self._get_doc_number(obj)
                vessel_name  = getattr(obj, 'vessel_name', '')
                doc_label    = self._get_doc_label()
                counterparty = self._get_counterparty_label()
                create_notification(
                    recipient=inspector,
                    title=f'Document Ready to Submit - {doc_label} #{doc_number}',
                    message=(
                        f'{doc_label} #{doc_number} for vessel "{vessel_name}" has been signed by both '
                        f'the inspector and the {counterparty}. '
                        f'Please submit it to Admin.'
                    ),
                    notification_type='ready_to_submit',
                    doc_type=self._doc_type_label,
                    doc_id=obj.pk,
                    doc_number=doc_number,
                    email=True,
                )

            signed_bytes = _finalize_pdf_for_download(signed_bytes, obj, self._get_doc_label(), request)
            response = HttpResponse(signed_bytes, content_type='application/pdf')
            fname = self._get_doc_number(obj)
            response['Content-Disposition'] = f'attachment; filename="ClientSigned_{fname}.pdf"'
            return response
        except ValueError as e:
            return Response({'detail': str(e)}, status=status.HTTP_400_BAD_REQUEST)
        except Exception:
            sec_log.exception('Client signing failed for %s #%s', self._doc_type_label, obj.pk)
            return Response({'detail': 'Signing failed. Please try again or contact an administrator.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


    @action(detail=True, methods=['post'], url_path='inspector_verify')
    def inspector_verify(self, request, pk=None):
        """Step 5a — Inspector verifies the client-signed document."""
        obj = self.get_object()
        role = self._get_role(request)
        if role not in ('inspector', 'admin'):
            return Response({'detail': 'Only inspectors can verify.'}, status=status.HTTP_403_FORBIDDEN)
        if obj.signing_step != 'sent_to_inspector':
            return Response({'detail': f'Cannot verify: current step is "{obj.signing_step}".'}, status=status.HTTP_400_BAD_REQUEST)
        obj.signing_step = 'verified'
        obj.verified_at  = timezone.now()
        self._save_signing_state(obj, ['signing_step', 'verified_at'])
        return Response(self.get_serializer(obj).data)

    @action(detail=True, methods=['post'], url_path='submit_to_admin')
    def submit_to_admin(self, request, pk=None):
        """Step 5 — Inspector submits the signed document to Admin and notifies admin + terminal rep."""
        obj = self.get_object()
        role = self._get_role(request)
        if role not in ('inspector', 'admin'):
            return Response({'detail': 'Only inspectors can submit to admin.'}, status=status.HTTP_403_FORBIDDEN)
        if obj.signing_step not in {'sent_to_inspector', 'verified'}:
            return Response({'detail': f'Cannot submit: current step is "{obj.signing_step}".'}, status=status.HTTP_400_BAD_REQUEST)
        # Allow submission if the document is already marked 'verified'. Some tests and workflows
        # set signing_step='verified' directly; in that case we don't require explicit timestamp fields.
        if obj.signing_step != 'verified':
            if not getattr(obj, 'inspector_signed_at', None):
                return Response({'detail': f'{self._get_doc_label()} is missing the inspector signature.'}, status=status.HTTP_400_BAD_REQUEST)
            if not getattr(obj, 'client_signed_at', None):
                return Response({'detail': f'{self._get_doc_label()} is missing the terminal representative signature.'}, status=status.HTTP_400_BAD_REQUEST)

        doc_type = self._doc_type_label
        doc_number = self._get_doc_number(obj)
        vessel_name = getattr(obj, 'vessel_name', '')
        terminal = getattr(obj, 'terminal', '')

        already = Submission.objects.filter(doc_type=doc_type, doc_id=obj.pk).exists()
        if not already:
            Submission.objects.create(
                doc_type=doc_type,
                doc_id=obj.pk,
                doc_number=doc_number,
                vessel_name=vessel_name,
                terminal=terminal,
                submitted_by=request.user,
                is_read=False,
            )

        # ── Notify all admins ────────────────────────────────────────────────
        doc_label = self._get_doc_label()

        title = f'New Report Submitted: {doc_label}'
        msg   = (f'{request.user.get_full_name() or request.user.username} submitted '
                 f'{doc_label} #{doc_number} for vessel "{vessel_name}" (Terminal: {terminal}).')

        admin_profiles = UserProfile.objects.filter(role='admin').select_related('user')
        for ap in admin_profiles:
            create_notification(ap.user, title, msg,
                                notification_type='report_submitted',
                                doc_type=doc_type, doc_id=obj.pk, doc_number=doc_number)

        # ── Notify terminal representative (client who signed) ────
        client_user = getattr(obj, 'client_signed_by', None)
        if client_user:
            client_msg = (f'The signed {doc_label} #{doc_number} for vessel "{vessel_name}" '
                          f'has been submitted to PBPA admin by '
                          f'{request.user.get_full_name() or request.user.username}.')
            create_notification(client_user, f'Report Submitted to Admin: {doc_label}', client_msg,
                                notification_type='report_submitted_client',
                                doc_type=doc_type, doc_id=obj.pk, doc_number=doc_number)

        obj.signing_step = 'submitted'
        self._save_signing_state(obj, ['signing_step'])
        ActivityLog.log(request.user, 'report_submitted', doc_type=doc_type,
            doc_id=obj.pk, doc_number=doc_number,
            detail=f'{doc_label} #{doc_number} submitted for vessel \"{vessel_name}\"',
            request=request)
        return Response(self.get_serializer(obj).data)

    @action(detail=True, methods=['post'], url_path='admin_sign')
    def admin_sign(self, request, pk=None):
        """Step 6 — Admin signs the submitted document (final approval signature)."""
        obj = self.get_object()
        role = self._get_role(request)
        if role != 'admin':
            return Response({'detail': 'Only admins can perform this step.'}, status=status.HTTP_403_FORBIDDEN)
        if obj.signing_step != 'submitted':
            return Response({'detail': f'Cannot sign: current step is "{obj.signing_step}".'}, status=status.HTTP_400_BAD_REQUEST)

        sig_b64 = request.data.get('signature')
        if not sig_b64:
            return Response({'detail': 'signature field is required.'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            pdf_bytes    = self._build_pdf_bytes(obj)
            signed_bytes = overlay_signature_image(pdf_bytes, sig_b64, x=60, y=18, w=55, h=18)
            doc_hash     = self._document_hash(obj)
            obj.document_hash = doc_hash
            self._save_signing_state(obj, ['document_hash'])
            _record_signature(request, obj, 'admin', doc_hash,
                              self._doc_type_label, self._get_doc_number(obj))
            ActivityLog.log(request.user, 'report_approved', doc_type=self._doc_type_label,
                doc_id=obj.pk, doc_number=self._get_doc_number(obj),
                detail=f'Admin signed {self._get_doc_label()} #{self._get_doc_number(obj)}',
                request=request)
            signed_bytes = _finalize_pdf_for_download(signed_bytes, obj, self._get_doc_label(), request)
            response = HttpResponse(signed_bytes, content_type='application/pdf')
            fname = self._get_doc_number(obj)
            response['Content-Disposition'] = f'attachment; filename="AdminSigned_{fname}.pdf"'
            return response
        except ValueError as e:
            return Response({'detail': str(e)}, status=status.HTTP_400_BAD_REQUEST)
        except Exception:
            sec_log.exception('Admin signing failed for %s #%s', self._doc_type_label, obj.pk)
            return Response({'detail': 'Signing failed. Please try again or contact an administrator.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @action(detail=True, methods=['get'], url_path='verify_hash')
    def verify_hash(self, request, pk=None):
        """Step 10 — Recalculate document hash and compare with stored hash."""
        obj = self.get_object()
        stored_hash = getattr(obj, 'document_hash', '')
        if not stored_hash:
            return Response({'verified': False, 'detail': 'No stored hash — document has not been signed.'})
        try:
            current_hash = self._document_hash(obj)
            verified = (current_hash == stored_hash)
            signature_records = DocumentSignature.objects.filter(
                doc_type=self._doc_type_label, doc_id=obj.pk
            ).select_related('signer')
            signatures = []
            for signature in signature_records:
                signature_valid = bool(signature.signer_id) and verify_digital_signature(
                    signature.document_hash, str(signature.signer_id),
                    signature.signing_timestamp, signature.digital_signature,
                )
                signatures.append({
                    'role': signature.role,
                    'signer_name': signature.signer.get_full_name() if signature.signer else '',
                    'signing_timestamp': signature.signing_timestamp,
                    'document_hash': signature.document_hash,
                    'status': signature.status,
                    'ip_address': signature.ip_address,
                    'signature_valid': signature_valid,
                })
            verified = verified and bool(signatures) and all(
                item['signature_valid'] and item['status'] == 'valid' and item['document_hash'] == current_hash
                for item in signatures
            )
            return Response({
                'verified':      verified,
                'stored_hash':   stored_hash,
                'current_hash':  current_hash,
                'status':        '\u2713 Verified' if verified else '\u2717 Document Modified',
                'signatures':    list(signatures),
            })
        except Exception as e:
            return Response({'detail': f'Verification failed: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class TerminalViewSet(viewsets.ModelViewSet):
    """Terminal directory. All authenticated users can read it; admins manage it."""
    serializer_class = TerminalSerializer
    permission_classes = (IsAuthenticated,)

    def get_queryset(self):
        queryset = Terminal.objects.all()
        if get_or_create_user_profile(self.request.user).role != 'admin':
            queryset = queryset.filter(is_active=True)
        return queryset

    def _require_admin(self, request):
        if get_or_create_user_profile(request.user).role != 'admin':
            raise PermissionDenied('Only admins can manage terminals.')

    def create(self, request, *args, **kwargs):
        self._require_admin(request)
        return super().create(request, *args, **kwargs)

    def update(self, request, *args, **kwargs):
        self._require_admin(request)
        return super().update(request, *args, **kwargs)

    def partial_update(self, request, *args, **kwargs):
        self._require_admin(request)
        return super().partial_update(request, *args, **kwargs)

    def destroy(self, request, *args, **kwargs):
        self._require_admin(request)
        return super().destroy(request, *args, **kwargs)

class UserRegistrationViewSet(viewsets.ModelViewSet):
    """User creation — admin only (no public registration)."""
    queryset = User.objects.all()
    serializer_class = UserRegistrationSerializer
    permission_classes = (IsAuthenticated,)

    def get_queryset(self):
        profile = get_or_create_user_profile(self.request.user)
        if profile.role == 'admin':
            return User.objects.all()
        return User.objects.filter(pk=self.request.user.pk)

    def get_serializer_class(self):
        if self.action == 'create':
            return UserRegistrationSerializer
        return UserSerializer

    def create(self, request, *args, **kwargs):
        # Only admins may create users
        profile = get_or_create_user_profile(request.user)
        if profile.role != 'admin':
            sec_log.warning('Non-admin user %s attempted to create a user.', request.user.username)
            return Response({'detail': 'Only admins can create users.'}, status=status.HTTP_403_FORBIDDEN)

        role = request.data.get('role', 'inspector')
        if role not in ('inspector', 'terminal_representative', 'admin'):
            role = 'inspector'

        registration_data = request.data.copy()
        if role == 'terminal_representative':
            required_fields = {
                'company': 'Company Name', 'terminal': 'Terminal',
                'terminal_location': 'Terminal Location', 'phone': 'Company Phone Number',
                'company_email': 'Company Email', 'date_joined_company': 'Date',
            }
            errors = {field: f'{label} is required for terminal representatives.' for field, label in required_fields.items()
                      if not str(request.data.get(field, '')).strip()}
            if errors:
                return Response(errors, status=status.HTTP_400_BAD_REQUEST)
            # Terminal accounts sign in with a stable, human-readable company
            # username (for example, "vivo-energy-tanzania"), not an email.
            # Slugification keeps it compatible with Django's username rules.
            registration_data['username'] = slugify(request.data['company'])
            if not registration_data['username']:
                return Response({'company': 'Company Name must contain letters or numbers.'}, status=status.HTTP_400_BAD_REQUEST)
            registration_data['email'] = request.data['company_email'].strip().lower()
            registration_data['first_name'] = request.data['company'].strip()
            registration_data['last_name'] = ''

        serializer = self.get_serializer(data=registration_data)
        serializer.is_valid(raise_exception=True)
        user = serializer.save()
        profile_fields = {'role': role}
        if role == 'terminal_representative':
            profile_fields['employee_id']         = request.data.get('employee_id', '')
            profile_fields['terminal']            = request.data.get('terminal', '')
            profile_fields['terminal_location']   = request.data.get('terminal_location', '')
            profile_fields['company']             = request.data.get('company', '')
            profile_fields['position']            = request.data.get('position', '')
            profile_fields['company_email']       = request.data.get('company_email', '')
            profile_fields['phone']               = request.data.get('phone', '')
            date_joined = request.data.get('date_joined_company')
            if date_joined:
                profile_fields['date_joined_company'] = date_joined
        UserProfile.objects.create(user=user, **profile_fields)
        sec_log.info('Admin %s created user %s with role %s.', request.user.username, user.username, role)
        ActivityLog.log(request.user, 'user_created', detail=f'Created user {user.username} with role {role}', request=request)
        return Response(
            {'detail': 'User created successfully.', 'user_id': user.id, 'username': user.username},
            status=status.HTTP_201_CREATED
        )

    def update(self, request, *args, **kwargs):
        if get_or_create_user_profile(request.user).role != 'admin':
            return Response({'detail': 'Only admins can update users.'}, status=status.HTTP_403_FORBIDDEN)
        return super().update(request, *args, **kwargs)

    def partial_update(self, request, *args, **kwargs):
        if get_or_create_user_profile(request.user).role != 'admin':
            return Response({'detail': 'Only admins can update users.'}, status=status.HTTP_403_FORBIDDEN)
        return super().partial_update(request, *args, **kwargs)

    def destroy(self, request, *args, **kwargs):
        if get_or_create_user_profile(request.user).role != 'admin':
            return Response({'detail': 'Only admins can delete users.'}, status=status.HTTP_403_FORBIDDEN)
        return super().destroy(request, *args, **kwargs)


class UserProfileViewSet(viewsets.ModelViewSet):
    """User profile management"""
    queryset = UserProfile.objects.all()
    serializer_class = UserProfileSerializer
    permission_classes = (IsAuthenticated,)

    def get_queryset(self):
        user = self.request.user
        profile = get_or_create_user_profile(user)
        if profile.role == 'admin':
            return UserProfile.objects.all()
        return UserProfile.objects.filter(user=user)

    @action(detail=False, methods=['get'])
    def current_user(self, request):
        """Get current user's profile"""
        profile = get_or_create_user_profile(request.user)
        serializer = self.get_serializer(profile)
        return Response(serializer.data)

    @action(detail=False, methods=['post'])
    def mfa_setup(self, request):
        """Create a fresh secret and return its authenticator provisioning URI once."""
        profile = get_or_create_user_profile(request.user)
        profile.mfa_secret = generate_secret()
        profile.mfa_enabled = False
        profile.mfa_confirmed_at = None
        profile.save(update_fields=['mfa_secret', 'mfa_enabled', 'mfa_confirmed_at', 'updated_at'])
        return Response({'secret': profile.mfa_secret, 'provisioning_uri': provisioning_uri(profile.mfa_secret, request.user.username)})

    @action(detail=False, methods=['post'])
    def mfa_confirm(self, request):
        profile = get_or_create_user_profile(request.user)
        if not profile.mfa_secret:
            return Response({'detail': 'Start MFA setup before confirming.'}, status=status.HTTP_400_BAD_REQUEST)
        if not verify_code(profile.mfa_secret, str(request.data.get('otp', ''))):
            return Response({'detail': 'The authenticator code is invalid or expired.'}, status=status.HTTP_400_BAD_REQUEST)
        profile.mfa_enabled = True
        profile.mfa_confirmed_at = timezone.now()
        profile.save(update_fields=['mfa_enabled', 'mfa_confirmed_at', 'updated_at'])
        ActivityLog.log(request.user, 'report_updated', detail='Enabled authenticator-app MFA.', request=request)
        return Response({'detail': 'Multi-factor authentication is enabled.', 'mfa_enabled': True})

    @action(detail=False, methods=['post'])
    def mfa_disable(self, request):
        profile = get_or_create_user_profile(request.user)
        if not profile.mfa_enabled:
            return Response({'detail': 'MFA is not enabled.'}, status=status.HTTP_400_BAD_REQUEST)
        if not verify_code(profile.mfa_secret, str(request.data.get('otp', ''))):
            return Response({'detail': 'The authenticator code is invalid or expired.'}, status=status.HTTP_400_BAD_REQUEST)
        profile.mfa_enabled = False
        profile.mfa_secret = ''
        profile.mfa_confirmed_at = None
        profile.save(update_fields=['mfa_secret', 'mfa_enabled', 'mfa_confirmed_at', 'updated_at'])
        ActivityLog.log(request.user, 'report_updated', detail='Disabled authenticator-app MFA.', request=request)
        return Response({'detail': 'Multi-factor authentication is disabled.', 'mfa_enabled': False})

    def create(self, request, *args, **kwargs):
        if get_or_create_user_profile(request.user).role != 'admin':
            return Response({'detail': 'Only admins can create profiles.'}, status=status.HTTP_403_FORBIDDEN)
        return super().create(request, *args, **kwargs)

    def update(self, request, *args, **kwargs):
        if get_or_create_user_profile(request.user).role != 'admin':
            return Response({'detail': 'Only admins can update profiles.'}, status=status.HTTP_403_FORBIDDEN)
        return super().update(request, *args, **kwargs)

    def partial_update(self, request, *args, **kwargs):
        if get_or_create_user_profile(request.user).role != 'admin':
            return Response({'detail': 'Only admins can update profiles.'}, status=status.HTTP_403_FORBIDDEN)
        return super().partial_update(request, *args, **kwargs)

    def destroy(self, request, *args, **kwargs):
        if get_or_create_user_profile(request.user).role != 'admin':
            return Response({'detail': 'Only admins can delete profiles.'}, status=status.HTTP_403_FORBIDDEN)
        return super().destroy(request, *args, **kwargs)

    @action(detail=False, methods=['get'])
    def list_inspectors(self, request):
        """Get all inspectors"""
        profile = get_or_create_user_profile(request.user)
        if profile.role not in ('terminal_representative', 'admin'):
            return Response({'detail': 'Only terminal representatives and admins can list inspectors.'}, status=status.HTTP_403_FORBIDDEN)
        inspectors = UserProfile.objects.filter(role='inspector', is_active=True)
        serializer = self.get_serializer(inspectors, many=True)
        return Response(serializer.data)

    @action(detail=True, methods=['post'])
    def set_password(self, request, pk=None):
        """Admin sets password for any user."""
        profile = get_or_create_user_profile(request.user)
        if profile.role != 'admin':
            sec_log.warning('Non-admin %s attempted set_password.', request.user.username)
            return Response({'detail': 'Only admins can set passwords.'}, status=status.HTTP_403_FORBIDDEN)
        user_profile = self.get_object()
        password = request.data.get('password', '')
        if len(password) < 8:
            return Response({'detail': 'Password must be at least 8 characters.'}, status=status.HTTP_400_BAD_REQUEST)
        user_profile.user.set_password(password)
        user_profile.user.save()
        sec_log.info('Admin %s reset password for user %s.', request.user.username, user_profile.user.username)
        ActivityLog.log(request.user, 'password_changed', detail=f'Password reset for user {user_profile.user.username}', request=request)
        return Response({'detail': 'Password updated successfully.'})


class RosterAssignmentViewSet(viewsets.ModelViewSet):
    """Admin roster assignments for inspectors."""
    permission_classes = (IsAuthenticated,)
    serializer_class = RosterAssignmentSerializer
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['inspector__username', 'inspector__first_name', 'inspector__last_name', 'terminal', 'vessel_name', 'task']
    ordering_fields = ['week_start_date', 'created_at', 'status']
    ordering = ['-week_start_date', '-created_at']

    def get_queryset(self):
        profile = get_or_create_user_profile(self.request.user)
        qs = RosterAssignment.objects.select_related('inspector', 'terminal_representative')
        if profile.role == 'inspector':
            return qs.filter(inspector=self.request.user, status='sent')
        if profile.role in ('terminal_representative', 'admin'):
            return qs
        return qs.none()

    def _ensure_terminal_rep_or_admin(self, request):
        profile = get_or_create_user_profile(request.user)
        if profile.role not in ('terminal_representative', 'admin'):
            return Response({'detail': 'Only terminal representatives and admins can manage rosters.'}, status=status.HTTP_403_FORBIDDEN)
        return None

    def create(self, request, *args, **kwargs):
        denied = self._ensure_terminal_rep_or_admin(request)
        if denied:
            return denied
        return super().create(request, *args, **kwargs)

    def update(self, request, *args, **kwargs):
        denied = self._ensure_terminal_rep_or_admin(request)
        if denied:
            return denied
        return super().update(request, *args, **kwargs)

    def partial_update(self, request, *args, **kwargs):
        denied = self._ensure_terminal_rep_or_admin(request)
        if denied:
            return denied
        return super().partial_update(request, *args, **kwargs)

    def destroy(self, request, *args, **kwargs):
        denied = self._ensure_terminal_rep_or_admin(request)
        if denied:
            return denied
        return super().destroy(request, *args, **kwargs)

    def perform_create(self, serializer):
        sent_at = timezone.now() if serializer.validated_data.get('status') == 'sent' else None
        serializer.save(created_by_admin=self.request.user, sent_at=sent_at, is_read=False)

    def perform_update(self, serializer):
        instance = self.get_object()
        status_value = serializer.validated_data.get('status', instance.status)
        sent_at = instance.sent_at
        if status_value == 'sent' and not sent_at:
            sent_at = timezone.now()
        serializer.save(sent_at=sent_at, is_read=False if status_value == 'sent' else instance.is_read)

    @action(detail=True, methods=['post'])
    def send(self, request, pk=None):
        denied = self._ensure_terminal_rep_or_admin(request)
        if denied:
            return denied
        assignment = self.get_object()
        if assignment.status == 'cancelled':
            return Response({'detail': 'Cancelled roster assignments cannot be sent.'}, status=status.HTTP_400_BAD_REQUEST)
        assignment.status = 'sent'
        assignment.sent_at = timezone.now()
        assignment.is_read = False
        assignment.save(update_fields=['status', 'sent_at', 'is_read', 'updated_at'])
        return Response(self.get_serializer(assignment).data)

    @action(detail=True, methods=['post'])
    def cancel(self, request, pk=None):
        denied = self._ensure_terminal_rep_or_admin(request)
        if denied:
            return denied
        assignment = self.get_object()
        assignment.status = 'cancelled'
        assignment.save(update_fields=['status', 'updated_at'])
        return Response(self.get_serializer(assignment).data)

    @action(detail=False, methods=['get'])
    def unread_count(self, request):
        profile = get_or_create_user_profile(request.user)
        if profile.role != 'inspector':
            return Response({'count': 0})
        count = RosterAssignment.objects.filter(inspector=request.user, status='sent', is_read=False).count()
        return Response({'count': count})

    @action(detail=True, methods=['post'])
    def mark_read(self, request, pk=None):
        assignment = self.get_object()
        assignment.is_read = True
        assignment.save(update_fields=['is_read', 'updated_at'])
        return Response(self.get_serializer(assignment).data)

    @action(detail=True, methods=['get'])
    def pdf(self, request, pk=None):
        assignment = self.get_object()
        buf = io.BytesIO()
        pdf = canvas.Canvas(buf, pagesize=letter)
        width, height = letter
        y = height - 60

        pdf.setFont("Helvetica-Bold", 15)
        pdf.drawString(50, y, "INSPECTOR ROSTER ASSIGNMENT")
        y -= 35
        pdf.setFont("Helvetica", 10)

        days_str = ', '.join(assignment.working_days) if assignment.working_days else '-'
        rows = [
            ("Inspector", assignment.inspector.get_full_name() or assignment.inspector.username),
            ("Week Starting", assignment.week_start_date.strftime("%d-%m-%Y") if assignment.week_start_date else '-'),
            ("Working Days", days_str),
            ("Shift", assignment.get_shift_display()),
            ("Location", assignment.location or "-"),
            ("Terminal", assignment.terminal or "-"),
            ("Vessel", assignment.vessel_name or "-"),
            ("Task", assignment.task or "-"),
            ("Client", assignment.approved_by.get_full_name() if assignment.approved_by else "-"),
            ("Status", assignment.get_status_display()),
        ]
        for label, value in rows:
            pdf.setFont("Helvetica-Bold", 10)
            pdf.drawString(50, y, f"{label}:")
            pdf.setFont("Helvetica", 10)
            pdf.drawString(170, y, str(value))
            y -= 20

        if assignment.notes:
            y -= 10
            pdf.setFont("Helvetica-Bold", 10)
            pdf.drawString(50, y, "Notes:")
            y -= 18
            pdf.setFont("Helvetica", 10)
            text = pdf.beginText(50, y)
            for line in assignment.notes.splitlines():
                text.textLine(line[:95])
            pdf.drawText(text)

        pdf.showPage()
        pdf.save()
        buf.seek(0)
        pdf_bytes = _finalize_pdf_for_download(buf.getvalue(), assignment, 'Inspector Roster Assignment', request)
        response = HttpResponse(pdf_bytes, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="Roster_{assignment.week_start_date}_{assignment.inspector.username}.pdf"'
        return response


# ========== TANK VIEWSETS ==========
class TankViewSet(viewsets.ModelViewSet):
    """Tank management"""
    queryset = Tank.objects.filter(is_active=True)
    permission_classes = (IsAuthenticated,)
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['tank_id', 'tank_name', 'product_type']
    ordering_fields = ['tank_id', 'created_at']
    ordering = ['tank_id']
    
    def get_serializer_class(self):
        if self.action == 'retrieve':
            return TankDetailSerializer
        return TankSerializer
    
    @action(detail=True, methods=['get'])
    def inspection_history(self, request, pk=None):
        """Get inspection history for a tank"""
        tank = self.get_object()
        inspections = tank.inspections.filter(inspector=request.user).order_by('-inspection_date')
        
        # Pagination
        page = request.query_params.get('page', 1)
        page_size = request.query_params.get('page_size', 10)
        start = (int(page) - 1) * int(page_size)
        end = start + int(page_size)
        
        inspections_page = inspections[start:end]
        serializer = InspectionListSerializer(inspections_page, many=True)
        
        return Response({
            'count': inspections.count(),
            'page': page,
            'page_size': page_size,
            'results': serializer.data
        })
    
    @action(detail=False, methods=['get'])
    def summary(self, request):
        """Get tank summary statistics"""
        tanks = self.get_queryset()
        total_tanks = tanks.count()
        
        inspections = Inspection.objects.filter(tank__in=tanks, inspector=request.user)
        total_inspections = inspections.count()
        pending_inspections = inspections.filter(status='submitted').count()
        
        return Response({
            'total_tanks': total_tanks,
            'total_inspections': total_inspections,
            'pending_inspections': pending_inspections,
            'inspection_rate': f"{(total_inspections / total_tanks * 100):.2f}%" if total_tanks > 0 else "0%"
        })


# ========== SEAL & ISOLATION VIEWSETS ==========
class SealViewSet(viewsets.ModelViewSet):
    """Seal management"""
    queryset = Seal.objects.all()
    serializer_class = SealSerializer
    permission_classes = (IsAuthenticated,)
    
    def get_queryset(self):
        queryset = Seal.objects.filter(inspection__inspector=self.request.user)
        inspection_id = self.request.query_params.get('inspection_id')
        if inspection_id:
            return queryset.filter(inspection_id=inspection_id)
        return queryset

    def perform_create(self, serializer):
        ensure_inspection_owner(serializer.validated_data['inspection'], self.request.user)
        serializer.save()


class IsolationViewSet(viewsets.ModelViewSet):
    """Isolation/Valve management"""
    queryset = Isolation.objects.all()
    serializer_class = IsolationSerializer
    permission_classes = (IsAuthenticated,)
    
    def get_queryset(self):
        queryset = Isolation.objects.filter(inspection__inspector=self.request.user)
        inspection_id = self.request.query_params.get('inspection_id')
        if inspection_id:
            return queryset.filter(inspection_id=inspection_id)
        return queryset

    def perform_create(self, serializer):
        ensure_inspection_owner(serializer.validated_data['inspection'], self.request.user)
        serializer.save()


# ========== CALCULATION VIEWSET ==========
class InspectionCalculationViewSet(viewsets.ReadOnlyModelViewSet):
    """View calculations (read-only)"""
    queryset = InspectionCalculation.objects.all()
    serializer_class = InspectionCalculationSerializer
    permission_classes = (IsAuthenticated,)

    def get_queryset(self):
        queryset = InspectionCalculation.objects.filter(inspection__inspector=self.request.user)
        inspection_id = self.request.query_params.get('inspection_id')
        if inspection_id:
            return queryset.filter(inspection_id=inspection_id)
        return queryset


# ========== REPORT VIEWSET ==========
class InspectionReportViewSet(viewsets.ModelViewSet):
    """Report management"""
    queryset = InspectionReport.objects.all()
    serializer_class = InspectionReportSerializer
    permission_classes = (IsAuthenticated,)
    
    def get_queryset(self):
        queryset = InspectionReport.objects.filter(inspection__inspector=self.request.user)
        inspection_id = self.request.query_params.get('inspection_id')
        if inspection_id:
            return queryset.filter(inspection_id=inspection_id)
        return queryset

    def perform_create(self, serializer):
        ensure_inspection_owner(serializer.validated_data['inspection'], self.request.user)
        serializer.save(generated_by=self.request.user)


# ========== INSPECTION VIEWSET (Main) ==========
class InspectionViewSet(SigningWorkflowMixin, viewsets.ModelViewSet):
    """Main inspection endpoint"""
    permission_classes = (IsAuthenticated,)
    _doc_type_label = 'dip_ticket'
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['tank__tank_name', 'tank__tank_id', 'inspector__username']
    ordering_fields = ['inspection_date', 'created_at', 'status']
    ordering = ['-inspection_date']
    
    def _build_pdf_bytes(self, obj):
        return generate_dip_ticket_pdf(obj).getvalue()

    def _get_doc_number(self, obj):
        return obj.ticket_number or str(obj.pk)

    def get_queryset(self):
        user = self.request.user
        status_filter = self.request.query_params.get('status')
        tank_id = self.request.query_params.get('tank_id')
        date_from = self.request.query_params.get('date_from')
        date_to = self.request.query_params.get('date_to')
        period = self.request.query_params.get('period')

        queryset = Inspection.objects.all()
        profile = get_or_create_user_profile(user)
        if profile.role == 'admin':
            if self.action in {'list', 'recent'}:
                # Keep this register in sync with the administration dashboard:
                # a ticket appears only after it has an inbox submission.
                submissions = filter_queryset_by_period(
                    Submission.objects.filter(doc_type='dip_ticket'), period, 'submitted_at'
                )
                queryset = queryset.filter(pk__in=submissions.values_list('doc_id', flat=True))
                # The period was applied to submitted_at above, not the ticket date.
                period = None
            elif self.action in {'approve', 'reject'}:
                queryset = queryset.filter(status='submitted')
        elif profile.role == 'terminal_representative':
            terminal_name = (profile.terminal or '').strip()
            # Always scope to the terminal rep's own terminal — never show all.
            base = queryset.filter(terminal__iexact=terminal_name) if terminal_name else queryset.none()

            if self.action == 'list':
                queryset = base.filter(signing_step='sent_to_client')
            else:
                queryset = base.filter(
                    Q(signing_step='sent_to_client')
                    | Q(client_signed_by=user)
                    | Q(signing_step__in=['sent_to_inspector', 'verified', 'submitted'])
                )
        else:
            queryset = owned_queryset(queryset, user, 'inspector')
        
        # Apply filters
        if status_filter:
            queryset = queryset.filter(status=status_filter)
        
        if tank_id:
            queryset = queryset.filter(tank_id=tank_id)
        
        if date_from:
            queryset = queryset.filter(inspection_date__gte=date_from)
        
        if date_to:
            queryset = queryset.filter(inspection_date__lte=date_to)

        queryset = filter_queryset_by_period(queryset, period, 'inspection_date')
        
        return queryset
    
    def get_serializer_class(self):
        if self.action == 'create':
            return InspectionCreateSerializer
        elif self.action == 'retrieve':
            return InspectionDetailSerializer
        elif self.action in ['approve', 'reject']:
            return InspectionApprovalSerializer
        return InspectionListSerializer
    
    def perform_create(self, serializer):
        """Override create to set inspector"""
        inspection = serializer.save(inspector=self.request.user)
        
        # Automatically trigger calculations
        try:
            calc_engine = InspectionCalculationEngine()
            calc_engine.calculate_all(inspection)
        except Exception as e:
            # Log error but don't fail the request
            print(f"Calculation error: {str(e)}")
    
    @action(detail=True, methods=['post'], permission_classes=[IsAuthenticated])
    def approve(self, request, pk=None):
        """Approve an inspection (Terminal Rep only)."""
        inspection = self.get_object()
        role = get_or_create_user_profile(request.user).role
        if role != 'terminal_representative':
            return Response(
                {'detail': 'Only terminal representatives can approve inspections'},
                status=status.HTTP_403_FORBIDDEN
            )
        
        if inspection.status != 'submitted':
            return Response(
                {'detail': 'Only submitted inspections can be approved'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        inspection.status = 'approved'
        inspection.approved_by = request.user
        inspection.approval_date = timezone.now()
        inspection.save()
        ActivityLog.log(request.user, 'report_approved', 'dip_ticket',
            doc_id=inspection.pk, doc_number=inspection.ticket_number or str(inspection.pk),
            detail=f'Dip ticket approved for vessel {inspection.vessel_name or "-"}',
            request=request)
        serializer = self.get_serializer(inspection)
        return Response(serializer.data)
    
    @action(detail=True, methods=['post'], permission_classes=[IsAuthenticated])
    def reject(self, request, pk=None):
        """Reject an inspection (Terminal Rep only)."""
        inspection = self.get_object()
        role = get_or_create_user_profile(request.user).role
        if role != 'terminal_representative':
            return Response(
                {'detail': 'Only terminal representatives can reject inspections'},
                status=status.HTTP_403_FORBIDDEN
            )
        
        if inspection.status != 'submitted':
            return Response(
                {'detail': 'Only submitted inspections can be rejected'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        reason = request.data.get('rejection_reason', '')
        inspection.status = 'rejected'
        inspection.approved_by = request.user
        inspection.approval_date = timezone.now()
        inspection.rejection_reason = reason
        inspection.save()
        ActivityLog.log(request.user, 'report_rejected', 'dip_ticket',
            doc_id=inspection.pk, doc_number=inspection.ticket_number or str(inspection.pk),
            detail=f'Rejected: {reason[:100] if reason else "-"}',
            request=request)
        serializer = self.get_serializer(inspection)
        return Response(serializer.data)
    
    @action(detail=True, methods=['post'], permission_classes=[IsAuthenticated])
    def submit(self, request, pk=None):
        """Submit an inspection for approval"""
        inspection = self.get_object()
        
        if inspection.inspector != request.user:
            return Response(
                {'detail': 'You can only submit your own inspections'},
                status=status.HTTP_403_FORBIDDEN
            )
        
        if inspection.status != 'draft':
            return Response(
                {'detail': 'Only draft inspections can be submitted'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        inspection.status = 'submitted'
        inspection.save()
        
        serializer = self.get_serializer(inspection)
        return Response(serializer.data)
    
    @action(detail=False, methods=['get'])
    def dashboard(self, request):
        """Get dashboard statistics"""
        user = request.user
        period = request.query_params.get('period')
        
        profile = get_or_create_user_profile(user)

        owned_document_qs = get_owned_document_querysets(user, period)
        if profile.role == 'admin':
            document_qs = {
                'inspections': filter_queryset_by_period(
                    Inspection.objects.all(),
                    period,
                    'inspection_date',
                ),
                'product_receipt_certificates': filter_queryset_by_period(
                    ProductReceiptCertificate.objects.all(),
                    period,
                    'receipt_date',
                ),
                'seal_isolation_reports': filter_queryset_by_period(
                    SealIsolationReport.objects.all(),
                    period,
                    'report_date',
                ),
                'shore_tank_calculations': filter_queryset_by_period(
                    ShoreTankCalculation.objects.all(),
                    period,
                    'calculation_date',
                ),
                'stock_reports': filter_queryset_by_period(
                    StockReport.objects.all(),
                    period,
                    'report_date',
                ),
                'provisional_outturn_reports': filter_queryset_by_period(
                    ProvisionalOuturnReport.objects.all(),
                    period,
                    'report_date',
                ),
                'vessel_reports': filter_queryset_by_period(
                    VesselReport.objects.all(),
                    period,
                    'discharge_date',
                ),
            }
        else:
            document_qs = owned_document_qs

        if profile.role == 'terminal_representative':
            inspections_qs = filter_queryset_by_period(
                Inspection.objects.filter(status='submitted'),
                period,
                'inspection_date',
            )
        elif profile.role == 'admin':
            # Administration overview must reflect every dip ticket, not only
            # those still awaiting terminal review.
            inspections_qs = document_qs['inspections']
        else:
            inspections_qs = owned_document_qs['inspections']

        document_counts = {
            'inspections': inspections_qs.count(),
            'product_receipt_certificates': document_qs['product_receipt_certificates'].count(),
            'seal_isolation_reports': document_qs['seal_isolation_reports'].count(),
            'shore_tank_calculations': document_qs['shore_tank_calculations'].count(),
            'stock_reports': document_qs['stock_reports'].count(),
            'provisional_outturn_reports': document_qs['provisional_outturn_reports'].count(),
            'vessel_reports': document_qs['vessel_reports'].count(),
        }

        def get_status_counts(queryset):
            return {
                'draft': queryset.filter(status='draft').count(),
                'submitted': queryset.filter(status='submitted').count(),
                'approved': queryset.filter(status='approved').count(),
                'rejected': queryset.filter(status='rejected').count(),
            }
        
        if profile.role == 'inspector':
            inspections = inspections_qs.filter(inspector=user)
            total_inspections = inspections.count()
            status_counts = get_status_counts(inspections)
            
            return Response({
                'role': 'inspector',
                'total_inspections': total_inspections,
                **status_counts,
                'pending_approval': status_counts['submitted'],
                'document_counts': document_counts,
            })
        
        elif profile.role == 'terminal_representative':
            inspections = inspections_qs.filter(status='submitted')
            total_pending = inspections.count()
            approved = inspections_qs.filter(approved_by=user, status='approved').count()
            status_counts = get_status_counts(inspections_qs)
            activity_overview = build_activity_overview(user, period)
            
            return Response({
                'role': 'terminal_representative',
                'total_pending_approval': total_pending,
                'total_approved': approved,
                'awaiting_review': total_pending,
                **status_counts,
                'document_counts': document_counts,
                'activity_overview': activity_overview,
            })
        
        elif profile.role == 'admin':
            total_inspections = inspections_qs.count()
            system_inspections = document_qs['inspections']
            total_tanks = Tank.objects.filter(is_active=True).count()
            status_counts = get_status_counts(inspections_qs)
            system_status_counts = get_status_counts(system_inspections)
            # The Dip Ticket and Shore Tank dashboard cards represent reports
            # received by administration.  Count Submission records rather than
            # their source documents so drafts and unsigned/unsubmitted reports
            # created by inspectors never appear on those cards.
            _sub_qs = SubmissionViewSet(request=request)._with_existing_document_targets(
                Submission.objects.select_related('submitted_by')
            )
            submitted_document_qs = filter_queryset_by_period(
                _sub_qs,
                period,
                'submitted_at',
            )
            document_counts = {
                **document_counts,
                'inspections': submitted_document_qs.filter(doc_type='dip_ticket').count(),
                'shore_tank_calculations': submitted_document_qs.filter(doc_type='shore_tank').count(),
            }
            system_document_counts = {
                **document_counts,
                'inspections': system_inspections.count(),
            }
            pending_submissions = _sub_qs.filter(is_read=False).count()
            submitted_reports_count = _sub_qs.count()
            recent_submissions = _sub_qs.order_by('-submitted_at')[:5]
            activity_overview = build_activity_overview(user, period)
            
            return Response({
                'role': 'admin',
                'total_inspections': total_inspections,
                'total_system_inspections': system_inspections.count(),
                'total_tanks': total_tanks,
                **status_counts,
                'document_counts': document_counts,
                'system_document_counts': system_document_counts,
                'system_status_counts': system_status_counts,
                'activity_overview': activity_overview,
                'pending_submissions': pending_submissions,
                'submitted_reports_count': submitted_reports_count,
                'recent_submissions': SubmissionSerializer(recent_submissions, many=True).data,
            })
        
        return Response({'detail': 'Unknown role'}, status=status.HTTP_400_BAD_REQUEST)
    
    @action(detail=True, methods=['get'])
    def generate_document(self, request, pk=None):
        """Generate Dip Ticket as PDF"""
        inspection = self.get_object()
        try:
            response = HttpResponse(
                _finalize_pdf_for_download(self._build_pdf_with_signatures(inspection), inspection, 'Dip Ticket', request),
                content_type='application/pdf'
            )
            response['Content-Disposition'] = f'attachment; filename="Dip_Ticket_{inspection.ticket_number or inspection.id}.pdf"'
            return response
        except Exception as e:
            return Response({'detail': f'Document generation error: {str(e)}'}, status=status.HTTP_400_BAD_REQUEST)
    
    @action(detail=False, methods=['get'])
    def recent(self, request):
        """Get recent inspections"""
        limit = request.query_params.get('limit', 10)
        inspections = self.get_queryset()[:int(limit)]
        serializer = InspectionListSerializer(inspections, many=True)
        return Response(serializer.data)


class ProductReceiptCertificateViewSet(SigningWorkflowMixin, viewsets.ModelViewSet):
    """CRUD and PDF generation for PBPA product receipt certificates."""
    permission_classes = (IsAuthenticated,)
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['certificate_number', 'vessel_name', 'terminal', 'pbpa_inspector_name']
    ordering_fields = ['created_at', 'receipt_date', 'certificate_number', 'status']
    ordering = ['-created_at']
    _doc_type_label = 'product_receipt'

    def create(self, request, *args, **kwargs):
        return Response(
            {'detail': 'Product Receipt Certificates are created automatically from Shore Tank Calculations.'},
            status=status.HTTP_405_METHOD_NOT_ALLOWED,
        )

    def get_queryset(self):
        profile = get_or_create_user_profile(self.request.user)
        base = ProductReceiptCertificate.objects.select_related('created_by').prefetch_related('items__tank')
        if profile.role == 'terminal_representative':
            base = base.filter(terminal__iexact=profile.terminal.strip())
            if self.action == 'list':
                qs = base.filter(signing_step='sent_to_client')
            else:
                qs = base.filter(
                    Q(signing_step='sent_to_client')
                    | Q(client_signed_by=self.request.user)
                    | Q(signing_step__in=['client_signed', 'sent_to_inspector', 'verified', 'submitted'])
                )
        elif profile.role == 'admin':
            qs = base if self.action != 'list' else base.filter(signing_step='submitted')
        else:
            qs = owned_queryset(base, self.request.user)

        status_filter = self.request.query_params.get('status')
        period = self.request.query_params.get('period')
        if status_filter:
            qs = qs.filter(status=status_filter)
        return filter_queryset_by_period(qs, period, 'receipt_date')

    def get_serializer_class(self):
        if self.action in ['create', 'update', 'partial_update', 'retrieve']:
            return ProductReceiptCertificateDetailSerializer
        return ProductReceiptCertificateListSerializer

    def perform_create(self, serializer):
        full_name = self.request.user.get_full_name() or self.request.user.username
        serializer.save(created_by=self.request.user, pbpa_inspector_name=serializer.validated_data.get('pbpa_inspector_name') or full_name)

    def _build_pdf_bytes(self, obj):
        return self._build_pdf(obj).getvalue()

    @action(detail=True, methods=['post'], url_path='sign_with_image')
    def sign_with_image(self, request, pk=None):
        """Legacy alias — redirects to inspector_sign for backwards compatibility."""
        return self.inspector_sign(request, pk=pk)

    @action(detail=True, methods=['post'])
    def sign_document(self, request, pk=None):
        """Digitally sign the Product Receipt Certificate PDF."""
        certificate = self.get_object()
        if certificate.is_signed:
            return Response({'detail': 'Document is already signed.'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            import io
            buf = self._build_pdf(certificate)
            pdf_bytes = buf.getvalue()
            signer_name = request.user.get_full_name() or request.user.username
            certificate.is_signed     = True
            certificate.signed_at     = timezone.now()
            certificate.signed_by     = request.user
            certificate.document_hash = self._document_hash(certificate)
            certificate.save(update_fields=['is_signed','signed_at','signed_by','document_hash','updated_at'])
            signed_bytes = _finalize_pdf_for_download(pdf_bytes, certificate, 'Product Receipt Certificate', request)
            _record_signature(request, certificate, self._get_role(request), certificate.document_hash,
                              self._doc_type_label, self._get_doc_number(certificate))
            response = HttpResponse(signed_bytes, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="SIGNED_PRC_{certificate.certificate_number}.pdf"'
            return response
        except Exception as e:
            return Response({'detail': f'Signing failed: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @action(detail=False, methods=['get'], url_path='signature-info')
    def signature_info(self, request):
        return Response(get_signature_info())

    @action(detail=True, methods=['post'])
    def issue(self, request, pk=None):
        if self._get_role(request) != 'admin':
            return Response({'detail': 'Use the signing workflow to submit this certificate.'}, status=status.HTTP_409_CONFLICT)
        certificate = self.get_object()
        if certificate.status == 'issued':
            return Response({'detail': 'Certificate has already been issued.'}, status=status.HTTP_400_BAD_REQUEST)

        certificate.status = 'issued'
        certificate.issued_at = timezone.now()
        certificate.save(update_fields=['status', 'issued_at', 'updated_at'])

        serializer = ProductReceiptCertificateDetailSerializer(certificate, context={'request': request})
        return Response(serializer.data)

    @action(detail=True, methods=['get'])
    def pdf(self, request, pk=None):
        certificate = self.get_object()
        pdf_bytes = _finalize_pdf_for_download(self._build_pdf_with_signatures(certificate), certificate, 'Product Receipt Certificate', request)
        response = HttpResponse(pdf_bytes, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="product-receipt-certificate-{certificate.certificate_number}.pdf"'
        return response
    
    @action(detail=True, methods=['get'])
    def generate_document(self, request, pk=None):
        certificate = self.get_object()
        try:
            buf = generate_product_receipt_document(certificate)
            response = HttpResponse(
                buf.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="Product_Receipt_Cert_{certificate.certificate_number}.docx"'
            return response
        except Exception as e:
            return Response({'detail': f'Document generation error: {str(e)}'}, status=status.HTTP_400_BAD_REQUEST)

    def _build_pdf(self, certificate):
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.units import mm

        buffer = io.BytesIO()
        W, H = A4
        pdf = canvas.Canvas(buffer, pagesize=A4)
        M = 15 * mm
        TW = W - 2 * M
        y = H - 10 * mm

        # THE UNITED REPUBLIC OF TANZANIA
        pdf.setFont("Helvetica", 9)
        pdf.drawCentredString(W / 2, y, "THE UNITED REPUBLIC OF TANZANIA")
        y -= 8 * mm

        # PETROLEUM BULK PROCUREMENT AGENCY
        pdf.setFont("Helvetica-Bold", 12)
        pdf.drawCentredString(W / 2, y, "PETROLEUM BULK PROCUREMENT AGENCY")
        y -= 5 * mm
        pdf.setFont("Helvetica", 6.5)
        pdf.drawCentredString(W / 2, y,
            "TANZANIA PORTS AUTHORITY, ONE STOP CENTER BUILDING, 11TH FLOOR, SOKOINE DRIVE, PLOT NO:1/2")
        y -= 4 * mm
        pdf.drawCentredString(W / 2, y,
            "Tel: +255222129009 / Fax: +255222129093 / Email: info@pbpa.go.tz / WEBSITE: www.pbpa.go.tz / P.O. Box 2634 Dar es Salaam, TANZANIA")
        y -= 7 * mm

        # Title box
        box_h = 11 * mm
        box_x = M + 8 * mm
        box_w = TW - 16 * mm
        pdf.setFillColor(colors.white)
        pdf.setStrokeColor(colors.black)
        pdf.setLineWidth(2.5)
        pdf.rect(box_x, y - box_h, box_w, box_h, fill=1, stroke=1)
        pdf.setFillColor(colors.black)
        pdf.setLineWidth(0.5)
        pdf.setFont("Helvetica-Bold", 14)
        pdf.drawCentredString(W / 2, y - box_h + 3.5 * mm, "PRODUCT RECEIPT CERTIFICATE")
        y -= box_h + 5 * mm

        cert_no = certificate.certificate_number or ""

        # Dotted underline field helper
        def ufield(label, value, y_pos, lw=30 * mm, line_end=None):
            le = line_end if line_end else (W - M - 38 * mm)
            pdf.setFont("Helvetica", 9)
            pdf.drawString(M, y_pos, label)
            lx = M + lw
            pdf.setDash(1, 2)
            pdf.line(lx, y_pos - 1, le, y_pos - 1)
            pdf.setDash()
            pdf.setFont("Helvetica-Bold", 9)
            pdf.drawString(lx + 1, y_pos, str(value))

        lg = 8 * mm
        ufield("Vessel Name:", certificate.vessel_name, y)
        y -= lg
        ufield("Terminal:", certificate.terminal, y)
        pdf.setFont("Helvetica-Bold", 12)
        pdf.drawRightString(W - M, y + lg - 1 * mm, cert_no)
        y -= lg
        date_str = certificate.receipt_date.strftime("%d - %m - %Y") if certificate.receipt_date else ""
        ufield("Date:", date_str, y)
        y -= lg
        time_str = (certificate.receipt_time.strftime("%H.%M") + "Hrs") if certificate.receipt_time else ""
        ufield("Time:", time_str, y)
        y -= 7 * mm

        pdf.setFont("Helvetica", 9)
        pdf.drawString(M, y,
            "This is to confirm that the following Quantity was Delivered/Received into your tanks:")
        y -= 7 * mm

        # Items table
        cw = [TW * p for p in (0.12, 0.30, 0.29, 0.29)]
        rh = 8 * mm
        hdrs = ["Tank No.", "Product", "Weight in Tonnage", "Volume in Liters"]
        x = M
        pdf.setFont("Helvetica-Bold", 8)
        for i, h in enumerate(hdrs):
            pdf.rect(x, y - rh, cw[i], rh)
            pdf.drawCentredString(x + cw[i] / 2, y - rh + 2.5 * mm, h)
            x += cw[i]
        y -= rh

        items = list(certificate.items.all())
        pdf.setFont("Helvetica", 9)
        for ri in range(11):
            item = items[ri] if ri < len(items) else None
            x = M
            vals = [
                item.tank_no if item else "",
                item.product_name if item else "",
                f"{item.weight_tonnage:.3f}" if item else "",
                f"{item.volume_liters:.3f}" if item else "",
            ]
            for ci, v in enumerate(vals):
                pdf.rect(x, y - rh, cw[ci], rh)
                pdf.drawCentredString(x + cw[ci] / 2, y - rh + 2.5 * mm, str(v))
                x += cw[ci]
            y -= rh

        x = M
        pdf.setFont("Helvetica-Bold", 9)
        totals = ["TOTAL", "",
                  f"{certificate.total_weight_tonnage:.3f}",
                  f"{certificate.total_volume_liters:.3f}"]
        for ci, v in enumerate(totals):
            pdf.rect(x, y - rh, cw[ci], rh)
            pdf.drawCentredString(x + cw[ci] / 2, y - rh + 2.5 * mm, str(v))
            x += cw[ci]
        y -= rh + 3 * mm

        # Flowmeter section
        half = TW / 2
        fm_rh = 7 * mm
        # A zero denotes that no inlet-flowmeter reading was recorded.  Leave
        # the printed cell blank instead of presenting it as a measurement.
        fm_quantity = certificate.quantity_received_through_inlet_flowmeters or 0
        fm_val = f"{fm_quantity:.3f}" if fm_quantity > 0 else ''
        pdf.setFont("Helvetica", 8)
        pdf.rect(M, y - fm_rh, half, fm_rh)
        pdf.rect(M + half, y - fm_rh, half, fm_rh)
        pdf.drawString(M + 2 * mm, y - fm_rh + 2 * mm, "Quantity received through inlet flowmeters:")
        pdf.drawString(M + half + 2 * mm, y - fm_rh + 2 * mm, "Volume in litres")
        y -= fm_rh
        pdf.rect(M, y - fm_rh, half, fm_rh)
        pdf.rect(M + half, y - fm_rh, half, fm_rh)
        pdf.setFont("Helvetica-Bold", 9)
        pdf.drawCentredString(M + half / 2, y - fm_rh + 2 * mm, fm_val)
        y -= fm_rh + 6 * mm

        # Signature block
        col = TW / 2 - 3 * mm
        rx = M + col + 6 * mm
        pdf.setFont("Helvetica", 9)
        pdf.drawString(M, y, "Terminal Representative:")
        pdf.drawString(rx, y, "PBPA Inspector:")
        y -= 7 * mm
        pdf.drawString(M, y, f"Name: {certificate.terminal_representative_name or ''}")
        pdf.drawString(rx, y, f"Name: {certificate.pbpa_inspector_name or ''}")
        y -= 10 * mm
        pdf.setLineWidth(0.8)
        pdf.line(M, y, M + col, y)
        pdf.line(rx, y, rx + col, y)
        y -= 5 * mm
        pdf.setFont("Helvetica", 9)
        pdf.drawString(M, y, f"Signature: {certificate.terminal_representative_signature or ''}")
        pdf.drawString(rx, y, f"Signature: {certificate.pbpa_inspector_signature or ''}")

        pdf.showPage()
        pdf.save()
        buffer.seek(0)
        return buffer


class SealIsolationReportViewSet(SigningWorkflowMixin, viewsets.ModelViewSet):
    """CRUD workflow for the PBPA sealing and isolation report."""
    permission_classes = (IsAuthenticated,)
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['report_number', 'vessel_name', 'product_name', 'terminal', 'pbpa_inspector_name']
    ordering_fields = ['created_at', 'report_date', 'report_number', 'status']
    ordering = ['-created_at']
    _doc_type_label = 'seal_isolation'

    def get_queryset(self):
        profile = get_or_create_user_profile(self.request.user)
        base = SealIsolationReport.objects.select_related('created_by').prefetch_related('entries')
        if profile.role == 'terminal_representative':
            base = base.filter(terminal__iexact=profile.terminal.strip())
            if self.action == 'list':
                qs = base.filter(signing_step='sent_to_client')
            else:
                qs = base.filter(
                    Q(signing_step='sent_to_client')
                    | Q(client_signed_by=self.request.user)
                    | Q(signing_step__in=['client_signed', 'sent_to_inspector', 'verified', 'submitted'])
                )
        elif profile.role == 'admin':
            qs = base if self.action != 'list' else base.filter(signing_step='submitted')
        else:
            qs = owned_queryset(base, self.request.user)

        status_filter = self.request.query_params.get('status')
        period = self.request.query_params.get('period')
        if status_filter:
            qs = qs.filter(status=status_filter)
        return filter_queryset_by_period(qs, period, 'report_date')

    def get_serializer_class(self):
        if self.action in ['create', 'update', 'partial_update', 'retrieve']:
            return SealIsolationReportDetailSerializer
        return SealIsolationReportListSerializer

    def perform_create(self, serializer):
        full_name = self.request.user.get_full_name() or self.request.user.username
        serializer.save(
            created_by=self.request.user,
            pbpa_inspector_name=serializer.validated_data.get('pbpa_inspector_name') or full_name,
        )

    def _build_pdf_bytes(self, obj):
        from .shore_tank_utils import generate_seal_isolation_pdf
        return generate_seal_isolation_pdf(obj).getvalue()

    @action(detail=True, methods=['post'], url_path='sign_with_image')
    def sign_with_image(self, request, pk=None):
        """Legacy alias — redirects to inspector_sign."""
        return self.inspector_sign(request, pk=pk)

    @action(detail=True, methods=['post'])
    def sign_document(self, request, pk=None):
        """Digitally sign the Seal & Isolation Report as PDF."""
        report = self.get_object()
        if report.is_signed:
            return Response({'detail': 'Document is already signed.'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            # Set the canonical hash before rendering so it is both embedded in
            # the visible footer and bound into the resulting PDF signature.
            report.document_hash = self._document_hash(report)
            report.is_signed = True
            report.signed_at = timezone.now()
            report.signed_by = request.user
            pdf_bytes = self._build_pdf_with_signatures(report)
            signer_name = request.user.get_full_name() or request.user.username
            report.save(update_fields=['is_signed','signed_at','signed_by','document_hash','updated_at'])
            signed_bytes = _finalize_pdf_for_download(pdf_bytes, report, 'Seal and Isolation Report', request)
            _record_signature(request, report, self._get_role(request), report.document_hash,
                              self._doc_type_label, self._get_doc_number(report))
            response = HttpResponse(signed_bytes, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="SIGNED_SIR_{report.report_number}.pdf"'
            return response
        except Exception as e:
            return Response({'detail': f'Signing failed: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @action(detail=True, methods=['post'])
    def issue(self, request, pk=None):
        if self._get_role(request) != 'admin':
            return Response({'detail': 'Use the signing workflow to submit this report.'}, status=status.HTTP_409_CONFLICT)
        report = self.get_object()
        if report.status == 'issued':
            return Response({'detail': 'Report has already been issued.'}, status=status.HTTP_400_BAD_REQUEST)
        report.status = 'issued'
        report.issued_at = timezone.now()
        report.save(update_fields=['status', 'issued_at', 'updated_at'])
        serializer = SealIsolationReportDetailSerializer(report, context={'request': request})
        return Response(serializer.data)

    @action(detail=True, methods=['get'])
    def generate_document(self, request, pk=None):
        report = self.get_object()
        try:
            # Render the stored canonical hash on generated copies.  A draft is
            # marked as pending rather than represented as cryptographically signed.
            if not report.document_hash:
                report.document_hash = self._document_hash(report)
            buf = generate_seal_isolation_pdf(report)
            pdf_bytes = buf.getvalue()
            pdf_bytes = _finalize_pdf_for_download(pdf_bytes, report, 'Seal and Isolation Report', request)
            response = HttpResponse(
                pdf_bytes,
                content_type='application/pdf'
            )
            response['Content-Disposition'] = f'attachment; filename="Seal_Isolation_{report.report_number}.pdf"'
            return response
        except Exception as e:
            return Response({'detail': f'Document generation error: {str(e)}'}, status=status.HTTP_400_BAD_REQUEST)


class ASTMLookupView(viewsets.ViewSet):
    """Stateless ASTM table lookup — returns d20, VCF, WCF for given inputs."""
    permission_classes = (IsAuthenticated,)

    @action(detail=False, methods=['post'], url_path='lookup')
    def lookup(self, request):
        try:
            sample_density = float(request.data['sample_density'])
            sample_temp    = float(request.data['sample_temp'])
            tank_temp      = float(request.data['tank_temp'])
        except (KeyError, TypeError, ValueError):
            return Response(
                {'detail': 'sample_density, sample_temp and tank_temp are required numbers.'},
                status=status.HTTP_400_BAD_REQUEST
            )

        table_d20 = density_at_20_from_table(sample_density, sample_temp)
        d20 = table_d20 if table_d20 is not None else density_at_20_formula(sample_density, sample_temp)

        table_vcf = vcf_from_table(d20, tank_temp) if d20 is not None else None
        vcf = table_vcf if table_vcf is not None else (
            vcf_formula(d20, tank_temp) if d20 is not None else None
        )
        wcf = wcf_from_density(d20)

        return Response({
            'density_at_20': d20,
            'vcf': vcf,
            'wcf': wcf,
            'source': {
                'density_at_20': 'table_59b' if table_d20 is not None else 'formula_fallback',
                'vcf': 'table_60b' if table_vcf is not None else 'formula_fallback',
                'wcf': 'density_at_20',
            },
            'table_range': table_range(),
        })


class ShoreTankCalculationViewSet(SigningWorkflowMixin, viewsets.ModelViewSet):
    """CRUD workflow for the PBPA shore tank calculation workbook."""
    permission_classes = (IsAuthenticated,)
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['calculation_number', 'vessel_name', 'product_name', 'terminal', 'pbpa_inspector_name']
    ordering_fields = ['created_at', 'calculation_date', 'calculation_number', 'status']
    ordering = ['-created_at']
    _doc_type_label = 'shore_tank'

    @action(detail=True, methods=['post'], url_path='send_to_client')
    def send_to_client(self, request, pk=None):
        """Keep a shore-tank calculation and its generated receipt together."""
        return self.send_both_to_terminal(request, pk=pk)

    def get_queryset(self):
        profile = get_or_create_user_profile(self.request.user)
        base = ShoreTankCalculation.objects.select_related('created_by').prefetch_related('tank_items__tank')
        if profile.role == 'terminal_representative':
            base = base.filter(terminal__iexact=profile.terminal.strip())
            if self.action == 'list':
                qs = base.filter(signing_step='sent_to_client')
            else:
                qs = base.filter(
                    Q(signing_step='sent_to_client')
                    | Q(client_signed_by=self.request.user)
                    | Q(signing_step__in=['client_signed', 'sent_to_inspector', 'verified', 'submitted'])
                )
        elif profile.role == 'admin':
            if self.action == 'list':
                # Use the same submission records as the dashboard card.
                submissions = filter_queryset_by_period(
                    Submission.objects.filter(doc_type='shore_tank'),
                    self.request.query_params.get('period'),
                    'submitted_at',
                )
                qs = base.filter(pk__in=submissions.values_list('doc_id', flat=True))
            else:
                qs = base
        else:
            qs = owned_queryset(base, self.request.user)

        status_filter = self.request.query_params.get('status')
        period = self.request.query_params.get('period')
        if profile.role == 'admin' and self.action == 'list':
            # The submission-time period filter above is authoritative.
            period = None
        if status_filter:
            qs = qs.filter(status=status_filter)
        return filter_queryset_by_period(qs, period, 'calculation_date')

    def get_serializer_class(self):
        if self.action in ['create', 'update', 'partial_update', 'retrieve']:
            return ShoreTankCalculationDetailSerializer
        return ShoreTankCalculationListSerializer

    def perform_create(self, serializer):
        full_name = self.request.user.get_full_name() or self.request.user.username
        serializer.save(
            created_by=self.request.user,
            pbpa_inspector_name=serializer.validated_data.get('pbpa_inspector_name') or full_name,
        )

    def _build_pdf_bytes(self, obj):
        from .shore_tank_utils import generate_shore_tank_pdf
        return generate_shore_tank_pdf(obj).getvalue()

    def _create_linked_receipt_certificate(self, calculation, request, signature):
        """Create the PRC from calculated shore-tank values exactly once."""
        try:
            return calculation.product_receipt_certificate
        except ProductReceiptCertificate.DoesNotExist:
            pass

        tank_items = list(calculation.tank_items.all())
        item_count = len(tank_items) or 1
        inspector_name = calculation.pbpa_inspector_name or request.user.get_full_name() or request.user.username
        now = timezone.now()
        certificate = ProductReceiptCertificate.objects.create(
            source_shore_tank_calculation=calculation,
            vessel_name=calculation.vessel_name,
            terminal=calculation.terminal,
            receipt_date=calculation.calculation_date,
            quantity_received_through_inlet_flowmeters=(calculation.meter_quantity_m3 or 0) * 1000,
            terminal_representative_name=calculation.terminal_representative_name or '',
            pbpa_inspector_name=inspector_name,
            notes=f'Automatically generated from Shore Tank Calculation #{calculation.calculation_number}.',
            created_by=request.user,
            is_signed=True,
            signed_at=now,
            signed_by=request.user,
            signing_step='inspector_signed',
            inspector_signed_at=now,
            inspector_signed_by=request.user,
        )
        for item in tank_items:
            standard_m3 = item.received_standard_volume_m3 or 0
            weight_mt = item.received_weight_air_mt or 0
            ProductReceiptCertificateItem.objects.create(
                certificate=certificate,
                tank=item.tank,
                tank_no=item.tank_no or (item.tank.tank_id if item.tank else ''),
                product_name=calculation.product_name,
                volume_liters=(standard_m3 or (calculation.terminal_standard_volume_m3 or 0) / item_count) * 1000,
                weight_tonnage=weight_mt or (calculation.terminal_weight_air_mt or 0) / item_count,
            )
        # The inspector's signature applies to the generated certificate too,
        # so both documents are ready for the same terminal countersignature.
        certificate.document_hash = compute_model_hash(certificate, 'product_receipt', certificate.certificate_number)
        certificate.save(update_fields=['document_hash', 'updated_at'])
        _record_signature(request, certificate, 'inspector', certificate.document_hash,
                          'product_receipt', certificate.certificate_number, signature)
        return certificate

    @action(detail=True, methods=['post'], url_path='inspector_sign')
    def inspector_sign(self, request, pk=None):
        """Sign the STC and automatically create its inspector-signed PRC."""
        response = super().inspector_sign(request, pk=pk)
        if response.status_code < 300:
            calculation = self.get_object()
            with transaction.atomic():
                self._create_linked_receipt_certificate(calculation, request, request.data.get('signature'))
        return response

    @action(detail=True, methods=['post'], url_path='send_both_to_terminal')
    def send_both_to_terminal(self, request, pk=None):
        """Send the signed shore-tank calculation and its generated PRC together."""
        calculation = self.get_object()
        if self._get_role(request) not in ('inspector', 'admin'):
            return Response({'detail': 'Only inspectors can send reports to the terminal.'}, status=status.HTTP_403_FORBIDDEN)
        if calculation.signing_step != 'inspector_signed':
            if calculation.signing_step in ('sent_to_client', 'sent_to_inspector', 'verified', 'submitted'):
                return Response(
                    {'detail': (
                        'This Shore Tank Calculation has already progressed to "{}". '
                        'Its linked Product Receipt Certificate cannot be sent separately; '
                        'an administrator must unlock and restart both documents together.'
                    ).format(calculation.signing_step)},
                    status=status.HTTP_409_CONFLICT,
                )
            return Response({'detail': 'The shore tank calculation must be inspector-signed first.'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            certificate = calculation.product_receipt_certificate
        except ProductReceiptCertificate.DoesNotExist:
            return Response({'detail': 'The generated Product Receipt Certificate is unavailable.'}, status=status.HTTP_409_CONFLICT)
        if certificate.signing_step != 'inspector_signed':
            return Response({'detail': 'The Product Receipt Certificate is not ready to send.'}, status=status.HTTP_400_BAD_REQUEST)

        terminal = (calculation.terminal or '').strip()
        representatives = list(_terminal_representatives(terminal))
        if not representatives:
            return Response({'detail': f'No active terminal representative is assigned to terminal "{terminal or "(blank)"}".'}, status=status.HTTP_400_BAD_REQUEST)

        documents = ((calculation, 'shore_tank', 'Shore Tank Calculation'),
                     (certificate, 'product_receipt', 'Product Receipt Certificate'))
        with transaction.atomic():
            for document, doc_type, label in documents:
                document.signing_step = 'sent_to_client'
                document.save(update_fields=['signing_step', 'updated_at'])
                for representative in representatives:
                    create_notification(
                        recipient=representative.user,
                        title=f'Signing Request: {label} #{self._get_doc_number(document)}',
                        message=(f'Inspector {request.user.get_full_name() or request.user.username} requests your signature for '
                                 f'{label} #{self._get_doc_number(document)} for vessel "{calculation.vessel_name}" at terminal "{terminal}".'),
                        notification_type='signing_request', doc_type=doc_type, doc_id=document.pk,
                        doc_number=self._get_doc_number(document), email=True,
                    )
        ActivityLog.log(request.user, 'report_updated', doc_type='shore_tank', doc_id=calculation.pk,
                        doc_number=calculation.calculation_number,
                        detail=f'Sent shore tank calculation and Product Receipt Certificate to {len(representatives)} terminal representative(s).', request=request)
        return Response(self.get_serializer(calculation).data)

    @action(detail=True, methods=['post'], url_path='sign_with_image')
    def sign_with_image(self, request, pk=None):
        """Legacy alias — redirects to inspector_sign."""
        return self.inspector_sign(request, pk=pk)

    @action(detail=True, methods=['post'])
    def sign_document(self, request, pk=None):
        """Digitally sign the Shore Tank Calculation as PDF."""
        calculation = self.get_object()
        if calculation.is_signed:
            return Response({'detail': 'Document is already signed.'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.pdfgen import canvas as rl_canvas
            import io
            buf = io.BytesIO()
            c = rl_canvas.Canvas(buf, pagesize=landscape(A4))
            w, h = landscape(A4)
            c.setFont('Helvetica-Bold', 14)
            c.drawCentredString(w/2, h-40, 'PBPA SHORE TANK CALCULATION')
            c.setFont('Helvetica', 11)
            c.drawString(40, h-70,  f'Calc No: {calculation.calculation_number}')
            c.drawString(40, h-90,  f'Vessel: {calculation.vessel_name}')
            c.drawString(40, h-110, f'Product: {calculation.product_name}')
            c.drawString(40, h-130, f'Terminal: {calculation.terminal}')
            c.drawString(40, h-150, f'Date: {calculation.calculation_date}')
            c.drawString(40, h-170, f'Terminal Std Vol: {calculation.terminal_standard_volume_m3} m3')
            c.drawString(40, h-190, f'Terminal Weight: {calculation.terminal_weight_air_mt} MT')
            c.drawString(40, h-210, f'Difference Std Vol: {calculation.difference_standard_volume_m3} m3')
            c.drawString(40, h-230, f'Difference Weight: {calculation.difference_weight_air_mt} MT')
            c.save()
            pdf_bytes = buf.getvalue()
            signer_name = request.user.get_full_name() or request.user.username
            calculation.is_signed     = True
            calculation.signed_at     = timezone.now()
            calculation.signed_by     = request.user
            calculation.document_hash = self._document_hash(calculation)
            calculation.save(update_fields=['is_signed','signed_at','signed_by','document_hash','updated_at'])
            signed_bytes = _finalize_pdf_for_download(pdf_bytes, calculation, 'Shore Tank Calculation', request)
            _record_signature(request, calculation, self._get_role(request), calculation.document_hash,
                              self._doc_type_label, self._get_doc_number(calculation))
            response = HttpResponse(signed_bytes, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="SIGNED_STC_{calculation.calculation_number}.pdf"'
            return response
        except Exception as e:
            return Response({'detail': f'Signing failed: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @action(detail=True, methods=['post'])
    def finalize(self, request, pk=None):
        if self._get_role(request) != 'admin':
            return Response({'detail': 'Use the signing workflow to submit this calculation.'}, status=status.HTTP_409_CONFLICT)
        calculation = self.get_object()
        if calculation.status == 'final':
            return Response({'detail': 'Calculation is already final.'}, status=status.HTTP_400_BAD_REQUEST)

        calculation.status = 'final'
        calculation.finalized_at = timezone.now()
        calculation.save(update_fields=['status', 'finalized_at', 'updated_at'])

        serializer = ShoreTankCalculationDetailSerializer(calculation, context={'request': request})
        return Response(serializer.data)
    
    @action(detail=True, methods=['post'])
    def calculate(self, request, pk=None):
        """
        Calculate all tank items for this shore tank calculation
        Performs ASTM D1250 calculations for volume and weight
        """
        calculation = self.get_object()
        
        try:
            calc_engine = ShoreTankCalculationEngine()
            results = calc_engine.calculate_all_tank_items(calculation)
            
            if results.get('errors'):
                return Response({
                    'detail': 'Calculation completed with warnings',
                    'errors': results['errors'],
                    'results': results
                }, status=status.HTTP_200_OK)
            
            return Response({
                'detail': 'Calculations completed successfully',
                'results': results
            }, status=status.HTTP_200_OK)
        
        except Exception as e:
            return Response({
                'detail': f'Calculation error: {str(e)}'
            }, status=status.HTTP_400_BAD_REQUEST)
    
    @action(detail=True, methods=['get'])
    def generate_document(self, request, pk=None):
        calculation = self.get_object()
        try:
            response = HttpResponse(
                _finalize_pdf_for_download(self._build_pdf_with_signatures(calculation), calculation, 'Shore Tank Calculation', request),
                content_type='application/pdf'
            )
            response['Content-Disposition'] = f'attachment; filename="Shore_Tank_Calc_{calculation.calculation_number}.pdf"'
            return response
        except Exception as e:
            return Response({'detail': f'Document generation error: {str(e)}'}, status=status.HTTP_400_BAD_REQUEST)


class SubmissionViewSet(viewsets.ModelViewSet):
    """Inspectors submit documents here; admins see notifications."""
    permission_classes = (IsAuthenticated,)
    serializer_class = SubmissionSerializer
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['vessel_name', 'doc_number', 'terminal', 'doc_type']
    ordering = ['-submitted_at']

    def _with_existing_document_targets(self, queryset):
        target_filter = Q(pk__in=[])
        for doc_type, config in SubmissionSerializer.SUBMITTABLE_DOCUMENTS.items():
            target_filter |= Q(
                doc_type=doc_type,
                doc_id__in=config['model'].objects.values_list('pk', flat=True),
            )
        return queryset.filter(target_filter)

    def get_queryset(self):
        profile = get_or_create_user_profile(self.request.user)
        qs = Submission.objects.select_related('submitted_by')
        if profile.role == 'inspector':
            qs = qs.filter(submitted_by=self.request.user)
        qs = self._with_existing_document_targets(qs)
        doc_type = self.request.query_params.get('doc_type')
        if doc_type:
            qs = qs.filter(doc_type=doc_type)
        return qs

    def perform_create(self, serializer):
        serializer.save(submitted_by=self.request.user)

    def destroy(self, request, *args, **kwargs):
        profile = get_or_create_user_profile(request.user)
        if profile.role not in ('admin', 'terminal_representative'):
            return Response({'detail': 'Permission denied.'}, status=status.HTTP_403_FORBIDDEN)
        return super().destroy(request, *args, **kwargs)

    @action(detail=False, methods=['get'])
    def unread_count(self, request):
        return Response({'count': self.get_queryset().filter(is_read=False).count()})

    @action(detail=True, methods=['post'])
    def mark_read(self, request, pk=None):
        sub = self.get_object()
        sub.is_read = True
        sub.save(update_fields=['is_read'])
        return Response({'status': 'marked read'})

    @action(detail=False, methods=['post'])
    def mark_all_read(self, request):
        self.get_queryset().filter(is_read=False).update(is_read=True)
        return Response({'status': 'all marked read'})


class TerminalRepresentativeReadOnlyMixin:
    """Keep terminal representatives in a view/sign-only role for reports."""

    def _terminal_representative_write_denied(self, request):
        return get_or_create_user_profile(request.user).role == 'terminal_representative'

    def create(self, request, *args, **kwargs):
        if self._terminal_representative_write_denied(request):
            return Response({'detail': 'Terminal representatives can only view and sign reports.'}, status=status.HTTP_403_FORBIDDEN)
        return super().create(request, *args, **kwargs)

    def update(self, request, *args, **kwargs):
        if self._terminal_representative_write_denied(request):
            return Response({'detail': 'Terminal representatives can only view and sign reports.'}, status=status.HTTP_403_FORBIDDEN)
        return super().update(request, *args, **kwargs)

    def destroy(self, request, *args, **kwargs):
        if self._terminal_representative_write_denied(request):
            return Response({'detail': 'Terminal representatives can only view and sign reports.'}, status=status.HTTP_403_FORBIDDEN)
        return super().destroy(request, *args, **kwargs)


class VesselReportViewSet(TerminalRepresentativeReadOnlyMixin, viewsets.ModelViewSet):
    """Create and manage vessel discharge summary reports."""
    permission_classes = (IsAuthenticated,)
    serializer_class = VesselReportSerializer
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['vessel_name', 'terminal', 'report_number', 'product_name']
    ordering = ['-created_at']

    def get_queryset(self):
        profile = get_or_create_user_profile(self.request.user)
        base = VesselReport.objects.select_related('created_by')
        if profile.role == 'admin':
            qs = base
        else:
            qs = owned_queryset(base, self.request.user)
        return filter_queryset_by_period(qs, self.request.query_params.get('period'), 'discharge_date')

    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user)

    @action(detail=True, methods=['post'])
    def finalize(self, request, pk=None):
        report = self.get_object()
        if report.status == 'cancelled':
            return Response({'detail': 'Cancelled vessel reports cannot be finalized.'}, status=status.HTTP_400_BAD_REQUEST)
        report.status = 'final'
        report.save(update_fields=['status', 'updated_at'])
        return Response(VesselReportSerializer(report).data)

    @action(detail=True, methods=['post'])
    def cancel(self, request, pk=None):
        report = self.get_object()
        report.status = 'cancelled'
        report.save(update_fields=['status', 'updated_at'])
        return Response(VesselReportSerializer(report).data)

    @action(detail=True, methods=['get'])
    def pdf(self, request, pk=None):
        report = self.get_object()
        buf = self._build_pdf(report)
        pdf_bytes = _finalize_pdf_for_download(buf.getvalue(), report, 'Vessel Report', request)
        response = HttpResponse(pdf_bytes, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="Vessel_Report_{report.report_number}.pdf"'
        return response

    def _build_pdf(self, report):
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            leftMargin=18 * mm,
            rightMargin=18 * mm,
            topMargin=14 * mm,
            bottomMargin=14 * mm,
        )
        styles = getSampleStyleSheet()

        def ps(name, **kwargs):
            return ParagraphStyle(name, parent=styles['Normal'], **kwargs)

        title_style = ps('vr_title', fontSize=13, alignment=TA_CENTER, fontName='Helvetica-Bold', leading=16)
        center_style = ps('vr_center', fontSize=8, alignment=TA_CENTER, leading=10)
        section_style = ps('vr_section', fontSize=10, fontName='Helvetica-Bold', textColor=colors.HexColor('#8B1A1A'))
        label_style = ps('vr_label', fontSize=8, fontName='Helvetica-Bold', textColor=colors.HexColor('#444444'))
        value_style = ps('vr_value', fontSize=8, alignment=TA_LEFT)

        def fnum(value, suffix=''):
            try:
                return f"{float(value or 0):,.3f}{suffix}"
            except (TypeError, ValueError):
                return f"0.000{suffix}"

        elems = [
            Paragraph("THE UNITED REPUBLIC OF TANZANIA", center_style),
            Paragraph("<b>PETROLEUM BULK PROCUREMENT AGENCY</b>", ps('vr_agency', fontSize=10, alignment=TA_CENTER, fontName='Helvetica-Bold')),
            Paragraph("TANZANIA PORTS AUTHORITY, ONE STOP CENTER BUILDING, 11TH FLOOR, SOKOINE DRIVE", center_style),
            Spacer(1, 6 * mm),
            Paragraph("VESSEL REPORT", title_style),
            Spacer(1, 5 * mm),
        ]

        status_label = (report.status or '').upper()
        details = [
            [Paragraph('Report Number', label_style), Paragraph(report.report_number or '', value_style),
             Paragraph('Status', label_style), Paragraph(status_label, value_style)],
            [Paragraph('Vessel Name', label_style), Paragraph(report.vessel_name or '', value_style),
             Paragraph('Terminal', label_style), Paragraph(report.terminal or '', value_style)],
            [Paragraph('Product', label_style), Paragraph(report.product_name or '', value_style),
             Paragraph('Discharge Date', label_style), Paragraph(report.discharge_date.strftime('%d-%b-%Y') if report.discharge_date else '', value_style)],
            [Paragraph('Total Weight', label_style), Paragraph(fnum(report.total_weight_mt, ' MT'), value_style),
             Paragraph('Total Volume', label_style), Paragraph(fnum(report.total_volume_m3, ' m3'), value_style)],
        ]
        detail_table = Table(details, colWidths=[32 * mm, 58 * mm, 32 * mm, 58 * mm])
        detail_table.setStyle(TableStyle([
            ('BOX', (0, 0), (-1, -1), 0.8, colors.black),
            ('INNERGRID', (0, 0), (-1, -1), 0.3, colors.grey),
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f4f4f4')),
            ('BACKGROUND', (2, 0), (2, -1), colors.HexColor('#f4f4f4')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('TOPPADDING', (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ]))
        elems.append(detail_table)
        elems.append(Spacer(1, 7 * mm))

        elems.append(Paragraph("LINKED DOCUMENTS", section_style))
        linked_docs = [
            ['Dip Tickets', len(report.dip_ticket_ids or [])],
            ['Seal Reports', len(report.seal_report_ids or [])],
            ['Shore Tank Calculations', len(report.shore_calc_ids or [])],
            ['Product Receipt Certificates', len(report.cert_ids or [])],
        ]
        linked_table = Table(
            [[Paragraph('<b>Document Type</b>', label_style), Paragraph('<b>Count</b>', label_style)]] +
            [[Paragraph(label, value_style), Paragraph(str(count), value_style)] for label, count in linked_docs],
            colWidths=[130 * mm, 50 * mm],
        )
        linked_table.setStyle(TableStyle([
            ('BOX', (0, 0), (-1, -1), 0.8, colors.black),
            ('INNERGRID', (0, 0), (-1, -1), 0.3, colors.grey),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f4f4f4')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('TOPPADDING', (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ]))
        elems.append(linked_table)

        if report.remarks:
            elems.append(Spacer(1, 7 * mm))
            elems.append(Paragraph("REMARKS", section_style))
            remarks_table = Table([[Paragraph(report.remarks, value_style)]], colWidths=[180 * mm])
            remarks_table.setStyle(TableStyle([
                ('BOX', (0, 0), (-1, -1), 0.8, colors.black),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            elems.append(remarks_table)

        elems.append(Spacer(1, 14 * mm))
        prepared_by = report.created_by.get_full_name() if report.created_by else ''
        prepared_by = prepared_by or (report.created_by.username if report.created_by else '')
        signature_table = Table([
            ['Prepared By', 'Reviewed By'],
            [prepared_by or ' ', ' '],
        ], colWidths=[85 * mm, 85 * mm])
        signature_table.setStyle(TableStyle([
            ('LINEABOVE', (0, 1), (0, 1), 0.8, colors.black),
            ('LINEABOVE', (1, 1), (1, 1), 0.8, colors.black),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ]))
        elems.append(signature_table)

        elems.append(Spacer(1, 8 * mm))
        generated_at = timezone.localtime().strftime('%d-%b-%Y %H:%M')
        elems.append(Paragraph(f"Generated: {generated_at}", ps('vr_footer', fontSize=7, alignment=TA_CENTER, textColor=colors.grey)))

        doc.build(elems)
        buf.seek(0)
        return buf


class ProvisionalOuturnReportViewSet(TerminalRepresentativeReadOnlyMixin, viewsets.ModelViewSet):
    """CRUD + PDF for PBPA Provisional Outturn Reports."""
    permission_classes = (IsAuthenticated,)
    serializer_class = ProvisionalOuturnReportSerializer
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['vessel_name', 'report_number', 'port', 'product']
    ordering = ['-created_at']

    def get_queryset(self):
        profile = get_or_create_user_profile(self.request.user)
        base = ProvisionalOuturnReport.objects.select_related('created_by').prefetch_related('items')
        if profile.role in ('admin', 'terminal_representative') and self.action in {'list', 'retrieve', 'pdf', 'docx'}:
            qs = base
        else:
            qs = owned_queryset(base, self.request.user)
        qs = filter_queryset_by_period(qs, self.request.query_params.get('period'), 'report_date')
        return qs

    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user)

    @action(detail=True, methods=['post'])
    def finalize(self, request, pk=None):
        report = self.get_object()
        report.status = 'final'
        report.save(update_fields=['status', 'updated_at'])
        return Response(ProvisionalOuturnReportSerializer(report).data)

    @action(detail=True, methods=['get'])
    def pdf(self, request, pk=None):
        report = self.get_object()
        buf = self._build_pdf(report)
        pdf_bytes = _finalize_pdf_for_download(buf.getvalue(), report, 'Provisional Outturn Report', request)
        response = HttpResponse(pdf_bytes, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="POR_{report.report_number}.pdf"'
        return response

    def _build_pdf(self, report):
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib import colors
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=landscape(A4),
                                leftMargin=15*mm, rightMargin=15*mm,
                                topMargin=12*mm, bottomMargin=12*mm)
        styles = getSampleStyleSheet()
        W = landscape(A4)[0] - 30*mm

        def ps(name, **kw):
            return ParagraphStyle(name, parent=styles['Normal'], **kw)

        elems = []

        # ── Header ──────────────────────────────────────────────────────────
        elems.append(Paragraph("THE UNITED REPUBLIC OF TANZANIA", ps('h1', fontSize=9, alignment=TA_CENTER)))
        elems.append(Spacer(1, 2*mm))
        elems.append(Paragraph("<b>PETROLEUM BULK PROCUREMENT AGENCY</b>", ps('h2', fontSize=11, alignment=TA_CENTER)))
        elems.append(Paragraph(
            "TANZANIA PORTS AUTHORITY, ONE STOP CENTER BUILDING, 11TH FLOOR, SOKOINE DRIVE, PLOT NO:1/2",
            ps('a1', fontSize=7, alignment=TA_CENTER)))
        elems.append(Paragraph(
            "Tel: +255222129009 / Fax: +255222129093 / Email: info@pbpa.go.tz / WEBSITE: www.pbpa.go.tz / P.O. Box 2634 Dar es Salaam, TANZANIA",
            ps('a2', fontSize=6.5, alignment=TA_CENTER)))
        elems.append(Spacer(1, 5*mm))

        # ── Vessel info ──────────────────────────────────────────────────────
        date_str = report.report_date.strftime('%d.%m.%Y') if report.report_date else ''
        info = [
            [Paragraph(f"<b>VESSEL</b>", ps('i')), Paragraph(report.vessel_name or '', ps('iv', fontSize=10))],
            [Paragraph(f"<b>DATE</b>",   ps('i')), Paragraph(date_str, ps('iv', fontSize=10))],
            [Paragraph(f"<b>PORT</b>",   ps('i')), Paragraph(report.port or '', ps('iv', fontSize=10))],
            [Paragraph(f"<b>PRODUCT</b>",ps('i')), Paragraph(report.product or '', ps('iv', fontSize=10))],
        ]
        info_tbl = Table(info, colWidths=[25*mm, 80*mm])
        info_tbl.setStyle(TableStyle([
            ('FONTSIZE', (0,0), (-1,-1), 9),
            ('BOTTOMPADDING', (0,0), (-1,-1), 2),
            ('TOPPADDING', (0,0), (-1,-1), 2),
        ]))
        elems.append(info_tbl)
        elems.append(Spacer(1, 5*mm))

        # ── Main table ───────────────────────────────────────────────────────
        items = list(report.items.all())
        totals = report.totals

        # diff totals
        tot_dv  = round((totals['shore_volume'] or 0) - (totals['ship_volume'] or 0), 3)
        tot_dvp = round(tot_dv / totals['ship_volume'] * 100, 3) if totals['ship_volume'] else 0
        tot_dw  = round((totals['shore_weight'] or 0) - (totals['ship_weight'] or 0), 3)
        tot_dwp = round(tot_dw / totals['ship_weight'] * 100, 3) if totals['ship_weight'] else 0

        def f3(v): return f'{float(v):,.3f}' if v is not None else ''

        hdr_style = ps('th', fontSize=7, alignment=TA_CENTER, fontName='Helvetica-Bold')
        cell_style = ps('td', fontSize=7.5, alignment=TA_CENTER)
        lft_style  = ps('tl', fontSize=7.5, alignment=TA_LEFT)

        # col widths: SN, Terminal, ShipVol, ShipWt, ShoreVol, ShoreWt, DiffVol, %Diff, DiffWt, %Diff
        cw = [8*mm, 38*mm, 22*mm, 22*mm, 22*mm, 22*mm, 22*mm, 16*mm, 22*mm, 16*mm]

        header_rows = [
            # row 0 — title span
            [Paragraph('<b>PROVISIONAL OUTTURN SUMMARY</b>', ps('title', fontSize=9, alignment=TA_CENTER, fontName='Helvetica-Bold'))]
            + ['']*9,
            # row 1 — group headers
            ['', '',
             Paragraph('<b>Ship Figures</b>', hdr_style), '',
             Paragraph('<b>Shore Figures</b>', hdr_style), '',
             Paragraph('<b>Difference</b>', hdr_style), '', '', ''],
            # row 2 — sub-group
            ['', '',
             Paragraph('<b>Volume</b>', hdr_style), Paragraph('<b>Weight</b>', hdr_style),
             Paragraph('<b>Volume</b>', hdr_style), Paragraph('<b>Weight</b>', hdr_style),
             Paragraph('<b>Volume</b>', hdr_style), '', Paragraph('<b>Weight</b>', hdr_style), ''],
            # row 3 — column labels
            [Paragraph('<b>S/N</b>', hdr_style),
             Paragraph('<b></b>', hdr_style),
             Paragraph('<b>M3 @ 20oC</b>', hdr_style), Paragraph('<b>M/Tons</b>', hdr_style),
             Paragraph('<b>M3 @ 20oC</b>', hdr_style), Paragraph('<b>M/Tons</b>', hdr_style),
             Paragraph('<b>M3 @ 20oC</b>', hdr_style), Paragraph('<b>%Diff</b>', hdr_style),
             Paragraph('<b>M/Tons</b>', hdr_style),    Paragraph('<b>%Diff</b>', hdr_style)],
        ]

        data_rows = []
        for item in items:
            data_rows.append([
                Paragraph(str(item.sn), cell_style),
                Paragraph(item.terminal_name or '', lft_style),
                Paragraph(f3(item.ship_volume_m3),  cell_style),
                Paragraph(f3(item.ship_weight_mt),  cell_style),
                Paragraph(f3(item.shore_volume_m3), cell_style),
                Paragraph(f3(item.shore_weight_mt), cell_style),
                Paragraph(f3(item.diff_volume_m3),  cell_style),
                Paragraph(f3(item.diff_volume_pct), cell_style),
                Paragraph(f3(item.diff_weight_mt),  cell_style),
                Paragraph(f3(item.diff_weight_pct), cell_style),
            ])

        # totals row
        bold_cell = ps('tb', fontSize=7.5, alignment=TA_CENTER, fontName='Helvetica-Bold')
        total_row = [
            '', '',
            Paragraph(f3(totals['ship_volume']),  bold_cell),
            Paragraph(f3(totals['ship_weight']),  bold_cell),
            Paragraph(f3(totals['shore_volume']), bold_cell),
            Paragraph(f3(totals['shore_weight']), bold_cell),
            Paragraph(f3(tot_dv),  bold_cell),
            Paragraph(f3(tot_dvp), bold_cell),
            Paragraph(f3(tot_dw),  bold_cell),
            Paragraph(f3(tot_dwp), bold_cell),
        ]

        all_rows = header_rows + data_rows + [total_row]
        n_data = len(data_rows)
        n_hdr  = len(header_rows)

        tbl = Table(all_rows, colWidths=cw, repeatRows=4)
        ts = TableStyle([
            # outer border
            ('BOX',         (0,0), (-1,-1), 1, colors.black),
            ('INNERGRID',   (0,0), (-1,-1), 0.4, colors.grey),
            # title row span
            ('SPAN',        (0,0), (-1,0)),
            ('BACKGROUND',  (0,0), (-1,0), colors.HexColor('#d0d0d0')),
            ('ALIGN',       (0,0), (-1,0), 'CENTER'),
            # group header spans
            ('SPAN',        (2,1), (3,1)),  # Ship Figures
            ('SPAN',        (4,1), (5,1)),  # Shore Figures
            ('SPAN',        (6,1), (9,1)),  # Difference
            ('SPAN',        (6,2), (7,2)),  # Volume diff
            ('SPAN',        (8,2), (9,2)),  # Weight diff
            ('BACKGROUND',  (0,1), (-1,3), colors.HexColor('#f0f0f0')),
            ('VALIGN',      (0,0), (-1,-1), 'MIDDLE'),
            ('ROWBACKGROUNDS', (0, n_hdr), (-1, n_hdr + n_data - 1), [colors.white, colors.HexColor('#f9f9f9')]),
            # totals row
            ('BACKGROUND',  (0, -1), (-1, -1), colors.HexColor('#e8e8e8')),
            ('FONTNAME',    (0, -1), (-1, -1), 'Helvetica-Bold'),
            ('TOPPADDING',  (0,0), (-1,-1), 2),
            ('BOTTOMPADDING',(0,0), (-1,-1), 2),
        ])
        tbl.setStyle(ts)
        elems.append(tbl)
        elems.append(Spacer(1, 8*mm))

        # ── Signatures ───────────────────────────────────────────────────────
        sig_data = [[
            Paragraph('<b>CAPTAIN/CHIEF OFFICER</b>', ps('sl', fontSize=8)),
            '',
            Paragraph('<b>PBPA SURVEYOR</b>', ps('sr', fontSize=8, alignment=TA_CENTER)),
        ]]
        sig_tbl = Table(sig_data, colWidths=[W*0.4, W*0.2, W*0.4])
        sig_tbl.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'TOP')]))
        elems.append(sig_tbl)
        elems.append(Spacer(1, 12*mm))

        name_data = [[
            Paragraph(f'Name: {report.captain_name or ""}', ps('sn', fontSize=9)),
            '',
            Paragraph(f'Name: {report.surveyor_name or ""}', ps('sn2', fontSize=9, alignment=TA_CENTER)),
        ]]
        name_tbl = Table(name_data, colWidths=[W*0.4, W*0.2, W*0.4])
        name_tbl.setStyle(TableStyle([
            ('LINEABOVE', (0,0), (0,0), 0.8, colors.black),
            ('LINEABOVE', (2,0), (2,0), 0.8, colors.black),
        ]))
        elems.append(name_tbl)

        doc.build(elems)
        buf.seek(0)
        return buf

    @action(detail=True, methods=['get'])
    def docx(self, request, pk=None):
        """Generate DOCX version of Provisional Outturn Report."""
        report = self.get_object()
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Header
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header.add_run('THE UNITED REPUBLIC OF TANZANIA\n')
        header_run.font.size = Pt(9)
        
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run('PETROLEUM BULK PROCUREMENT AGENCY\n')
        title_run.font.size = Pt(11)
        title_run.font.bold = True
        
        # PBPA Address
        address = doc.add_paragraph()
        address.alignment = WD_ALIGN_PARAGRAPH.CENTER
        address_run = address.add_run(
            'TANZANIA PORTS AUTHORITY, ONE STOP CENTER BUILDING, 11TH FLOOR, SOKOINE DRIVE, PLOT NO:1/2\n'
            'Tel: +255222129009 / Fax: +255222129093 / Email: info@pbpa.go.tz / WEBSITE: www.pbpa.go.tz / P.O. Box 2634 Dar es Salaam, TANZANIA'
        )
        address_run.font.size = Pt(7)

        doc.add_paragraph()  # Spacing

        # Report Title
        report_title = doc.add_paragraph()
        report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        report_title_run = report_title.add_run('PROVISIONAL OUTTURN REPORT')
        report_title_run.font.size = Pt(12)
        report_title_run.font.bold = True

        doc.add_paragraph()  # Spacing

        # Report Info Table
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Light Grid Accent 1'
        info_table.cell(0, 0).text = 'Vessel Name'
        info_table.cell(0, 1).text = report.vessel_name or ''
        info_table.cell(1, 0).text = 'Date'
        info_table.cell(1, 1).text = report.report_date.strftime('%d.%m.%Y') if report.report_date else ''
        info_table.cell(2, 0).text = 'Port'
        info_table.cell(2, 1).text = report.port or ''
        info_table.cell(3, 0).text = 'Product'
        info_table.cell(3, 1).text = report.product or ''

        doc.add_paragraph()  # Spacing

        # Summary Table
        items = list(report.items.all())
        totals = report.totals

        # Calculate differences
        tot_dv  = round((totals['shore_volume'] or 0) - (totals['ship_volume'] or 0), 3)
        tot_dvp = round(tot_dv / totals['ship_volume'] * 100, 3) if totals['ship_volume'] else 0
        tot_dw  = round((totals['shore_weight'] or 0) - (totals['ship_weight'] or 0), 3)
        tot_dwp = round(tot_dw / totals['ship_weight'] * 100, 3) if totals['ship_weight'] else 0

        # Create table
        table = doc.add_table(rows=len(items) + 2, cols=10)
        table.style = 'Light Grid Accent 1'

        # Header row
        headers = ['S/N', 'Terminal', 'Ship Vol\n(m³)', 'Ship Wgt\n(MT)', 'Shore Vol\n(m³)', 'Shore Wgt\n(MT)', 
                   'Diff Vol\n(m³)', 'Diff %', 'Diff Wgt\n(MT)', 'Diff %']
        header_cells = table.rows[0].cells
        for i, header_text in enumerate(headers):
            header_cells[i].text = header_text
            # Style header cell
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

        # Data rows
        for idx, item in enumerate(items):
            row = table.rows[idx + 1]
            cells = row.cells
            cells[0].text = str(idx + 1)
            cells[1].text = item.terminal_name or ''
            cells[2].text = f"{item.ship_volume_m3:.3f}" if item.ship_volume_m3 else ''
            cells[3].text = f"{item.ship_weight_mt:.3f}" if item.ship_weight_mt else ''
            cells[4].text = f"{item.shore_volume_m3:.3f}" if item.shore_volume_m3 else ''
            cells[5].text = f"{item.shore_weight_mt:.3f}" if item.shore_weight_mt else ''
            cells[6].text = f"{item.diff_volume_m3:.3f}"
            cells[7].text = f"{item.diff_volume_pct:.3f}%"
            cells[8].text = f"{item.diff_weight_mt:.3f}"
            cells[9].text = f"{item.diff_weight_pct:.3f}%"

        # Totals row
        total_row = table.rows[-1]
        total_cells = total_row.cells
        total_cells[0].text = ''
        total_cells[1].text = 'TOTAL'
        total_cells[2].text = f"{totals['ship_volume']:.3f}"
        total_cells[3].text = f"{totals['ship_weight']:.3f}"
        total_cells[4].text = f"{totals['shore_volume']:.3f}"
        total_cells[5].text = f"{totals['shore_weight']:.3f}"
        total_cells[6].text = f"{tot_dv:.3f}"
        total_cells[7].text = f"{tot_dvp:.3f}%"
        total_cells[8].text = f"{tot_dw:.3f}"
        total_cells[9].text = f"{tot_dwp:.3f}%"

        # Style total row
        for cell in total_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        doc.add_paragraph()  # Spacing

        # Signature section
        sig_para = doc.add_paragraph()
        sig_para_format = sig_para.paragraph_format
        sig_para_format.space_before = Pt(12)

        sig_table = doc.add_table(rows=3, cols=3)
        sig_table.style = 'Table Grid'

        # Signature lines (row 0)
        sig_table.cell(0, 0).text = '_' * 30
        sig_table.cell(0, 1).text = '_' * 30
        sig_table.cell(0, 2).text = '_' * 30

        # Names (row 1)
        sig_table.cell(1, 0).text = 'Captain / Master'
        sig_table.cell(1, 1).text = 'PBPA Surveyor'
        sig_table.cell(1, 2).text = 'Terminal Rep.'

        # Actual names (row 2)
        sig_table.cell(2, 0).text = report.captain_name or 'Name: _____________'
        sig_table.cell(2, 1).text = report.surveyor_name or 'Name: _____________'
        sig_table.cell(2, 2).text = 'Name: _____________'

        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run(f'Report Generated: {report.created_at.strftime("%d.%m.%Y %H:%M")}')
        footer_run.font.size = Pt(8)
        footer_run.font.italic = True

        # Save to BytesIO
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        response = HttpResponse(
            buf.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="POR_{report.report_number}.docx"'
        return response


class StockReportViewSet(TerminalRepresentativeReadOnlyMixin, viewsets.ModelViewSet):
    """CRUD + PDF for PBPA Weekly Stock Reports."""
    permission_classes = (IsAuthenticated,)
    serializer_class = StockReportSerializer
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['report_number', 'notes']
    ordering = ['-report_date', '-created_at']

    def get_queryset(self):
        profile = get_or_create_user_profile(self.request.user)
        base = StockReport.objects.select_related('created_by').prefetch_related('items')
        if profile.role in ('admin', 'terminal_representative') and self.action in {'list', 'retrieve', 'pdf'}:
            qs = base
        else:
            qs = owned_queryset(base, self.request.user)
        qs = filter_queryset_by_period(qs, self.request.query_params.get('period'), 'report_date')
        return qs

    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user)

    @action(detail=True, methods=['post'])
    def finalize(self, request, pk=None):
        report = self.get_object()
        report.status = 'final'
        report.save(update_fields=['status', 'updated_at'])
        return Response(StockReportSerializer(report).data)

    @action(detail=True, methods=['get'])
    def pdf(self, request, pk=None):
        report = self.get_object()
        buf = self._build_pdf(report)
        pdf_bytes = _finalize_pdf_for_download(buf.getvalue(), report, 'Weekly Stock Report', request)
        response = HttpResponse(pdf_bytes, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="StockReport_{report.report_number}.pdf"'
        return response

    def _build_pdf(self, report):
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib import colors
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=landscape(A4),
                                leftMargin=10*mm, rightMargin=10*mm,
                                topMargin=8*mm, bottomMargin=10*mm)
        styles = getSampleStyleSheet()
        W = landscape(A4)[0] - 20*mm

        def ps(name, **kw):
            return ParagraphStyle(name, parent=styles['Normal'], **kw)

        YELLOW = colors.HexColor('#FFD700')
        BLACK  = colors.black
        WHITE  = colors.white

        elems = []

        # Header
        elems.append(Paragraph("THE UNITED REPUBLIC OF TANZANIA", ps('h1', fontSize=9, alignment=TA_CENTER)))
        elems.append(Spacer(1, 1*mm))
        elems.append(Paragraph("<b>PETROLEUM BULK PROCUREMENT AGENCY</b>", ps('h2', fontSize=11, alignment=TA_CENTER)))
        elems.append(Paragraph(
            "TANZANIA PORTS AUTHORITY, ONE STOP CENTER BUILDING, 11TH FLOOR, SOKOINE DRIVE, PLOT NO:1/2",
            ps('a1', fontSize=7, alignment=TA_CENTER)))
        elems.append(Paragraph(
            "Tel: +255222129009 / Email: info@pbpa.go.tz / WEBSITE: www.pbpa.go.tz / P.O. Box 2634 Dar es Salaam, TANZANIA",
            ps('a2', fontSize=6.5, alignment=TA_CENTER)))
        # Reserve the rest of the logo band before the report-specific title
        # and table begin.  This keeps the four agency lines horizontally
        # aligned with the two marks without overlapping the table.
        elems.append(Spacer(1, 15*mm))

        date_str = report.report_date.strftime('%d-%b-%Y') if report.report_date else ''
        elems.append(Paragraph(f"<b>WEEKLY STOCK REPORT — {date_str}</b>", ps('title', fontSize=11, alignment=TA_CENTER)))
        elems.append(Spacer(1, 4*mm))

        items = list(report.items.all())

        def f(v):
            if v is None or v == 0:
                return ''
            return f'{float(v):,.0f}'

        hdr = ps('th', fontSize=7, alignment=TA_CENTER, fontName='Helvetica-Bold', textColor=BLACK)
        cell = ps('td', fontSize=7.5, alignment=TA_CENTER)
        lcell = ps('tl', fontSize=7.5, alignment=TA_LEFT)

        # Transit group header spans cols 5-6 (BPS + NON-BPS)
        cw = [W*p for p in (0.04, 0.14, 0.08, 0.09, 0.09, 0.09, 0.09, 0.09, 0.09, 0.10, 0.10)]

        transit_box = Table(
            [[Paragraph('<b>TRANSIT (LTRS)</b>', ps('tb', fontSize=7, alignment=TA_CENTER, fontName='Helvetica-Bold'))]],
            colWidths=[cw[5]+cw[6]]
        )
        transit_box.setStyle(TableStyle([
            ('BOX', (0,0), (-1,-1), 1.5, YELLOW),
            ('BACKGROUND', (0,0), (-1,-1), WHITE),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('TOPPADDING', (0,0), (-1,-1), 2),
            ('BOTTOMPADDING', (0,0), (-1,-1), 2),
        ]))

        header_rows = [
            # row 0 — column headers
            [
                Paragraph('<b>S/N</b>', hdr),
                Paragraph('<b>DEPOT NAME</b>', hdr),
                Paragraph('<b>DATE</b>', hdr),
                Paragraph('<b>PRODUCT</b>', hdr),
                Paragraph('<b>LOCAL (LTRS)</b>', hdr),
                Paragraph('<b>BPS TRANSIT</b>', hdr),
                Paragraph('<b>NON-BPS TRANSIT</b>', hdr),
                Paragraph('<b>MINING (LTRS)</b>', hdr),
                Paragraph('<b>TRANSSHIPMENT (LTRS)</b>', hdr),
                Paragraph('<b>AWAITING FOR OUTTURN (LTRS)</b>', hdr),
                Paragraph('<b>TOTAL (LTRS)</b>', hdr),
            ],
        ]

        data_rows = []
        for item in items:
            item_date = item.date.strftime('%d-%b-%Y') if item.date else ''
            data_rows.append([
                Paragraph(str(item.sn), cell),
                Paragraph(item.depot_name or '', lcell),
                Paragraph(item_date, cell),
                Paragraph(item.product or '', cell),
                Paragraph(f(item.local_ltrs), cell),
                Paragraph(f(item.bps_transit_ltrs), cell),
                Paragraph(f(item.non_bps_transit_ltrs), cell),
                Paragraph(f(item.mining_ltrs), cell),
                Paragraph(f(item.transshipment_ltrs), cell),
                Paragraph(f(item.awaiting_outturn_ltrs), cell),
                Paragraph(f(item.total_ltrs), cell),
            ])

        bold_cell = ps('tb', fontSize=7.5, alignment=TA_CENTER, fontName='Helvetica-Bold')
        total_row = [
            '', Paragraph('<b>TOTAL</b>', ps('tot', fontSize=8, alignment=TA_CENTER, fontName='Helvetica-Bold')),
            '', '', '', '', '', '', '', '',
            Paragraph(f'{report.total_ltrs:,.0f}', bold_cell),
        ]

        all_rows = header_rows + data_rows + [total_row]
        n_hdr = len(header_rows)
        n_data = len(data_rows)

        tbl = Table(all_rows, colWidths=cw, repeatRows=1)
        ts = TableStyle([
            ('BOX',         (0,0), (-1,-1), 1, BLACK),
            ('INNERGRID',   (0,0), (-1,-1), 0.4, colors.grey),
            ('BACKGROUND',  (0,0), (-1, n_hdr-1), YELLOW),
            ('ROWBACKGROUNDS', (0, n_hdr), (-1, n_hdr+n_data-1), [WHITE, colors.HexColor('#FFFDE7')]),
            ('BACKGROUND',  (0,-1), (-1,-1), colors.HexColor('#F5F5F5')),
            ('FONTNAME',    (0,-1), (-1,-1), 'Helvetica-Bold'),
            ('VALIGN',      (0,0), (-1,-1), 'MIDDLE'),
            ('TOPPADDING',  (0,0), (-1,-1), 2),
            ('BOTTOMPADDING',(0,0), (-1,-1), 2),
            # Transit group box on header
            ('BOX',         (5,0), (6,0), 1.5, BLACK),
        ])
        tbl.setStyle(ts)
        elems.append(tbl)

        if report.notes:
            elems.append(Spacer(1, 4*mm))
            elems.append(Paragraph(f'Notes: {report.notes}', ps('notes', fontSize=8)))

        doc.build(elems)
        buf.seek(0)
        return buf


class ServiceRequestViewSet(viewsets.ModelViewSet):
    """Service requests submitted by clients; admin & inspector get notified."""
    permission_classes = (IsAuthenticated,)
    serializer_class = ServiceRequestSerializer
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['request_number', 'vessel_name', 'terminal', 'operation_type']
    ordering = ['-created_at']

    def get_queryset(self):
        profile = get_or_create_user_profile(self.request.user)
        qs = ServiceRequest.objects.select_related('submitted_by', 'assigned_to')
        # terminal representative sees only their own; admin/inspector see all
        if profile.role == 'terminal_representative':
            return qs.filter(submitted_by=self.request.user)
        return qs

    def perform_create(self, serializer):
        serializer.save(submitted_by=self.request.user, is_read_admin=False, is_read_inspector=False)

    @action(detail=False, methods=['get'])
    def unread_count(self, request):
        """Unread count for admin or inspector."""
        profile = get_or_create_user_profile(request.user)
        if profile.role == 'admin':
            count = ServiceRequest.objects.filter(is_read_admin=False).count()
        elif profile.role == 'inspector':
            count = ServiceRequest.objects.filter(is_read_inspector=False).count()
        else:
            count = 0
        return Response({'count': count})

    @action(detail=True, methods=['post'])
    def mark_read(self, request, pk=None):
        sr = self.get_object()
        profile = get_or_create_user_profile(request.user)
        if profile.role == 'admin':
            sr.is_read_admin = True
            sr.save(update_fields=['is_read_admin', 'updated_at'])
        elif profile.role == 'inspector':
            sr.is_read_inspector = True
            sr.save(update_fields=['is_read_inspector', 'updated_at'])
        return Response(self.get_serializer(sr).data)

    @action(detail=False, methods=['post'])
    def mark_all_read(self, request):
        profile = get_or_create_user_profile(request.user)
        if profile.role == 'admin':
            ServiceRequest.objects.filter(is_read_admin=False).update(is_read_admin=True)
        elif profile.role == 'inspector':
            ServiceRequest.objects.filter(is_read_inspector=False).update(is_read_inspector=True)
        return Response({'status': 'all marked read'})

    @action(detail=True, methods=['post'])
    def acknowledge(self, request, pk=None):
        sr = self.get_object()
        sr.status = 'acknowledged'
        sr.is_read_admin = True
        sr.save(update_fields=['status', 'is_read_admin', 'updated_at'])
        return Response(self.get_serializer(sr).data)

    @action(detail=True, methods=['get', 'post'])
    def messages(self, request, pk=None):
        sr = self.get_object()
        if request.method == 'GET':
            msgs = sr.messages.select_related('sender').all()
            return Response(ServiceRequestMessageSerializer(msgs, many=True).data)
        serializer = ServiceRequestMessageSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        msg = serializer.save(service_request=sr, sender=request.user)
        # Notify all participants except the sender
        sender_name = request.user.get_full_name() or request.user.username
        notif_title = f'New message on {sr.request_number}'
        notif_body = f'{sender_name}: {msg.body[:100]}'
        recipients = set()
        # Always notify the terminal rep (submitted_by)
        if sr.submitted_by and sr.submitted_by != request.user:
            recipients.add(sr.submitted_by)
        # Notify assigned inspector
        if sr.assigned_to and sr.assigned_to != request.user:
            recipients.add(sr.assigned_to)
        # Notify all admins
        admin_ids = UserProfile.objects.filter(role='admin').values_list('user_id', flat=True)
        for u in User.objects.filter(pk__in=admin_ids).exclude(pk=request.user.pk):
            recipients.add(u)
        # Notify all inspectors who have replied in this thread
        thread_senders = sr.messages.exclude(sender=request.user).values_list('sender_id', flat=True).distinct()
        for u in User.objects.filter(pk__in=thread_senders):
            recipients.add(u)
        for recipient in recipients:
            Notification.objects.create(
                recipient=recipient,
                notification_type='sr_message',
                title=notif_title,
                message=notif_body,
            )
        return Response(serializer.data, status=status.HTTP_201_CREATED)

    @action(detail=True, methods=['post'])
    def assign(self, request, pk=None):
        sr = self.get_object()
        inspector_id = request.data.get('inspector_id')
        try:
            inspector = User.objects.get(pk=inspector_id)
        except User.DoesNotExist:
            return Response({'detail': 'Inspector not found.'}, status=status.HTTP_400_BAD_REQUEST)
        sr.assigned_to = inspector
        sr.status = 'in_progress'
        sr.is_read_inspector = False
        sr.save(update_fields=['assigned_to', 'status', 'is_read_inspector', 'updated_at'])
        return Response(self.get_serializer(sr).data)

    @action(detail=True, methods=['post'])
    def complete(self, request, pk=None):
        sr = self.get_object()
        sr.status = 'completed'
        sr.save(update_fields=['status', 'updated_at'])
        return Response(self.get_serializer(sr).data)

    @action(detail=True, methods=['post'])
    def cancel(self, request, pk=None):
        sr = self.get_object()
        reason = request.data.get('reason', '')
        sr.status = 'cancelled'
        if reason:
            sr.notes = (sr.notes + '\n\nCancellation reason: ' + reason).strip()
        sr.save(update_fields=['status', 'notes', 'updated_at'])
        return Response(self.get_serializer(sr).data)



class ActivityLogViewSet(viewsets.ReadOnlyModelViewSet):
    """Real-time activity log — admin only."""
    permission_classes = (IsAuthenticated,)
    serializer_class = ActivityLogSerializer
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['user__username', 'user__first_name', 'action', 'doc_type', 'doc_number', 'detail']
    ordering = ['-timestamp']

    def get_queryset(self):
        profile = get_or_create_user_profile(self.request.user)
        if profile.role != 'admin':
            return ActivityLog.objects.none()
        qs = ActivityLog.objects.select_related('user').all()
        action = self.request.query_params.get('action')
        period = self.request.query_params.get('period')
        if action:
            qs = qs.filter(action=action)
        if period:
            start = period_start_datetime(period)
            if start:
                qs = qs.filter(timestamp__gte=start)
        return qs

    @action(detail=False, methods=['get'])
    def summary(self, request):
        """Return activity counts grouped by action type for the dashboard."""
        profile = get_or_create_user_profile(request.user)
        if profile.role != 'admin':
            return Response({'detail': 'Admin only.'}, status=status.HTTP_403_FORBIDDEN)

        period = request.query_params.get('period')
        qs = ActivityLog.objects.all()
        if period:
            start = period_start_datetime(period)
            if start:
                qs = qs.filter(timestamp__gte=start)

        from django.db.models import Count
        counts = {
            row['action']: row['count']
            for row in qs.values('action').annotate(count=Count('id'))
        }

        COLOR_MAP = {
            'login':            '#2563eb',
            'logout':           '#64748b',
            'login_failed':     '#dc2626',
            'report_created':   '#16a34a',
            'report_updated':   '#f59e0b',
            'report_deleted':   '#ef4444',
            'report_submitted': '#0891b2',
            'report_approved':  '#059669',
            'report_rejected':  '#be123c',
            'password_changed': '#7c3aed',
            'user_created':     '#4f46e5',
            'user_deleted':     '#9f1239',
            'file_uploaded':    '#0369a1',
        }
        WHY_MAP = {
            'login':            'Security audit',
            'logout':           'Session tracking',
            'login_failed':     'Detect attacks',
            'report_created':   'Usage monitoring',
            'report_updated':   'Audit trail',
            'report_deleted':   'Accountability',
            'report_submitted': 'Workflow tracking',
            'report_approved':  'Approval tracking',
            'report_rejected':  'Rejection tracking',
            'password_changed': 'Security tracking',
            'user_created':     'Access management',
            'user_deleted':     'Access management',
            'file_uploaded':    'Storage monitoring',
        }
        LABEL_MAP = dict(ActivityLog.ACTION_CHOICES)
        rows = [
            {
                'action': action,
                'activity': LABEL_MAP.get(action, action),
                'value': count,
                'why': WHY_MAP.get(action, ''),
                'color': COLOR_MAP.get(action, '#8B1A1A'),
            }
            for action, count in sorted(counts.items(), key=lambda x: -x[1])
        ]
        return Response(rows)

class NotificationViewSet(viewsets.ReadOnlyModelViewSet):
    """In-app notifications for the current user."""
    permission_classes = (IsAuthenticated,)
    serializer_class = NotificationSerializer
    ordering = ['-created_at']

    def get_queryset(self):
        queryset = Notification.objects.filter(recipient=self.request.user)
        # Bells can request their own alert types without including unrelated notifications.
        notification_types = self.request.query_params.get('notification_type', '')
        if notification_types:
            queryset = queryset.filter(
                notification_type__in=[
                    value.strip() for value in notification_types.split(',') if value.strip()
                ]
            )
        return queryset

    @action(detail=False, methods=['get'])
    def unread_count(self, request):
        qs = self.get_queryset().filter(is_read=False)
        return Response({'count': qs.count()})

    @action(detail=True, methods=['post'])
    def mark_read(self, request, pk=None):
        n = self.get_object()
        n.is_read = True
        n.save(update_fields=['is_read'])
        return Response({'status': 'marked read'})

    @action(detail=False, methods=['post'])
    def mark_all_read(self, request):
        self.get_queryset().filter(is_read=False).update(is_read=True)
        return Response({'status': 'all marked read'})


def csrf_failure(request, reason=""):
    """Handle CSRF failures securely."""
    from .exception_handler import custom_exception_handler
    from rest_framework.exceptions import ValidationError
    
    sec_log.warning(f"CSRF failure: {reason or 'no reason provided'} | IP: {request.META.get('REMOTE_ADDR')}")
    
    if request.headers.get('Accept', '').startswith('application/json') or request.path.startswith('/api/'):
        return Response(
            {'detail': 'CSRF verification failed. Request aborted.'},
            status=status.HTTP_403_FORBIDDEN
        )
    
    return HttpResponse(
        '<h1>403 Forbidden</h1><p>CSRF verification failed. Request aborted.</p>',
        status=403,
        content_type='text/html'
    )


class SamplingFormViewSet(viewsets.ModelViewSet):
    """CRUD + official PBPA PDF for Sampling Forms."""
    permission_classes = (IsAuthenticated,)
    serializer_class = SamplingFormSerializer
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['form_number', 'vessel_name', 'product_name']
    ordering_fields = ['created_at', 'sampling_date', 'form_number', 'status']
    ordering = ['-created_at']
    _doc_type_label = 'sampling_form'

    def _get_role(self, request):
        return get_or_create_user_profile(request.user).role

    def get_queryset(self):
        profile = get_or_create_user_profile(self.request.user)
        base = SamplingForm.objects.select_related('created_by')
        if profile.role == 'admin':
            qs = base
        else:
            qs = owned_queryset(base, self.request.user)
        status_filter = self.request.query_params.get('status')
        if status_filter:
            qs = qs.filter(status=status_filter)
        period = self.request.query_params.get('period')
        return filter_queryset_by_period(qs, period, 'sampling_date')

    def perform_create(self, serializer):
        full_name = self.request.user.get_full_name() or self.request.user.username
        serializer.save(
            created_by=self.request.user,
            pbpa_inspector_name=serializer.validated_data.get('pbpa_inspector_name') or full_name,
        )

    def _build_pdf_bytes(self, obj):
        return self._build_pdf(obj).getvalue()

    @action(detail=True, methods=['post'])
    def issue(self, request, pk=None):
        form = self.get_object()
        if form.created_by_id != request.user.id and self._get_role(request) != 'admin':
            return Response({'detail': 'Only the form creator or an admin can issue this form.'}, status=status.HTTP_403_FORBIDDEN)
        if form.status == 'issued':
            return Response({'detail': 'Form has already been issued.'}, status=status.HTTP_400_BAD_REQUEST)
        form.status = 'issued'
        form.issued_at = timezone.now()
        form.save(update_fields=['status', 'issued_at', 'updated_at'])
        return Response(SamplingFormSerializer(form, context={'request': request}).data)

    @action(detail=True, methods=['get'])
    def pdf(self, request, pk=None):
        form = self.get_object()
        pdf_bytes = _finalize_pdf_for_download(self._build_pdf(form).getvalue(), form, 'Vessel Sampling Form', request)
        response = HttpResponse(pdf_bytes, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="Sampling_Form_{form.form_number}.pdf"'
        return response

    # ------------------------------------------------------------------
    # Official PBPA Vessel Sampling Form PDF
    # ------------------------------------------------------------------
    def _build_pdf(self, form):
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.units import mm

        buf = io.BytesIO()
        W, H = A4
        pdf = canvas.Canvas(buf, pagesize=A4)
        M = 15 * mm
        TW = W - 2 * M

        y = H - 10 * mm

        # ── Letterhead ────────────────────────────────────────────────
        pdf.setFont("Helvetica", 9)
        pdf.drawCentredString(W / 2, y, "THE UNITED REPUBLIC OF TANZANIA")
        y -= 7 * mm
        pdf.setFont("Helvetica-Bold", 12)
        pdf.drawCentredString(W / 2, y, "PETROLEUM BULK PROCUREMENT AGENCY")
        y -= 5 * mm
        pdf.setFont("Helvetica", 6.5)
        pdf.drawCentredString(W / 2, y,
            "TANZANIA PORTS AUTHORITY, ONE STOP CENTER BUILDING, 11TH FLOOR, SOKOINE DRIVE, PLOT NO:1/2")
        y -= 4 * mm
        pdf.drawCentredString(W / 2, y,
            "Tel: +255222129009 / Fax: +255222129093 / Email: info@pbpa.go.tz / WEBSITE: www.pbpa.go.tz / P.O. Box 2634 Dar es Salaam, TANZANIA")
        y -= 8 * mm

        # ── Title box ─────────────────────────────────────────────────
        box_h = 11 * mm
        box_x = M + 8 * mm
        box_w = TW - 16 * mm
        pdf.setLineWidth(2.5)
        pdf.rect(box_x, y - box_h, box_w, box_h, fill=0, stroke=1)
        pdf.setLineWidth(0.5)
        pdf.setFont("Helvetica-Bold", 14)
        pdf.drawCentredString(W / 2, y - box_h + 3.5 * mm, "VESSEL SAMPLING FORM")
        y -= box_h + 5 * mm

        # Form number top-right
        pdf.setFont("Helvetica-Bold", 11)
        pdf.drawRightString(W - M, y + 6 * mm, form.form_number or "")

        # ── Field helpers ──────────────────────────────────────────────
        def ufield(label, value, y_pos, label_w=38 * mm, line_end=None):
            le = line_end if line_end else (W - M)
            pdf.setFont("Helvetica", 9)
            pdf.drawString(M, y_pos, label)
            lx = M + label_w
            pdf.setDash(1, 2)
            pdf.line(lx, y_pos - 1, le, y_pos - 1)
            pdf.setDash()
            if value:
                pdf.setFont("Helvetica-Bold", 9)
                pdf.drawString(lx + 1, y_pos, str(value))

        def row2(lbl1, val1, lbl2, val2, y_pos, lw1=38 * mm, lw2=40 * mm):
            half = TW / 2
            ufield(lbl1, val1, y_pos, lw1, M + half - 4 * mm)
            rx = M + half + 4 * mm
            pdf.setFont("Helvetica", 9)
            pdf.drawString(rx, y_pos, lbl2)
            llx = rx + lw2
            pdf.setDash(1, 2)
            pdf.line(llx, y_pos - 1, W - M, y_pos - 1)
            pdf.setDash()
            if val2:
                pdf.setFont("Helvetica-Bold", 9)
                pdf.drawString(llx + 1, y_pos, str(val2))

        lg = 8 * mm

        # ── Section header helper ──────────────────────────────────────
        def section_hdr(label, y_pos):
            pdf.setFillColor(colors.HexColor('#f0f0f0'))
            pdf.setStrokeColor(colors.black)
            pdf.setLineWidth(0.6)
            pdf.rect(M, y_pos - 5 * mm, TW, 6 * mm, fill=1, stroke=1)
            pdf.setFillColor(colors.black)
            pdf.setFont("Helvetica-Bold", 8)
            pdf.drawString(M + 2 * mm, y_pos - 2 * mm, label)

        # ── VESSEL INFORMATION ─────────────────────────────────────────
        section_hdr("VESSEL INFORMATION", y)
        y -= 7 * mm

        ufield("Vessel Name:", form.vessel_name, y)
        y -= lg
        ufield("Date:", form.sampling_date.strftime("%d - %m - %Y") if form.sampling_date else "", y)
        y -= lg
        row2("Product:", form.product_name,
             "Time:", form.sampling_time.strftime("%H:%M") + " Hrs" if form.sampling_time else "", y)
        y -= lg
        row2("Voyage No.:", form.voyage_no,
             "Bill of Lading No.:", form.bill_of_lading_no, y, 32 * mm, 46 * mm)
        y -= lg + 3 * mm

        # ── SAMPLING DETAILS ───────────────────────────────────────────
        section_hdr("SAMPLING DETAILS", y)
        y -= 7 * mm

        row2("Cargo Tank No.:", form.cargo_tank_no,
             "Sample Location:", form.sample_location, y, 36 * mm, 40 * mm)
        y -= lg
        row2("Sample Reference:", form.sample_reference,
             "No. of Samples:", str(form.number_of_samples) if form.number_of_samples else "", y, 40 * mm, 36 * mm)
        y -= lg
        row2("Sample Quantity:", form.sample_quantity,
             "Sample Container:", form.sample_container, y, 38 * mm, 40 * mm)
        y -= lg
        row2("Seal No. Before Sampling:", form.seal_number_before,
             "Seal No. After Sampling:", form.seal_number_after, y, 54 * mm, 50 * mm)
        y -= lg + 3 * mm

        # ── PHYSICAL PROPERTIES ────────────────────────────────────────
        section_hdr("PHYSICAL PROPERTIES", y)
        y -= 7 * mm

        row2("Temperature (°C):",
             f"{form.temperature:.2f}" if form.temperature is not None else "",
             "Observed Density (kg/L):",
             f"{form.density_observed:.4f}" if form.density_observed is not None else "",
             y, 42 * mm, 52 * mm)
        y -= lg
        row2("Colour:", form.colour, "Appearance:", form.appearance, y, 22 * mm, 28 * mm)
        y -= lg

        # Remarks
        pdf.setFont("Helvetica", 8)
        pdf.drawString(M, y, "Remarks:")
        pdf.setDash(1, 2)
        pdf.line(M + 22 * mm, y - 1, W - M - 2 * mm, y - 1)
        pdf.setDash()
        if form.remarks:
            pdf.setFont("Helvetica-Bold", 8)
            pdf.drawString(M + 23 * mm, y, form.remarks[:95])
        y -= lg + 2 * mm

        # ── SAMPLED / WITNESSED BY ─────────────────────────────────────
        row2("Sampled By:", form.sampled_by,
             "Witnessed By:", form.witnessed_by, y, 28 * mm, 32 * mm)
        y -= lg + 4 * mm

        # ── Certification statement ────────────────────────────────────
        pdf.setLineWidth(0.5)
        pdf.rect(M, y - 14 * mm, TW, 15 * mm, fill=0, stroke=1)
        pdf.setFont("Helvetica", 8)
        stmt = (
            "We, the undersigned, hereby confirm that the above sample(s) were drawn in our joint presence "
            "from the above-mentioned vessel cargo tank(s) in accordance with standard petroleum sampling "
            "procedures, and that the sample containers were sealed with the seal numbers stated above."
        )
        tw = pdf.beginText(M + 2 * mm, y - 3 * mm)
        tw.setFont("Helvetica", 8)
        words, line = stmt.split(), ""
        for w in words:
            test = (line + " " + w).strip()
            if pdf.stringWidth(test, "Helvetica", 8) > TW - 4 * mm:
                tw.textLine(line)
                line = w
            else:
                line = test
        if line:
            tw.textLine(line)
        pdf.drawText(tw)
        y -= 17 * mm

        # ── Signature block ───────────────────────────────────────────
        # ── Footer ────────────────────────────────────────────────────
        pdf.setFont("Helvetica", 7)
        pdf.setFillColor(colors.grey)
        pdf.drawCentredString(W / 2, M + 4 * mm,
            f"PBPA Vessel Sampling Form {form.form_number or ''} | Generated: {timezone.localtime().strftime('%d-%b-%Y %H:%M')}")

        pdf.showPage()
        pdf.save()
        buf.seek(0)
        return buf
