"""
Document generators that fill the original PBPA Word templates exactly,
preserving every table, paragraph, and layout from the source .docx files.
"""
import os
from io import BytesIO
from copy import deepcopy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .calculations import ShoreTankCalculationEngine, ASTMD1250Calculator

BASE = os.path.dirname(os.path.abspath(__file__))
TEMPLATES = {
    'shore_tank':   os.path.join(BASE, '..', 'SHORE TANK CALCULATIONS.docx'),
    'dip_ticket':   os.path.join(BASE, '..', 'DIP TICKET.docx'),
    'product_receipt': os.path.join(BASE, '..', 'PRODUCT RECEIPT CERTIFICATE.docx'),
    'seal_isolation':  os.path.join(BASE, '..', 'SEAL AND ISOLATION.docx'),
}


# ── helpers ──────────────────────────────────────────────────────────────────

def _set_unique_cell(table, row_idx, unique_col_idx, text):
    """
    Write to the Nth unique (non-merged) cell in a row.
    unique_col_idx: 0=label, 1=first data cell, 2=second data cell, etc.
    """
    try:
        row = table.rows[row_idx]
        seen = {}
        unique_cells = []
        for cell in row.cells:
            cid = id(cell._tc)
            if cid not in seen:
                seen[cid] = True
                unique_cells.append(cell)
        cell = unique_cells[unique_col_idx]
        para = cell.paragraphs[0]
        para.clear()
        para.add_run(str(text) if text is not None else '')
    except (IndexError, Exception):
        pass


def _set_cell(table, row, col, text, bold=False, align=None):
    """Write text into a table cell by absolute column index."""
    try:
        cell = table.rows[row].cells[col]
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(str(text) if text is not None else '')
        run.bold = bold
        if align:
            para.alignment = align
    except IndexError:
        pass


def _set_para(doc, para_index, text):
    """Replace the text of a paragraph by index."""
    try:
        p = doc.paragraphs[para_index]
        for run in p.runs:
            run.text = ''
        if p.runs:
            p.runs[0].text = text
        else:
            p.add_run(text)
    except IndexError:
        pass

def _f(v, d=3):
    """Format a number to d decimal places, empty string if None."""
    if v is None or v == '':
        return ''
    try:
        return f'{float(v):.{d}f}'
    except (TypeError, ValueError):
        return str(v)


def _f4(v):
    return _f(v, 4)


def _load(key):
    path = os.path.abspath(TEMPLATES[key])
    if not os.path.exists(path):
        raise FileNotFoundError(f'Template not found: {path}')
    return Document(path)


def _save(doc):
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Shore Tank Calculation ────────────────────────────────────────────────────

def generate_shore_tank_document(shore_calc):
    """
    Fill SHORE TANK CALCULATIONS.docx template.

    Template structure (verified from docx):
      Para [0]  THE UNITED REPUBLIC OF TANZANIA
      Para [1]  PETROLEUM BULK PROCUREMENT AGENCY
      Para [7]  TERMINAL REPRESENTATIVE  (signature block)

      Table 0 (4r x 4c) – header
        R0: VESSEL NAME | | Date |
        R1: PRODUCT     | | Vessel Density |
        R2: TERMINAL    | | Vessel Temperature |

      Table 1 (19r x 12c) – measurements (cols 1-2 = tank1 initial/final, 3-4 = tank2, etc.)
        R0:  Tank No. headers
        R1:  PARTICULARS | INITIAL | L.DISP/PROV/FINAL | ...
        R2:  Overall Dip(mm)
        R3:  Product Dip(mm)
        R4:  Water Dip(mm)
        R5:  Tank Temperature
        R6:  Specific Gravity(SG)
        R7:  Sample Temperature
        R8:  Density @20
        R9:  VCF
        R10: WCF
        R11: Gross Observed Volume
        R12: Roof Displacement Vol
        R13: Water Volume
        R14: Net Obs Volume
        R15: Standard Vol@20
        R16: Standard Vol@20 Received
        R17: Weight in Air (Mt)
        R18: Weight in Air Received (Mt)

      Table 2 (7r x 4c) – summary
        R0: STATUS | Obs Vol | M3 | M/Tons
        R1: TERMINAL
        R2: VESSEL
        R3: METER QUANTITY
        R4: DIFFERENCE(TERMINAL VS VESSEL)
        R6: INLET MANIFOLD SEAL

      Table 3 (4r x 6c) – Terminal Rep signatures
      Table 4 (4r x 6c) – PBPA Inspector signatures
    """
    doc = _load('shore_tank')
    items = list(shore_calc.tank_items.all())

    # ── Table 0: header ──
    t0 = doc.tables[0]
    _set_cell(t0, 0, 1, shore_calc.vessel_name)
    _set_cell(t0, 0, 3, str(shore_calc.calculation_date))
    _set_cell(t0, 1, 1, shore_calc.product_name)
    _set_cell(t0, 1, 3, _f4(shore_calc.vessel_density_kg_m3) if shore_calc.vessel_density_kg_m3 else '')
    _set_cell(t0, 2, 1, shore_calc.terminal)
    _set_cell(t0, 2, 3, _f(shore_calc.vessel_temperature_c, 1) if shore_calc.vessel_temperature_c else '')

    # ── Table 1: measurements ──
    # Unique cell layout per data row (rows 2-15, 17):
    #   unique[0]=label, unique[1]=tank1_initial, unique[2]=tank1_final,
    #   unique[3]=tank2_initial, unique[4]=tank2_final
    # Row 16 (Std Vol Received): unique[0]=label, unique[1]=tank1_recv, unique[3]=tank2_recv
    # Row 18 (Wt Received):      unique[0]=label, unique[1]=tank1_recv, unique[3]=tank2_recv
    t1 = doc.tables[1]

    # Row 0: single merged cell — write tank numbers as combined label
    tank_labels = '  |  '.join(item.tank_no or f'Tank {i+1}' for i, item in enumerate(items[:2]))
    _set_unique_cell(t1, 0, 0, tank_labels)

    # Data rows per tank pair (template supports 2 tanks in unique cells 1-4)
    STANDARD_ROWS = [
        # (row_idx, attr_initial, attr_final)
        (2,  'overall_dip_initial_mm',       'overall_dip_final_mm'),
        (3,  'product_dip_initial_mm',        'product_dip_final_mm'),
        (4,  'water_dip_initial_mm',          'water_dip_final_mm'),
        (5,  'tank_temperature_initial_c',    'tank_temperature_final_c'),
        (6,  'density_initial_kg_l',          'density_final_kg_l'),
        (7,  'sample_temperature_initial_c',  'sample_temperature_final_c'),
        (11, 'gross_observed_initial_m3',     'gross_observed_final_m3'),
        (12, 'roof_displacement_initial_m3',  'roof_displacement_final_m3'),
        (13, 'water_volume_initial_m3',       'water_volume_final_m3'),
        (14, 'net_observed_initial_m3',       'net_observed_final_m3'),
        (15, 'standard_volume_initial_m3',    'standard_volume_final_m3'),
        (17, 'weight_air_initial_mt',         'weight_air_final_mt'),
    ]

    for idx, item in enumerate(items[:2]):
        ui = 1 + idx * 2   # unique cell index for initial
        uf = 2 + idx * 2   # unique cell index for final

        for row_idx, attr_i, attr_f in STANDARD_ROWS:
            _set_unique_cell(t1, row_idx, ui, _f(getattr(item, attr_i, None)))
            _set_unique_cell(t1, row_idx, uf, _f(getattr(item, attr_f, None)))

        # ASTM rows (4 decimal places)
        vcf_i = item.effective_vcf_initial
        vcf_f = item.effective_vcf_final
        wcf_i = item.effective_wcf_initial
        wcf_f = item.effective_wcf_final
        _set_unique_cell(t1, 8,  ui, _f4(item.density_initial_kg_l))
        _set_unique_cell(t1, 8,  uf, _f4(item.density_final_kg_l))
        _set_unique_cell(t1, 9,  ui, _f4(vcf_i))
        _set_unique_cell(t1, 9,  uf, _f4(vcf_f))
        _set_unique_cell(t1, 10, ui, _f4(wcf_i))
        _set_unique_cell(t1, 10, uf, _f4(wcf_f))

        # Received rows have different merge patterns per row:
        # Row 16 unique: [0=label, 1=t1, 3=t2, 6=t3, 8=t4, 11=t5]
        # Row 18 unique: [0=label, 1=t1, 2=t2, 5=t3, 8=t4, 10=t5]
        recv16_idx = [1, 3, 6, 8, 11]
        recv18_idx = [1, 2, 5, 8, 10]
        if idx < len(recv16_idx):
            _set_unique_cell(t1, 16, recv16_idx[idx], _f(item.received_standard_volume_m3))
        if idx < len(recv18_idx):
            _set_unique_cell(t1, 18, recv18_idx[idx], _f(item.received_weight_air_mt))

    # ── Table 2: summary ──
    t2 = doc.tables[2]
    _set_cell(t2, 1, 1, _f(shore_calc.terminal_observed_volume_m3))
    _set_cell(t2, 1, 2, _f(shore_calc.terminal_standard_volume_m3))
    _set_cell(t2, 1, 3, _f(shore_calc.terminal_weight_air_mt))

    _set_cell(t2, 2, 1, _f(shore_calc.vessel_observed_volume_m3))
    _set_cell(t2, 2, 2, _f(shore_calc.vessel_standard_volume_m3))
    _set_cell(t2, 2, 3, _f(shore_calc.vessel_weight_air_mt))

    _set_cell(t2, 3, 2, _f(shore_calc.meter_quantity_m3))

    _set_cell(t2, 4, 1, _f(shore_calc.difference_observed_volume_m3))
    _set_cell(t2, 4, 2, _f(shore_calc.difference_standard_volume_m3))
    _set_cell(t2, 4, 3, _f(shore_calc.difference_weight_air_mt))

    # ── Table 3: Terminal Rep signatures ──
    t3 = doc.tables[3]
    _set_cell(t3, 1, 0, 'NAME')
    _set_cell(t3, 1, 1, shore_calc.terminal_representative_name or '')

    # ── Table 4: PBPA Inspector signatures ──
    t4 = doc.tables[4]
    _set_cell(t4, 1, 0, 'NAME')
    _set_cell(t4, 1, 1, shore_calc.pbpa_inspector_name or '')

    return _save(doc)


# ── Dip Ticket ────────────────────────────────────────────────────────────────

def generate_dip_ticket_document(inspection):
    """
    Fill DIP TICKET.docx template.

    Template structure:
      Para [0]  THE UNITED REPUBLIC OF TANZANIA
      Para [9]  Terminal representative: ... PBPA Inspector:
      Para [11] Name ... Name
      Para [13] Signature ... Signature

      Table 0 (3r x 2c) – vessel/product/terminal header
        R0: Vessel: | Product:
        R1: Terminal: | Tank No:
        R2: Date: | Time:

      Table 1 (10r x 5c) – measurements
        R0: PARTICULARS | MEASUREMENTS x4
        R1:  | 1st | 2nd | 3rd | Average
        R2:  Overall Dip (mm)
        R3:  Product Dip (mm)
        R4:  Product volume (L)
        R5:  Free water volume (L)
        R6:  Tank temperature
        R7:  Specific gravity (SG)
        R8:  Sample temperature
        R9:  (spare)

      Table 2 (4r x 4c) – seals & meter
        R0: PBPA SEAL POSITION | PBPA SEAL NUMBER | METER READINGS x2
        R1: Outlet valve seal  | | OBS |
        R2: Water valve seal   | | @20 |
        R3: Other branches seal| | MTS |
    """
    doc = _load('dip_ticket')

    # ── Table 0: header ──
    t0 = doc.tables[0]
    _set_cell(t0, 0, 0, f'Vessel: {inspection.vessel_name or ""}')
    _set_cell(t0, 0, 1, f'Product: {inspection.product_name or ""}')
    _set_cell(t0, 1, 0, f'Terminal: {inspection.terminal or ""}')
    _set_cell(t0, 1, 1, f'Tank No: {inspection.tank.tank_id if inspection.tank else ""}')
    date_str = inspection.inspection_date.strftime('%d-%m-%Y') if inspection.inspection_date else ''
    time_str = inspection.inspection_time.strftime('%H:%M') if inspection.inspection_time else ''
    _set_cell(t0, 2, 0, f'Date: {date_str}')
    _set_cell(t0, 2, 1, f'Time: {time_str}')

    # ── Table 1: measurements ──
    t1 = doc.tables[1]

    def _avg(vals):
        nums = [v for v in vals if v is not None]
        return round(sum(nums) / len(nums), 3) if nums else None

    rows = [
        # (row_idx, [v1, v2, v3])
        (2, [inspection.overall_dip_1_mm,       inspection.overall_dip_2_mm,       inspection.overall_dip_3_mm]),
        (3, [inspection.product_dip_1_mm,        inspection.product_dip_2_mm,       inspection.product_dip_3_mm]),
        (4, [inspection.product_volume_1_l,      inspection.product_volume_2_l,     inspection.product_volume_3_l]),
        (5, [inspection.free_water_volume_1_l,   inspection.free_water_volume_2_l,  inspection.free_water_volume_3_l]),
        (6, [inspection.tank_temperature_1_c,    inspection.tank_temperature_2_c,   inspection.tank_temperature_3_c]),
        (7, [inspection.specific_gravity_1,      inspection.specific_gravity_2,     inspection.specific_gravity_3]),
        (8, [inspection.sample_temperature_1_c,  inspection.sample_temperature_2_c, inspection.sample_temperature_3_c]),
    ]
    for row_idx, vals in rows:
        for col, v in enumerate(vals, start=1):
            _set_cell(t1, row_idx, col, _f(v))
        _set_cell(t1, row_idx, 4, _f(_avg(vals)))

    # ── Table 2: seals & meter ──
    t2 = doc.tables[2]
    _set_cell(t2, 1, 1, inspection.outlet_valve_seal_number or '')
    _set_cell(t2, 2, 1, inspection.water_valve_seal_number or '')
    _set_cell(t2, 3, 1, inspection.other_branches_seal_number or '')
    _set_cell(t2, 1, 3, _f(inspection.meter_reading_obs))
    _set_cell(t2, 2, 3, _f(inspection.meter_reading_at_20))
    _set_cell(t2, 3, 3, _f(inspection.meter_reading_mts))

    # ── Signature paragraphs ──
    # Para 11: Name line
    try:
        p11 = doc.paragraphs[11]
        p11.clear()
        r = p11.add_run(f'Name: {inspection.terminal_representative_name or ""}')
        r2_text = f'Name: {inspection.pbpa_inspector_name or ""}'
        # add spacing then second name
        p11.add_run('                                                    ' + r2_text)
    except Exception:
        pass

    return _save(doc)


# ── Product Receipt Certificate ───────────────────────────────────────────────

def generate_product_receipt_document(certificate):
    """
    Fill PRODUCT RECEIPT CERTIFICATE.docx template.

    Template structure:
      Para [6]  Vessel Name:
      Para [7]  Terminal:..
      Para [8]  Date:..
      Para [9]  Time:.
      Para [10] This is to certify...
      Para [12] Quantity received through inlet flowmeters:
      Para [14] Terminal representative: ... PBPA Inspector:
      Para [16] Name ... Name
      Para [18] Signature ... Signature

      Table 0 (13r x 4c) – items
        R0:  Tank No. | Product | Weight in Tonnage | Volume in Liters
        R1-R11: data rows
        R12: TOTAL | TOTAL | total_weight | total_volume

      Table 1 (2r x 2c) – totals summary
        R0: Weight in Tonnage | Volume in Litres
        R1: value | value
    """
    doc = _load('product_receipt')

    # ── Header paragraphs ──
    paras = doc.paragraphs
    _set_para_text(paras, 6, f'Vessel Name: {certificate.vessel_name}')
    _set_para_text(paras, 7, f'Terminal: {certificate.terminal}')
    date_str = certificate.receipt_date.strftime('%d-%m-%Y') if certificate.receipt_date else ''
    _set_para_text(paras, 8, f'Date: {date_str}')
    time_str = certificate.receipt_time.strftime('%H:%M') if certificate.receipt_time else ''
    _set_para_text(paras, 9, f'Time: {time_str}')
    flowmeter_quantity = certificate.quantity_received_through_inlet_flowmeters or 0
    _set_para_text(
        paras, 12,
        ('Quantity received through inlet flowmeters: '
         f'{_f(flowmeter_quantity)} Liters') if flowmeter_quantity > 0
        else 'Quantity received through inlet flowmeters:'
    )

    # ── Table 0: items ──
    t0 = doc.tables[0]
    items = list(certificate.items.all())
    data_rows = 11  # rows 1-11

    for i in range(data_rows):
        if i < len(items):
            item = items[i]
            _set_cell(t0, i + 1, 0, item.tank_no or '')
            _set_cell(t0, i + 1, 1, item.product_name or '')
            _set_cell(t0, i + 1, 2, _f(item.weight_tonnage))
            _set_cell(t0, i + 1, 3, _f(item.volume_liters))
        else:
            for col in range(4):
                _set_cell(t0, i + 1, col, '')

    # Total row (row 12)
    _set_cell(t0, 12, 2, _f(certificate.total_weight_tonnage))
    _set_cell(t0, 12, 3, _f(certificate.total_volume_liters))

    # ── Table 1: summary totals ──
    t1 = doc.tables[1]
    _set_cell(t1, 1, 0, _f(certificate.total_weight_tonnage))
    _set_cell(t1, 1, 1, _f(certificate.total_volume_liters))

    # ── Signature paragraphs ──
    _set_para_text(paras, 16,
        f'Name: {certificate.terminal_representative_name or ""}' +
        '                                                    ' +
        f'Name: {certificate.pbpa_inspector_name or ""}'
    )

    return _save(doc)


# ── Seal and Isolation ────────────────────────────────────────────────────────

def generate_seal_isolation_document(report):
    """
    Fill SEAL AND ISOLATION.docx template.

    Template structure:
      Para [4]  Vessel Name:.
      Para [5]  Product:.
      Para [6]  Terminal:...
      Para [7]  Date:..
      Para [14] Name ... Name
      Para [16] Signature ... Signature

      Table 0 (16r x 2c) – seal entries
        R0:  Location | Seal Number
        R1-R15: data rows
    """
    doc = _load('seal_isolation')
    paras = doc.paragraphs

    _set_para_text(paras, 4, f'Vessel Name: {report.vessel_name}')
    _set_para_text(paras, 5, f'Product: {report.product_name}')
    _set_para_text(paras, 6, f'Terminal: {report.terminal}')
    date_str = report.report_date.strftime('%d-%m-%Y') if report.report_date else ''
    _set_para_text(paras, 7, f'Date: {date_str}')

    # ── Table 0: entries ──
    t0 = doc.tables[0]
    entries = list(report.entries.all())
    data_rows = 15  # rows 1-15

    for i in range(data_rows):
        if i < len(entries):
            _set_cell(t0, i + 1, 0, entries[i].location or '')
            _set_cell(t0, i + 1, 1, entries[i].seal_number or '')
        else:
            _set_cell(t0, i + 1, 0, '')
            _set_cell(t0, i + 1, 1, '')

    # ── Signature paragraphs ──
    _set_para_text(paras, 14,
        f'Name: {report.terminal_representative_name or ""}' +
        '                                                    ' +
        f'Name: {report.pbpa_inspector_name or ""}'
    )

    return _save(doc)


# ── internal helper ───────────────────────────────────────────────────────────

def _set_para_text(paras, idx, text):
    """Replace all text in a paragraph while keeping the first run's formatting."""
    try:
        p = paras[idx]
        # clear all runs
        for run in p.runs:
            run.text = ''
        if p.runs:
            p.runs[0].text = text
        else:
            p.add_run(text)
    except IndexError:
        pass


# ── Compatibility shim (keeps existing views.py calls working) ────────────────

class ShoreTankDocumentGenerator:
    def fill_shore_tank_calculation_document(self, shore_calc, template_path=None):
        return generate_shore_tank_document(shore_calc)

    def fill_product_receipt_certificate_document(self, certificate, template_path=None):
        return generate_product_receipt_document(certificate)

    def fill_seal_isolation_document(self, report, template_path=None):
        return generate_seal_isolation_document(report)

    def fill_dip_ticket_document(self, inspection, template_path=None):
        return generate_dip_ticket_document(inspection)


# ── PDF Generation Functions ──────────────────────────────────────────────

from reportlab.lib.pagesizes import A4, landscape
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT


def _draw_signature_in_field(canvas, image_data, field_x, field_y, field_w, field_h,
                             padding=1.5 * mm):
    """Draw a signature centred and proportionally scaled inside one field."""
    if not image_data:
        return
    try:
        import base64
        from reportlab.lib.utils import ImageReader
        payload = image_data.split(',', 1)[1] if ',' in image_data else image_data
        image = ImageReader(BytesIO(base64.b64decode(payload)))
        image_w, image_h = image.getSize()
        available_w = max(field_w - 2 * padding, 0)
        available_h = max(field_h - 2 * padding, 0)
        if not image_w or not image_h or not available_w or not available_h:
            return
        scale = min(available_w / image_w, available_h / image_h)
        draw_w, draw_h = image_w * scale, image_h * scale
        canvas.drawImage(
            image,
            field_x + (field_w - draw_w) / 2,
            field_y + (field_h - draw_h) / 2,
            draw_w, draw_h,
            mask='auto',
        )
    except Exception:
        # An invalid optional image must not prevent an otherwise valid report.
        pass


def _stored_signature_image(document_type, document_id, role):
    """Return the latest valid drawn signature for the person/role requested."""
    try:
        from .models import DocumentSignature
        signature = DocumentSignature.objects.filter(
            doc_type=document_type, doc_id=document_id, role=role, status='valid'
        ).exclude(signature_image='').order_by('-created_at').first()
        return signature.signature_image if signature else ''
    except Exception:
        return ''
import io


def _generate_dip_ticket_pdf_legacy(inspection):
    """Generate Dip Ticket as PDF using ReportLab."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=10*mm, leftMargin=10*mm, topMargin=10*mm, bottomMargin=10*mm)
    
    elements = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#1a1a1a'),
        spaceAfter=6,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    elements.append(Paragraph("THE UNITED REPUBLIC OF TANZANIA", title_style))
    elements.append(Paragraph("PETROLEUM BULK PROCUREMENT AGENCY - DIP TICKET", title_style))
    elements.append(Spacer(1, 8*mm))
    
    # Header info table - use Paragraph for HTML support
    header_style = ParagraphStyle(
        'HeaderCell',
        parent=styles['Normal'],
        fontSize=9,
        alignment=TA_LEFT
    )
    
    header_data = [
        [
            Paragraph(f"<b>Vessel:</b> {inspection.vessel_name or ''}", header_style),
            Paragraph(f"<b>Product:</b> {inspection.product_name or ''}", header_style)
        ],
        [
            Paragraph(f"<b>Terminal:</b> {inspection.terminal or ''}", header_style),
            Paragraph(f"<b>Tank No:</b> {inspection.tank.tank_id if inspection.tank else ''}", header_style)
        ],
        [
            Paragraph(f"<b>Date:</b> {inspection.inspection_date.strftime('%d-%m-%Y') if inspection.inspection_date else ''}", header_style),
            Paragraph(f"<b>Time:</b> {inspection.inspection_time.strftime('%H:%M') if inspection.inspection_time else ''}", header_style)
        ]
    ]
    
    header_table = Table(header_data, colWidths=[doc.width/2-5*mm, doc.width/2-5*mm])
    header_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f5f5f5')),
    ]))
    elements.append(header_table)
    elements.append(Spacer(1, 8*mm))
    
    # Measurements table
    def _avg(vals):
        nums = [v for v in vals if v is not None]
        return round(sum(nums) / len(nums), 3) if nums else None
    
    normal_style = ParagraphStyle('Normal2', parent=styles['Normal'], fontSize=9, alignment=TA_CENTER)
    
    measurements_data = [
        ['PARTICULARS', '1st', '2nd', '3rd', 'Average'],
        ['Overall Dip (mm)', _f(inspection.overall_dip_1_mm), _f(inspection.overall_dip_2_mm), _f(inspection.overall_dip_3_mm), _f(_avg([inspection.overall_dip_1_mm, inspection.overall_dip_2_mm, inspection.overall_dip_3_mm]))],
        ['Product Dip (mm)', _f(inspection.product_dip_1_mm), _f(inspection.product_dip_2_mm), _f(inspection.product_dip_3_mm), _f(_avg([inspection.product_dip_1_mm, inspection.product_dip_2_mm, inspection.product_dip_3_mm]))],
        ['Product Volume (L)', _f(inspection.product_volume_1_l), _f(inspection.product_volume_2_l), _f(inspection.product_volume_3_l), _f(_avg([inspection.product_volume_1_l, inspection.product_volume_2_l, inspection.product_volume_3_l]))],
        ['Free Water Volume (L)', _f(inspection.free_water_volume_1_l), _f(inspection.free_water_volume_2_l), _f(inspection.free_water_volume_3_l), _f(_avg([inspection.free_water_volume_1_l, inspection.free_water_volume_2_l, inspection.free_water_volume_3_l]))],
        ['Tank Temperature (°C)', _f(inspection.tank_temperature_1_c, 1), _f(inspection.tank_temperature_2_c, 1), _f(inspection.tank_temperature_3_c, 1), _f(_avg([inspection.tank_temperature_1_c, inspection.tank_temperature_2_c, inspection.tank_temperature_3_c]), 1)],
        ['Specific Gravity (SG)', _f(inspection.specific_gravity_1), _f(inspection.specific_gravity_2), _f(inspection.specific_gravity_3), _f(_avg([inspection.specific_gravity_1, inspection.specific_gravity_2, inspection.specific_gravity_3]))],
        ['Sample Temperature (°C)', _f(inspection.sample_temperature_1_c, 1), _f(inspection.sample_temperature_2_c, 1), _f(inspection.sample_temperature_3_c, 1), _f(_avg([inspection.sample_temperature_1_c, inspection.sample_temperature_2_c, inspection.sample_temperature_3_c]), 1)],
    ]
    
    measurements_table = Table(measurements_data, colWidths=[doc.width*0.35, doc.width*0.13, doc.width*0.13, doc.width*0.13, doc.width*0.13])
    measurements_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#8B1A1A')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGNMENT', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f9f9f9')]),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
    ]))
    elements.append(measurements_table)
    elements.append(Spacer(1, 8*mm))
    
    # Seals & meter readings table
    seals_data = [
        ['PBPA SEAL POSITION', 'SEAL NUMBER', 'METER READINGS', ''],
        ['', '', 'OBS', '@20'],
        ['Outlet valve seal', inspection.outlet_valve_seal_number or '', _f(inspection.meter_reading_obs), ''],
        ['Water valve seal', inspection.water_valve_seal_number or '', _f(inspection.meter_reading_at_20), ''],
        ['Other branches seal', inspection.other_branches_seal_number or '', _f(inspection.meter_reading_mts), ''],
    ]
    
    seals_table = Table(seals_data, colWidths=[doc.width*0.35, doc.width*0.25, doc.width*0.2, doc.width*0.2])
    seals_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#8B1A1A')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGNMENT', (1, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ROWBACKGROUNDS', (0, 2), (-1, -1), [colors.white, colors.HexColor('#f9f9f9')]),
    ]))
    elements.append(seals_table)
    elements.append(Spacer(1, 12*mm))
    
    # Signatures
    sig_data = [
        ['Terminal Representative', '', 'PBPA Inspector'],
        ['Name: ' + (inspection.terminal_representative_name or ''), '', 'Name: ' + (inspection.pbpa_inspector_name or '')],
        ['Signature: _____________________', '', 'Signature: _____________________'],
    ]
    
    sig_table = Table(sig_data, colWidths=[doc.width/3-5*mm, doc.width/3-5*mm, doc.width/3-5*mm])
    sig_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGNMENT', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
    ]))
    elements.append(sig_table)
    
    doc.build(elements)
    buffer.seek(0)
    return buffer


def generate_dip_ticket_pdf(inspection):
    """Generate a PBPA dip ticket PDF in the official fixed-sheet format."""
    buffer = io.BytesIO()
    pdf = rl_canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    margin = 13 * mm

    def draw_text(x, y, value, size=8, bold=False, align='left'):
        pdf.setFont('Helvetica-Bold' if bold else 'Helvetica', size)
        value = '' if value is None else str(value)
        if align == 'center':
            pdf.drawCentredString(x, y, value)
        elif align == 'right':
            pdf.drawRightString(x, y, value)
        else:
            pdf.drawString(x, y, value)

    def draw_line(x1, y1, x2, y2, line_width=0.5):
        pdf.setLineWidth(line_width)
        pdf.line(x1, y1, x2, y2)
        pdf.setLineWidth(0.5)

    def draw_rect(x, y, w, h, fill=0):
        pdf.rect(x, y, w, h, stroke=1, fill=fill)

    def average(values):
        numeric = [v for v in values if v is not None]
        return round(sum(numeric) / len(numeric), 3) if numeric else None

    def date_value():
        return inspection.inspection_date.strftime('%d-%m-%Y') if inspection.inspection_date else ''

    def time_value():
        return inspection.inspection_time.strftime('%H:%M') if inspection.inspection_time else ''

    def signature_value_line(x, y, w, value):
        draw_line(x, y - 1.2, x + w, y - 1.2)
        draw_text(x + 2, y, value or '', 8, bold=True)

    pdf.setLineWidth(0.6)

    # Header mirrors the official PBPA ticket sheet.
    draw_text(width / 2, height - 18 * mm, 'THE UNITED REPUBLIC OF TANZANIA', 7, align='center')
    draw_text(width / 2, height - 28 * mm, 'PETROLEUM BULK PROCUREMENT AGENCY', 14, bold=True, align='center')
    draw_text(width / 2, height - 33 * mm, 'TANZANIA PORTS AUTHORITY, ONE STOP CENTRE BUILDING 11TH FLOOR, SOKOINE DRIVE, PLOT NO. 1/2', 5.5, align='center')
    draw_text(width / 2, height - 37 * mm, 'P.O. BOX 2314 DAR ES SALAAM, TANZANIA', 5.5, align='center')
    # Official Tanzania and PBPA artwork is applied to every downloaded PDF
    # by ``add_pbpa_letterhead``.  Do not draw the old text/circle stand-ins.
    draw_line(margin, height - 42 * mm, width - margin, height - 42 * mm, 1.1)

    title_x = margin + 55 * mm
    title_y = height - 58 * mm
    title_w = 54 * mm
    title_h = 13 * mm
    draw_rect(title_x, title_y, title_w, title_h)
    pdf.setFillColor(colors.black)
    pdf.roundRect(title_x + 2 * mm, title_y + 2 * mm, title_w - 4 * mm, title_h - 4 * mm, 2 * mm, fill=1, stroke=0)
    pdf.setFillColor(colors.white)
    draw_text(title_x + title_w / 2, title_y + 4.2 * mm, 'DIP TICKET', 15, bold=True, align='center')
    pdf.setFillColor(colors.black)
    serial = str(inspection.ticket_number or inspection.id or '').replace('DIP-', '')
    draw_text(width - margin - 7 * mm, title_y + 4.2 * mm, serial, 12, align='right')

    # Ticket header details.
    grid_x = margin
    grid_y = height - 88 * mm
    grid_w = 152 * mm
    grid_h = 24 * mm
    row_h = grid_h / 3
    col_w = grid_w / 2
    draw_rect(grid_x, grid_y, grid_w, grid_h)
    draw_line(grid_x + col_w, grid_y, grid_x + col_w, grid_y + grid_h)
    for i in range(1, 3):
        draw_line(grid_x, grid_y + i * row_h, grid_x + grid_w, grid_y + i * row_h)
    details = [
        ('Vessel:', inspection.vessel_name or '', 'Product:', inspection.product_name or ''),
        ('Terminal:', inspection.terminal or '', 'Tank No:', inspection.tank_no or (inspection.tank.tank_id if inspection.tank else '')),
        ('Date:', date_value(), 'Time:', time_value()),
    ]
    for idx, (left_label, left_value, right_label, right_value) in enumerate(details):
        y = grid_y + grid_h - (idx + 0.68) * row_h
        draw_text(grid_x + 2 * mm, y, left_label, 7, bold=True)
        draw_text(grid_x + 20 * mm, y, left_value, 9, bold=True)
        if right_label:
            draw_text(grid_x + col_w + 2 * mm, y, right_label, 7, bold=True)
            draw_text(grid_x + col_w + 22 * mm, y, right_value, 9, bold=True)

    # Measurements table.
    table_x = margin
    table_y = height - 170 * mm
    table_w = width - 2 * margin
    table_h = 73 * mm
    header_h = 13 * mm
    subheader_h = 10 * mm
    data_h = (table_h - header_h - subheader_h) / 7
    particulars_w = 52 * mm
    measurement_w = (table_w - particulars_w) / 4
    draw_rect(table_x, table_y, table_w, table_h)
    draw_line(table_x + particulars_w, table_y, table_x + particulars_w, table_y + table_h)
    draw_line(table_x, table_y + table_h - header_h, table_x + table_w, table_y + table_h - header_h)
    draw_line(table_x + particulars_w, table_y + table_h - header_h - subheader_h, table_x + table_w, table_y + table_h - header_h - subheader_h)
    for i in range(1, 4):
        draw_line(table_x + particulars_w + i * measurement_w, table_y, table_x + particulars_w + i * measurement_w, table_y + table_h - header_h)
    for i in range(1, 8):
        draw_line(table_x, table_y + i * data_h, table_x + table_w, table_y + i * data_h)

    draw_text(table_x + particulars_w / 2, table_y + table_h - 8 * mm, 'PARTICULARS', 7, bold=True, align='center')
    draw_text(table_x + particulars_w + (table_w - particulars_w) / 2, table_y + table_h - 6 * mm, 'MEASUREMENTS', 8, bold=True, align='center')
    for i, (line_one, line_two) in enumerate((('1st', 'Measurement'), ('2nd', 'Measurement'), ('3rd', 'Measurement'), ('Average', 'Measurement'))):
        cx = table_x + particulars_w + i * measurement_w + measurement_w / 2
        draw_text(cx, table_y + table_h - header_h - 4 * mm, line_one, 6, align='center')
        draw_text(cx, table_y + table_h - header_h - 7.2 * mm, line_two, 5.2, align='center')

    rows = [
        ('Overall Dip (mm)', [inspection.overall_dip_1_mm, inspection.overall_dip_2_mm, inspection.overall_dip_3_mm], 0),
        ('Product Dip (mm)', [inspection.product_dip_1_mm, inspection.product_dip_2_mm, inspection.product_dip_3_mm], 0),
        ('Product Volume (L)', [inspection.product_volume_1_l, inspection.product_volume_2_l, inspection.product_volume_3_l], 0),
        ('Free Water Dip (L)', [inspection.free_water_volume_1_l, inspection.free_water_volume_2_l, inspection.free_water_volume_3_l], 0),
        ('Tank Temperature (C)', [inspection.tank_temperature_1_c, inspection.tank_temperature_2_c, inspection.tank_temperature_3_c], 1),
        ('Specific Gravity (SG)', [inspection.specific_gravity_1, inspection.specific_gravity_2, inspection.specific_gravity_3], 3),
        ('Sample Temperature (C)', [inspection.sample_temperature_1_c, inspection.sample_temperature_2_c, inspection.sample_temperature_3_c], 1),
    ]
    for idx, (label, values, precision) in enumerate(rows):
        y = table_y + table_h - header_h - subheader_h - (idx + 0.68) * data_h
        draw_text(table_x + 2 * mm, y, label, 7)
        for col, value in enumerate(values + [average(values)]):
            draw_text(table_x + particulars_w + col * measurement_w + measurement_w / 2, y, _f(value, precision), 9, bold=True, align='center')

    # PBPA seal block.
    seal_x = margin
    seal_y = height - 205 * mm
    seal_w = width - 2 * margin
    seal_h = 28 * mm
    seal_label_w = 58 * mm
    seal_number_w = 48 * mm
    meter_w = (seal_w - seal_label_w - seal_number_w) / 3
    seal_row_h = seal_h / 5
    meter_x = seal_x + seal_label_w + seal_number_w
    draw_rect(seal_x, seal_y, seal_w, seal_h)
    draw_line(seal_x + seal_label_w, seal_y, seal_x + seal_label_w, seal_y + seal_h)
    draw_line(meter_x, seal_y, meter_x, seal_y + seal_h)
    for i in range(1, 3):
        x = meter_x + i * meter_w
        # OBS, @20 and MTS are subcolumns only; keep the METER READINGS
        # heading merged across them in the row above.
        draw_line(x, seal_y, x, seal_y + seal_h - seal_row_h)
    # First two columns span the two header rows.  The first horizontal rule
    # therefore begins at METER READINGS, while the remaining rules cross the
    # entire table.
    draw_line(meter_x, seal_y + seal_h - seal_row_h, seal_x + seal_w, seal_y + seal_h - seal_row_h)
    for i in range(2, 5):
        y = seal_y + seal_h - i * seal_row_h
        draw_line(seal_x, y, seal_x + seal_w, y)

    header_y = seal_y + seal_h - seal_row_h
    draw_text(seal_x + 2 * mm, header_y, 'PBPA SEAL POSITION', 6, bold=True)
    draw_text(seal_x + seal_label_w + seal_number_w / 2, header_y, 'PBPA SEAL NUMBER', 6, bold=True, align='center')
    draw_text(meter_x + 1.5 * meter_w, seal_y + seal_h - 0.7 * seal_row_h,
              'METER READINGS', 6, bold=True, align='center')
    for idx, heading in enumerate(('OBS', '@20', 'MTS')):
        x = meter_x + idx * meter_w + meter_w / 2
        draw_text(x, seal_y + seal_h - 1.7 * seal_row_h, heading, 6, bold=True, align='center')
    for idx, (label, seal_value) in enumerate((
        ('Outlet valve seal', inspection.outlet_valve_seal_number),
        ('Water valve seal', inspection.water_valve_seal_number),
        ('Other branches seal', inspection.other_branches_seal_number),
    )):
        y = seal_y + seal_h - (idx + 2.7) * seal_row_h
        draw_text(seal_x + 2 * mm, y, label, 7)
        draw_text(seal_x + seal_label_w + seal_number_w / 2, y, seal_value or '-', 9, bold=True, align='center')
        # Meter readings describe one reading set, so keep OBS, @20 and MTS
        # together on a single row rather than placing them diagonally beside
        # the three independent seal-position records.
        meter_values = (
            (_f(inspection.meter_reading_obs),
             _f(inspection.meter_reading_at_20),
             _f(inspection.meter_reading_mts))
            if idx == 0 else ('', '', '')
        )
        for col, meter_value in enumerate(meter_values):
            x = meter_x + col * meter_w + meter_w / 2
            draw_text(x, y, meter_value, 8, bold=True, align='center')

    # Signature area.
    sig_y = height - 255 * mm
    left_x = margin + 4 * mm
    right_x = width / 2 + 15 * mm
    draw_text(left_x, sig_y + 25 * mm, 'Terminal Representative:', 7, bold=True)
    draw_text(right_x, sig_y + 25 * mm, 'Petroleum Bulk Procurement Agency Inspector:', 7, bold=True)
    signature_value_line(left_x + 25 * mm, sig_y + 13 * mm, 45 * mm, inspection.terminal_representative_name)
    signature_value_line(right_x + 22 * mm, sig_y + 13 * mm, 48 * mm, inspection.pbpa_inspector_name)
    draw_text(left_x + 42 * mm, sig_y + 8 * mm, 'Name', 6, align='center')
    draw_text(right_x + 45 * mm, sig_y + 8 * mm, 'Name', 6, align='center')
    signature_value_line(left_x + 25 * mm, sig_y - 2 * mm, 45 * mm, inspection.terminal_representative_signature)
    signature_value_line(right_x + 22 * mm, sig_y - 2 * mm, 48 * mm, inspection.pbpa_inspector_signature)
    draw_text(left_x + 42 * mm, sig_y - 7 * mm, 'Signature', 6, align='center')
    draw_text(right_x + 45 * mm, sig_y - 7 * mm, 'Signature', 6, align='center')

    # These boxes are derived from the two printed signature rules above.
    # They preserve the official order: Terminal Representative, then PBPA
    # Inspector, and avoid later page-coordinate stamping.
    terminal_signature_field = (left_x + 25 * mm, sig_y - 2 * mm, 45 * mm, 9 * mm)
    inspector_signature_field = (right_x + 22 * mm, sig_y - 2 * mm, 48 * mm, 9 * mm)
    _draw_signature_in_field(
        pdf, _stored_signature_image('dip_ticket', inspection.pk, 'terminal_representative'),
        *terminal_signature_field,
    )
    _draw_signature_in_field(
        pdf, _stored_signature_image('dip_ticket', inspection.pk, 'inspector'),
        *inspector_signature_field,
    )
    draw_text(margin, 12 * mm, 'QF-17', 5)

    pdf.showPage()
    pdf.save()
    buffer.seek(0)
    return buffer


def generate_seal_isolation_pdf(report):
    """Generate the standalone PBPA Sealing and Isolation Report PDF."""
    buffer = io.BytesIO()
    pdf = rl_canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    margin = 14 * mm
    content_w = width - 2 * margin

    def draw_line(x1, y1, x2, y2, lw=0.55):
        pdf.setLineWidth(lw)
        pdf.line(x1, y1, x2, y2)

    def draw_rect(x, y, w, h, lw=0.55):
        pdf.setLineWidth(lw)
        pdf.rect(x, y, w, h, stroke=1, fill=0)

    def draw_text(x, y, text, size=7, bold=False, align='left'):
        value = '' if text is None else str(text)
        pdf.setFont('Helvetica-Bold' if bold else 'Helvetica', size)
        if align == 'center':
            pdf.drawCentredString(x, y, value)
        elif align == 'right':
            pdf.drawRightString(x, y, value)
        else:
            pdf.drawString(x, y, value)

    def draw_wrapped(x, y, text, max_width, size=7, bold=False, line_gap=3.4 * mm):
        words = ('' if text is None else str(text)).split()
        lines = []
        line = ''
        pdf.setFont('Helvetica-Bold' if bold else 'Helvetica', size)
        for word in words:
            candidate = f'{line} {word}'.strip()
            if pdf.stringWidth(candidate, 'Helvetica-Bold' if bold else 'Helvetica', size) <= max_width:
                line = candidate
            else:
                if line:
                    lines.append(line)
                line = word
        if line:
            lines.append(line)
        if not lines:
            lines = ['']
        for idx, line_text in enumerate(lines[:2]):
            draw_text(x, y - idx * line_gap, line_text, size=size, bold=bold)

    def line_field(label, value, x, y, w, size=8):
        label_text = f'{label}:'
        draw_text(x, y, label_text, size=size, bold=True)
        label_w = pdf.stringWidth(label_text, 'Helvetica-Bold', size) + 2 * mm
        draw_text(x + label_w, y, value or '', size=size)
        draw_line(x + label_w, y - 1.1 * mm, x + w, y - 1.1 * mm, 0.35)

    y = height - 12 * mm

    # The shared PDF letterhead applies the official Tanzania and PBPA marks.
    pdf.setStrokeColor(colors.black)
    draw_text(width / 2, y, 'THE UNITED REPUBLIC OF TANZANIA', size=6, align='center')
    y -= 4.8 * mm
    draw_text(width / 2, y, 'PETROLEUM BULK PROCUREMENT AGENCY', size=11.5, bold=True, align='center')
    y -= 3.6 * mm
    draw_text(width / 2, y, 'TANZANIA PORTS AUTHORITY, ONE STOP CENTER BUILDING 11TH FLOOR, SOKOINE DRIVE, PLOT NO:1/2', size=4.4, align='center')
    y -= 3 * mm
    draw_text(width / 2, y, 'Tel: +255222129009 / Fax: +255222129093 / info@pbpa.go.tz / WEBSITE: www.pbpa.go.tz / P.O. BOX 2634 Dar es Salaam, TANZANIA', size=4.2, align='center')
    y -= 2.4 * mm
    draw_line(margin, y, width - margin, y, 0.7)

    # Header fields and serial number.
    y -= 6 * mm
    left_w = 116 * mm
    date_str = report.report_date.strftime('%d-%m-%Y') if report.report_date else ''
    line_field('Vessel name', report.vessel_name, margin, y, left_w)
    draw_text(width - margin - 4 * mm, y - 0.5 * mm, report.report_number or '', size=10.5, align='right')
    y -= 6 * mm
    line_field('Product', report.product_name, margin, y, left_w)
    y -= 6 * mm
    line_field('Terminal', report.terminal, margin, y, left_w)
    y -= 6 * mm
    line_field('Date', date_str, margin, y, left_w)

    # Title.
    y -= 9 * mm
    title_w = 72 * mm
    title_h = 7 * mm
    title_x = (width - title_w) / 2
    draw_rect(title_x, y - title_h + 1.5 * mm, title_w, title_h, 1)
    draw_text(width / 2, y - 3 * mm, 'SEALING AND ISOLATION REPORT', size=9, bold=True, align='center')

    # Entries table.
    y -= 11 * mm
    table_x = margin
    table_w = content_w
    location_w = table_w * 0.56
    seal_w = table_w - location_w
    header_h = 7 * mm
    row_h = 7 * mm
    rows = 24
    table_h = header_h + rows * row_h

    draw_rect(table_x, y - table_h, table_w, table_h, 0.7)
    draw_line(table_x + location_w, y, table_x + location_w, y - table_h, 0.55)
    draw_line(table_x, y - header_h, table_x + table_w, y - header_h, 0.55)
    draw_text(table_x + location_w / 2, y - 4.5 * mm, 'Location', size=7, bold=True, align='center')
    draw_text(table_x + location_w + seal_w / 2, y - 4.5 * mm, 'Seal Number', size=7, bold=True, align='center')

    for idx in range(rows):
        row_y = y - header_h - idx * row_h
        draw_line(table_x, row_y - row_h, table_x + table_w, row_y - row_h, 0.35)

    entries = list(report.entries.all())
    for idx, entry in enumerate(entries[:rows]):
        row_top = y - header_h - idx * row_h
        text_y = row_top - 4.4 * mm
        draw_wrapped(table_x + 3 * mm, text_y, entry.location or '', location_w - 6 * mm, size=7.5, bold=True)
        draw_text(table_x + location_w + 3 * mm, text_y, entry.seal_number or '', size=8, bold=True)

    # Signatures.
    y = y - table_h - 8 * mm
    block_w = (content_w - 16 * mm) / 2
    left_x = margin
    right_x = margin + block_w + 16 * mm
    draw_text(left_x, y, 'TPA/TIPER/Terminal representative', size=6.5, bold=True)
    draw_text(right_x, y, 'PBPA Inspector', size=6.5, bold=True)
    y -= 8 * mm
    line_field('Name', report.terminal_representative_name, left_x, y, block_w, size=6.5)
    line_field('Name', report.pbpa_inspector_name, right_x, y, block_w, size=6.5)
    y -= 10 * mm
    line_field('Signature', report.terminal_representative_signature, left_x, y, block_w, size=6.5)
    line_field('Signature', report.pbpa_inspector_signature, right_x, y, block_w, size=6.5)

    # A compact, user-visible integrity section makes the document's security
    # state clear in printouts as well as in the API verification endpoint.
    document_hash = (getattr(report, 'document_hash', '') or '').strip()
    if document_hash:
        footer_y = 7 * mm
        pdf.setStrokeColor(colors.HexColor('#9CA3AF'))
        draw_line(margin, footer_y + 15 * mm, width - margin, footer_y + 15 * mm, 0.35)
        draw_text(margin, footer_y + 11 * mm, 'DOCUMENT INTEGRITY', size=5.8, bold=True)
        integrity = 'VERIFIED SIGNATURE RECORD' if getattr(report, 'is_signed', False) else 'PENDING SIGNATURE'
        draw_text(margin, footer_y + 7 * mm, f'Integrity: {integrity}', size=5.5, bold=True)
        draw_text(margin, footer_y + 3.5 * mm, f'SHA-256: {document_hash[:40]}...', size=5.1)
        # The downloadable original receives one authoritative, URL-based QR
        # code in its bottom-right corner (see finalize_pdf_bytes).

    draw_text(width - margin, 8 * mm, 'SIR/7', size=5, align='right')

    pdf.showPage()
    pdf.save()
    buffer.seek(0)
    return buffer


def generate_shore_tank_pdf(shore_calc):
    """
    Generate Shore Tank Calculation PDF matching the official PBPA format:
    - Portrait A4
    - PBPA letterhead
    - Vessel/Product/Terminal header box
    - SHORE TANK CALCULATIONS title with document number
    - Full measurements table: PARTICULARS | INITIAL | L.DISP/PROV/FINAL (per tank)
    - Summary table: TERMINAL / VESSEL / METER QUANTITY / DIFFERENCE
    - PBPA Inspector & Terminal Representative signature blocks
    """
    buf = io.BytesIO()
    W, H = A4
    M = 12 * mm
    TW = W - 2 * M
    c = rl_canvas.Canvas(buf, pagesize=A4)

    def line(x1, y1, x2, y2, lw=0.5):
        c.setLineWidth(lw)
        c.line(x1, y1, x2, y2)

    def rect(x, y, w, h, fill=0, lw=0.5):
        c.setLineWidth(lw)
        c.rect(x, y, w, h, fill=fill)

    def cell(x, y, w, h, text, font='Helvetica', size=7, align='C', bold=False):
        fn = 'Helvetica-Bold' if bold else font
        c.setFont(fn, size)
        tx = x + w / 2 if align == 'C' else (x + w - 1*mm if align == 'R' else x + 1*mm)
        ty = y + h / 2 - size * 0.35
        c.drawString(tx, ty, str(text)) if align == 'L' else (
            c.drawRightString(tx, ty, str(text)) if align == 'R' else
            c.drawCentredString(tx, ty, str(text))
        )

    items = list(shore_calc.tank_items.all())
    n = min(len(items), 2)

    # Compute terminal totals from tank items (no stored summary fields on model)
    def _sum_attr(attr):
        vals = [getattr(it, attr, None) for it in items]
        nums = [float(v) for v in vals if v is not None]
        return round(sum(nums), 3) if nums else None

    terminal_obs = _sum_attr('gross_observed_final_m3')
    terminal_std = _sum_attr('standard_volume_final_m3')
    terminal_wt  = _sum_attr('weight_air_final_mt')
    vessel_obs   = shore_calc.vessel_observed_volume_m3
    vessel_std   = shore_calc.vessel_standard_volume_m3
    vessel_wt    = shore_calc.vessel_weight_air_mt
    meter_qty    = shore_calc.meter_quantity_m3

    def _diff(a, b):
        if a is not None and b is not None:
            return round(float(a) - float(b), 3)
        return None

    diff_obs = _diff(terminal_obs, vessel_obs)
    diff_std = _diff(terminal_std, vessel_std)
    diff_wt  = _diff(terminal_wt, vessel_wt)

    y = H - 8 * mm

    # ── Letterhead ────────────────────────────────────────────────────────────
    c.setFont('Helvetica', 7.5)
    c.drawCentredString(W / 2, y, 'THE UNITED REPUBLIC OF TANZANIA')
    y -= 5 * mm
    c.setFont('Helvetica-Bold', 10)
    c.drawCentredString(W / 2, y, 'PETROLEUM BULK PROCUREMENT AGENCY')
    y -= 4 * mm
    c.setFont('Helvetica', 6)
    c.drawCentredString(W / 2, y, 'TANZANIA PORTS AUTHORITY, ONE STOP CENTER BUILDING 11TH FLOOR, SOKOINE DRIVE, PLOT NO:1/2')
    y -= 3.5 * mm
    c.drawCentredString(W / 2, y,
        'Tel: +255222129009 / Fax: +255222129093 / info@pbpa.go.tz / WEBSITE: www.pbpa.go.tz / P.O. BOX 2634 Dar es Salaam, TANZANIA')
    y -= 4 * mm
    line(M, y, W - M, y)
    y -= 3 * mm

    # ── Vessel / Product / Terminal header box ────────────────────────────────
    rh = 6 * mm
    lw1 = TW * 0.28
    lw2 = TW * 0.32
    lw3 = TW * 0.40
    date_str = shore_calc.calculation_date.strftime('%d-%m-%Y') if shore_calc.calculation_date else ''

    rows_hdr = [
        ('VESSEL NAME', shore_calc.vessel_name or '', 'Date', date_str),
        ('PRODUCT',     shore_calc.product_name or '', 'Vessel Density', _f4(shore_calc.vessel_density_kg_m3) if shore_calc.vessel_density_kg_m3 else ''),
        ('TERMINAL',    shore_calc.terminal or '',    'Vessel Temperature', _f(shore_calc.vessel_temperature_c, 1) if shore_calc.vessel_temperature_c else ''),
    ]
    for label, val, label2, val2 in rows_hdr:
        rect(M, y - rh, lw1, rh)
        rect(M + lw1, y - rh, lw2, rh)
        rect(M + lw1 + lw2, y - rh, lw3 * 0.45, rh)
        rect(M + lw1 + lw2 + lw3 * 0.45, y - rh, lw3 * 0.55, rh)
        cell(M, y - rh, lw1, rh, label, bold=True, size=7)
        cell(M + lw1, y - rh, lw2, rh, val, size=7, align='L')
        cell(M + lw1 + lw2, y - rh, lw3 * 0.45, rh, label2, bold=True, size=7)
        cell(M + lw1 + lw2 + lw3 * 0.45, y - rh, lw3 * 0.55, rh, val2, size=7, align='L')
        y -= rh
    y -= 2 * mm

    # ── Title + Document Number ───────────────────────────────────────────────
    title_h = 8 * mm
    rect(M, y - title_h, TW * 0.65, title_h, lw=1.5)
    c.setFont('Helvetica-Bold', 11)
    c.drawCentredString(M + TW * 0.65 / 2, y - title_h + 2.5 * mm, 'SHORE TANK CALCULATIONS')
    # Doc number box top-right
    rect(M + TW * 0.65 + 2 * mm, y - title_h, TW * 0.35 - 2 * mm, title_h, lw=1.5)
    c.setFont('Helvetica-Bold', 9)
    c.drawCentredString(M + TW * 0.65 + 2 * mm + (TW * 0.35 - 2 * mm) / 2,
                        y - title_h + 2.5 * mm,
                        shore_calc.calculation_number or '')
    y -= title_h + 2 * mm

    # ── Measurements table ────────────────────────────────────────────────────
    # Columns: PARTICULARS | Tank1 INITIAL | Tank1 L.DISP/PROV/FINAL | Tank2 INITIAL | Tank2 L.DISP/PROV/FINAL
    # (if only 1 tank, last 2 cols are blank)
    col0 = TW * 0.30
    col_pair = (TW - col0) / 2  # width per tank pair
    col_val = col_pair / 2

    # Tank number header row
    rh2 = 5.5 * mm
    rect(M, y - rh2, col0, rh2)
    cell(M, y - rh2, col0, rh2, 'Tank No.', bold=True, size=7)
    for i in range(2):
        x0 = M + col0 + i * col_pair
        rect(x0, y - rh2, col_pair, rh2)
        tank_no = items[i].tank_no if i < n else ''
        cell(x0, y - rh2, col_pair, rh2, tank_no, bold=True, size=8)
    y -= rh2

    # Sub-header: INITIAL / selected operation (per tank)
    rect(M, y - rh2, col0, rh2)
    cell(M, y - rh2, col0, rh2, 'PARTICULARS', bold=True, size=7)
    for i in range(2):
        x0 = M + col0 + i * col_pair
        rect(x0, y - rh2, col_val, rh2)
        rect(x0 + col_val, y - rh2, col_val, rh2)
        cell(x0, y - rh2, col_val, rh2, 'INITIAL', bold=True, size=6.5)
        operation_label = items[i].operation_label if i < n else ''
        cell(x0 + col_val, y - rh2, col_val, rh2, operation_label, bold=True, size=5.5)
    y -= rh2

    def mrow(label, attr_i, attr_f, dp=3):
        nonlocal y
        rect(M, y - rh2, col0, rh2)
        cell(M, y - rh2, col0, rh2, label, size=6.5, align='L')
        for i in range(2):
            x0 = M + col0 + i * col_pair
            rect(x0, y - rh2, col_val, rh2)
            rect(x0 + col_val, y - rh2, col_val, rh2)
            vi = _f(getattr(items[i], attr_i, None), dp) if i < n else ''
            vf = _f(getattr(items[i], attr_f, None), dp) if i < n else ''
            cell(x0, y - rh2, col_val, rh2, vi, size=6.5)
            cell(x0 + col_val, y - rh2, col_val, rh2, vf, size=6.5)
        y -= rh2

    def mrow4(label, attr_i, attr_f, dp=4):
        """4 decimal places for ASTM values."""
        nonlocal y
        rect(M, y - rh2, col0, rh2)
        cell(M, y - rh2, col0, rh2, label, size=6.5, align='L')
        for i in range(2):
            x0 = M + col0 + i * col_pair
            rect(x0, y - rh2, col_val, rh2)
            rect(x0 + col_val, y - rh2, col_val, rh2)
            vi = _f(getattr(items[i], attr_i, None), dp) if i < n else ''
            vf = _f(getattr(items[i], attr_f, None), dp) if i < n else ''
            cell(x0, y - rh2, col_val, rh2, vi, size=6.5)
            cell(x0 + col_val, y - rh2, col_val, rh2, vf, size=6.5)
        y -= rh2

    def mrow_recv(label, attr, dp=3):
        """Single value per tank (received row)."""
        nonlocal y
        rect(M, y - rh2, col0, rh2)
        cell(M, y - rh2, col0, rh2, label, size=6.5, align='L')
        for i in range(2):
            x0 = M + col0 + i * col_pair
            rect(x0, y - rh2, col_pair, rh2)
            v = _f(getattr(items[i], attr, None), dp) if i < n else ''
            cell(x0, y - rh2, col_pair, rh2, v, size=6.5)
        y -= rh2

    mrow('Overall Dip (Mm)',        'overall_dip_initial_mm',      'overall_dip_final_mm')
    mrow('Product Dip (Mm)',        'product_dip_initial_mm',      'product_dip_final_mm')
    mrow('Water Dip (Mm)',          'water_dip_initial_mm',        'water_dip_final_mm')
    mrow('Tank Temperature',        'tank_temperature_initial_c',  'tank_temperature_final_c', 1)
    mrow('Specific Gravity',        'density_initial_kg_l',        'density_final_kg_l', 4)
    mrow('Sample Temperature',      'sample_temperature_initial_c','sample_temperature_final_c', 1)
    mrow4('Density @20',            'density_initial_kg_l',        'density_final_kg_l')
    mrow4('VCF',                    'effective_vcf_initial',       'effective_vcf_final')
    mrow4('WCF',                    'effective_wcf_initial',       'effective_wcf_final')
    mrow('Gross Obs Volume',        'gross_observed_initial_m3',   'gross_observed_final_m3')
    mrow('Roof Displacement Vol.',  'roof_displacement_initial_m3','roof_displacement_final_m3')
    mrow('Water Volume',            'water_volume_initial_m3',     'water_volume_final_m3')
    mrow('Net Obs Volume',          'net_observed_initial_m3',     'net_observed_final_m3')
    mrow('Standard Vol@20',         'standard_volume_initial_m3',  'standard_volume_final_m3')
    mrow_recv('Standard Vol@20 Received', 'received_standard_volume_m3')
    mrow('Weight in Air (Mt)',      'weight_air_initial_mt',       'weight_air_final_mt')
    mrow_recv('Weight in Air Received (Mt)', 'received_weight_air_mt')

    y -= 2 * mm

    # ── Summary table ─────────────────────────────────────────────────────────
    sc0 = TW * 0.28
    sc1 = TW * 0.24
    sc2 = TW * 0.24
    sc3 = TW * 0.24
    srh = 6 * mm

    # Header
    for xi, (lbl, w) in enumerate([('STATUS', sc0), ('Obs Vol', sc1), ('M³', sc2), ('M/Tons', sc3)]):
        rx = M + sum([sc0, sc1, sc2, sc3][:xi])
        rect(rx, y - srh, w, srh)
        cell(rx, y - srh, w, srh, lbl, bold=True, size=7)
    y -= srh

    def srow(label, obs, std, wt):
        nonlocal y
        for xi, (val, w) in enumerate([(label, sc0), (_f(obs), sc1), (_f(std), sc2), (_f(wt), sc3)]):
            rx = M + sum([sc0, sc1, sc2, sc3][:xi])
            rect(rx, y - srh, w, srh)
            cell(rx, y - srh, w, srh, val, bold=(xi == 0), size=7, align='L' if xi == 0 else 'C')
        y -= srh

    srow('TERMINAL',  terminal_obs, terminal_std, terminal_wt)
    srow('VESSEL',    vessel_obs,   vessel_std,   vessel_wt)

    # Meter quantity row (spans obs+std columns)
    rect(M, y - srh, sc0, srh)
    rect(M + sc0, y - srh, sc1 + sc2, srh)
    rect(M + sc0 + sc1 + sc2, y - srh, sc3, srh)
    cell(M, y - srh, sc0, srh, 'METER QUANTITY', bold=True, size=7, align='L')
    cell(M + sc0, y - srh, sc1 + sc2, srh, _f(meter_qty), size=7)
    cell(M + sc0 + sc1 + sc2, y - srh, sc3, srh, '', size=7)
    y -= srh

    srow('DIFFERENCE (TERMINAL VS VESSEL)', diff_obs, diff_std, diff_wt)

    y -= 2 * mm

    # Inlet manifold seal
    rect(M, y - srh, TW, srh)
    cell(M, y - srh, TW * 0.35, srh, 'INLET MANIFOLD SEAL', bold=True, size=7, align='L')
    cell(M + TW * 0.35, y - srh, TW * 0.65, srh, '', size=7, align='L')
    y -= srh + 3 * mm

    # ── Signature blocks ──────────────────────────────────────────────────────
    # The terminal representative table is deliberately below the inspector
    # table.  Both extend across the report width for clear handwritten names
    # and signatures.
    lbl_w  = TW * 0.15
    data_cols = ['INITIAL', '1ST DISPL', '2ND DISPL', 'PROVISIONAL', 'FINAL']
    dcw = (TW - lbl_w) / len(data_cols)
    sig_rh = 5.5 * mm

    operation_label = next((
        (item.operation_label or '').strip().upper()
        for item in items if (item.operation_label or '').strip().upper() in {
            'INITIAL', '1ST DISPL', '2ND DISPL', 'L.DISPL', 'PROVISIONAL', 'FINAL'
        }
    ), 'INITIAL')
    operation_column = {
        'INITIAL': 0, '1ST DISPL': 1, 'L.DISPL': 1,
        '2ND DISPL': 2, 'PROVISIONAL': 3, 'FINAL': 4,
    }[operation_label]

    def draw_signature(image_data, x, row_y, w, h):
        _draw_signature_in_field(c, image_data, x, row_y - h, w, h)

    def signature_table(title, name, signed_at, role):
        nonlocal y
        c.setFont('Helvetica-Bold', 7)
        c.drawCentredString(M + TW / 2, y, title)
        y -= 5 * mm
        rect(M, y - sig_rh, lbl_w, sig_rh)
        cell(M, y - sig_rh, lbl_w, sig_rh, 'STATUS', bold=True, size=6)
        for j, dc in enumerate(data_cols):
            rect(M + lbl_w + j * dcw, y - sig_rh, dcw, sig_rh)
            cell(M + lbl_w + j * dcw, y - sig_rh, dcw, sig_rh, dc, bold=True, size=5.5)
        y -= sig_rh
        signed_date = signed_at.strftime('%d-%m-%Y') if signed_at else date_str
        row_data = [('STATUS', operation_label), ('NAME', name or ''), ('SIGNATURE', ''), ('DATE', signed_date)]
        image_data = _stored_signature_image('shore_tank', shore_calc.pk, role)
        for row_label, value in row_data:
            rect(M, y - sig_rh, lbl_w, sig_rh)
            cell(M, y - sig_rh, lbl_w, sig_rh, row_label, bold=True, size=6.5, align='L')
            for j in range(len(data_cols)):
                cell_x = M + lbl_w + j * dcw
                rect(cell_x, y - sig_rh, dcw, sig_rh)
                if j == operation_column and value:
                    cell(cell_x, y - sig_rh, dcw, sig_rh, value, size=6)
                if row_label == 'SIGNATURE' and j == operation_column:
                    draw_signature(image_data, cell_x, y, dcw, sig_rh)
            y -= sig_rh
        y -= 3 * mm

    signature_table('PETROLEUM BULK PROCUREMENT AGENCY INSPECTOR', shore_calc.pbpa_inspector_name,
                    getattr(shore_calc, 'inspector_signed_at', None), 'inspector')
    signature_table('TERMINAL REPRESENTATIVE', shore_calc.terminal_representative_name,
                    getattr(shore_calc, 'client_signed_at', None), 'terminal_representative')

    c.showPage()
    c.save()
    buf.seek(0)
    return buf
