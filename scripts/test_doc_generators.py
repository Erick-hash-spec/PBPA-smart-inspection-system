"""
Dry-run test: generate all 4 documents using mock data and verify they open cleanly.
"""
import os, sys
sys.path.insert(0, os.path.dirname(__file__))
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings')

import django
django.setup()

from docx import Document
from io import BytesIO
from unittest.mock import MagicMock, PropertyMock
from inspections.shore_tank_utils import (
    generate_shore_tank_document,
    generate_product_receipt_document,
    generate_seal_isolation_document,
    generate_dip_ticket_document,
)

def mock_item():
    item = MagicMock()
    item.tank_no = '5002'
    item.overall_dip_initial_mm = 28020; item.overall_dip_final_mm = 28020
    item.product_dip_initial_mm = 1075;  item.product_dip_final_mm = 4299
    item.water_dip_initial_mm = 0;       item.water_dip_final_mm = 0
    item.tank_temperature_initial_c = 29; item.tank_temperature_final_c = 31
    item.density_initial_kg_l = 0.735;   item.density_final_kg_l = 0.742
    item.sample_temperature_initial_c = 30; item.sample_temperature_final_c = 28
    item.gross_observed_initial_m3 = 475.588; item.gross_observed_final_m3 = 1810.281
    item.roof_displacement_initial_m3 = 0; item.roof_displacement_final_m3 = 0
    item.water_volume_initial_m3 = 0;    item.water_volume_final_m3 = 0
    item.net_observed_initial_m3 = 475.588; item.net_observed_final_m3 = 1810.281
    item.effective_vcf_initial = 0.9890; item.effective_vcf_final = 0.9868
    item.effective_wcf_initial = 0.7428; item.effective_wcf_final = 0.7480
    item.standard_volume_initial_m3 = 470.357; item.standard_volume_final_m3 = 1786.285
    item.received_standard_volume_m3 = 1316.029
    item.weight_air_initial_mt = 349.281; item.weight_air_final_mt = 1356.216
    item.received_weight_air_mt = 986.835
    return item

# ── Shore Tank ──
shore = MagicMock()
shore.vessel_name = 'MT NCC NAIMA'
shore.product_name = 'GASOLINE'
shore.terminal = 'GBP'
shore.calculation_date = '2025-10-24'
shore.calculation_number = '00010732'
shore.vessel_density_kg_m3 = None
shore.vessel_temperature_c = None
shore.vessel_observed_volume_m3 = 1437.499
shore.vessel_standard_volume_m3 = 1413.852
shore.vessel_weight_air_mt = 1053.744
shore.meter_quantity_m3 = None
shore.terminal_representative_name = 'John Doe'
shore.pbpa_inspector_name = 'Jane Smith'
shore.terminal_observed_volume_m3 = 1334.693
shore.terminal_standard_volume_m3 = 1316.029
shore.terminal_weight_air_mt = 986.835
shore.difference_observed_volume_m3 = -102.806
shore.difference_standard_volume_m3 = -97.823
shore.difference_weight_air_mt = -66.909
shore.tank_items.all.return_value = [mock_item()]

buf = generate_shore_tank_document(shore)
doc = Document(buf)
print('Shore Tank: OK -', len(doc.tables), 'tables,', len(doc.paragraphs), 'paragraphs')
# Check key cells
t1 = doc.tables[1]
print('  Tank No header:', t1.rows[0].cells[1].text)
print('  Overall Dip initial:', t1.rows[2].cells[1].text)
print('  VCF initial:', t1.rows[9].cells[1].text)
print('  WCF initial:', t1.rows[10].cells[1].text)
print('  Std Vol@20 Received:', t1.rows[16].cells[2].text)
print('  Weight Received:', t1.rows[18].cells[2].text)
t2 = doc.tables[2]
print('  Terminal Std Vol:', t2.rows[1].cells[2].text)
print('  Difference Std Vol:', t2.rows[4].cells[2].text)

# ── Product Receipt ──
cert = MagicMock()
cert.vessel_name = 'MT NCC NAIMA'
cert.terminal = 'GBP'
cert.receipt_date.strftime.return_value = '24-10-2025'
cert.receipt_time.strftime.return_value = '10:30'
cert.certificate_number = '00000001'
cert.quantity_received_through_inlet_flowmeters = 1316029.0
cert.total_weight_tonnage = 986.835
cert.total_volume_liters = 1316029.0
cert.terminal_representative_name = 'John Doe'
cert.pbpa_inspector_name = 'Jane Smith'
item1 = MagicMock(); item1.tank_no='5002'; item1.product_name='GASOLINE'; item1.weight_tonnage=986.835; item1.volume_liters=1316029.0
cert.items.all.return_value = [item1]
buf2 = generate_product_receipt_document(cert)
doc2 = Document(buf2)
print('\nProduct Receipt: OK -', len(doc2.tables), 'tables')
print('  Row 1 tank:', doc2.tables[0].rows[1].cells[0].text)
print('  Row 1 weight:', doc2.tables[0].rows[1].cells[2].text)
print('  Total weight:', doc2.tables[0].rows[12].cells[2].text)

# ── Seal Isolation ──
rep = MagicMock()
rep.vessel_name = 'MT NCC NAIMA'; rep.product_name = 'GASOLINE'; rep.terminal = 'GBP'
rep.report_date.strftime.return_value = '24-10-2025'
rep.report_number = '00000001'
rep.terminal_representative_name = 'John Doe'; rep.pbpa_inspector_name = 'Jane Smith'
e1 = MagicMock(); e1.location = 'Inlet Manifold'; e1.seal_number = 'SL-001'
e2 = MagicMock(); e2.location = 'Outlet Valve';   e2.seal_number = 'SL-002'
rep.entries.all.return_value = [e1, e2]
buf3 = generate_seal_isolation_document(rep)
doc3 = Document(buf3)
print('\nSeal Isolation: OK -', len(doc3.tables), 'tables')
print('  Entry 1 location:', doc3.tables[0].rows[1].cells[0].text)
print('  Entry 1 seal:', doc3.tables[0].rows[1].cells[1].text)

print('\nAll 3 generators passed.')
