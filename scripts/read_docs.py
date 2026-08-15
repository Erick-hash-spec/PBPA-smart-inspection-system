from docx import Document
from docx.shared import Pt
import sys

docs = [
    'DIP TICKET.docx',
    'SHORE TANK CALCULATIONS.docx',
    'PRODUCT RECEIPT CERTIFICATE.docx',
    'SEAL AND ISOLATION.docx',
]

for doc_name in docs:
    sep = '='*60
    print(sep)
    print(f'DOCUMENT: {doc_name}')
    print(sep)
    doc = Document(doc_name)
    print(f'Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}')

    print('\n-- PARAGRAPHS --')
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t:
            print(f'  [{i}] {p.style.name}: {t[:100]}')

    print('\n-- TABLES --')
    for ti, table in enumerate(doc.tables):
        print(f'\n  Table {ti}: {len(table.rows)}r x {len(table.columns)}c')
        for ri, row in enumerate(table.rows):
            cells = []
            for ci, cell in enumerate(row.cells):
                ct = cell.text.strip().replace('\n', ' ')[:40]
                cells.append(f'[{ci}]{ct}')
            print(f'    R{ri}: {" | ".join(cells)}')
    print()
