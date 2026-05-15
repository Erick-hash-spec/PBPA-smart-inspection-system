import os,sys; sys.path.insert(0,'.')
os.environ['DJANGO_SETTINGS_MODULE']='config.settings'
import django; django.setup()
from docx import Document
from unittest.mock import MagicMock
from inspections.shore_tank_utils import generate_shore_tank_document

item=MagicMock()
item.tank_no='5002'
for a in ['overall_dip_initial_mm','overall_dip_final_mm','product_dip_initial_mm','product_dip_final_mm','water_dip_initial_mm','water_dip_final_mm']:
    setattr(item,a,0)
item.tank_temperature_initial_c=29; item.tank_temperature_final_c=31
item.density_initial_kg_l=0.735; item.density_final_kg_l=0.742
item.sample_temperature_initial_c=30; item.sample_temperature_final_c=28
item.gross_observed_initial_m3=475.588; item.gross_observed_final_m3=1810.281
item.roof_displacement_initial_m3=0; item.roof_displacement_final_m3=0
item.water_volume_initial_m3=0; item.water_volume_final_m3=0
item.net_observed_initial_m3=475.588; item.net_observed_final_m3=1810.281
item.effective_vcf_initial=0.9890; item.effective_vcf_final=0.9868
item.effective_wcf_initial=0.7428; item.effective_wcf_final=0.7480
item.standard_volume_initial_m3=470.357; item.standard_volume_final_m3=1786.285
item.received_standard_volume_m3=1316.029
item.weight_air_initial_mt=349.281; item.weight_air_final_mt=1356.216
item.received_weight_air_mt=986.835

shore=MagicMock()
shore.vessel_name='MT NCC NAIMA'; shore.product_name='GASOLINE'; shore.terminal='GBP'
shore.calculation_date='2025-10-24'; shore.calculation_number='00010732'
shore.vessel_density_kg_m3=None; shore.vessel_temperature_c=None
shore.vessel_observed_volume_m3=1437.499; shore.vessel_standard_volume_m3=1413.852
shore.vessel_weight_air_mt=1053.744; shore.meter_quantity_m3=None
shore.terminal_representative_name='John Doe'; shore.pbpa_inspector_name='Jane Smith'
shore.terminal_observed_volume_m3=1334.693; shore.terminal_standard_volume_m3=1316.029
shore.terminal_weight_air_mt=986.835; shore.difference_observed_volume_m3=-102.806
shore.difference_standard_volume_m3=-97.823; shore.difference_weight_air_mt=-66.909
shore.tank_items.all.return_value=[item]

buf=generate_shore_tank_document(shore)
doc=Document(buf)
t1=doc.tables[1]

def uc(row):
    seen={}; cells=[]
    for c in row.cells:
        cid=id(c._tc)
        if cid not in seen:
            seen[cid]=True; cells.append(c)
    return cells

print('=== SHORE TANK DOCUMENT CELL VERIFICATION ===')
print('Header vessel:', doc.tables[0].rows[0].cells[1].text)
print('VCF initial:', uc(t1.rows[9])[1].text, '(expected 0.9890)')
print('VCF final:  ', uc(t1.rows[9])[2].text, '(expected 0.9868)')
print('WCF initial:', uc(t1.rows[10])[1].text, '(expected 0.7428)')
print('WCF final:  ', uc(t1.rows[10])[2].text, '(expected 0.7480)')
print('D@20 initial:', uc(t1.rows[8])[1].text, '(expected 0.7350)')
print('GOV initial:', uc(t1.rows[11])[1].text, '(expected 475.588)')
print('GOV final:  ', uc(t1.rows[11])[2].text, '(expected 1810.281)')
print('NOV initial:', uc(t1.rows[14])[1].text, '(expected 475.588)')
print('StdVol initial:', uc(t1.rows[15])[1].text, '(expected 470.357)')
print('StdVol final:  ', uc(t1.rows[15])[2].text, '(expected 1786.285)')
print('StdVol Received:', uc(t1.rows[16])[1].text, '(expected 1316.029)')
print('Wt initial:', uc(t1.rows[17])[1].text, '(expected 349.281)')
print('Wt final:  ', uc(t1.rows[17])[2].text, '(expected 1356.216)')
print('Wt Received:', uc(t1.rows[18])[1].text, '(expected 986.835)')
t2=doc.tables[2]
print('Terminal StdVol:', t2.rows[1].cells[2].text, '(expected 1316.029)')
print('Difference StdVol:', t2.rows[4].cells[2].text, '(expected -97.823)')
