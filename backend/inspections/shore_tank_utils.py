"""
Document generators that fill the original PBPA Word templates exactly,
preserving every table, paragraph, and layout from the source .docx files.
"""
import os
from io import BytesIO
from copy import deepcopy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .calculations import ShoreTankCalculationEngine, ASTMD1250Calculator

BASE = os.path.dirname(os.path.abspath(__file__))
TEMPLATES = {
    'shore_tank':   os.path.join(BASE, '..', 'SHORE TANK CALCULATIONS.docx'),
    'dip_ticket':   os.path.join(BASE, '..', 'DIP TICKET.docx'),
    'product_receipt': os.path.join(BASE, '..', 'PRODUCT RECEIPT CERTIFICATE.docx'),
    'seal_isolation':  os.path.join(BASE, '..', 'SEAL AND ISOLATION.docx'),
}


# ── helpers ──────────────────────────────────────────────────────────────────

def _set_unique_cell(table, row_idx, unique_col_idx, text):
    """
    Write to the Nth unique (non-merged) cell in a row.
    unique_col_idx: 0=label, 1=first data cell, 2=second data cell, etc.
    """
    try:
        row = table.rows[row_idx]
        seen = {}
        unique_cells = []
        for cell in row.cells:
            cid = id(cell._tc)
            if cid not in seen:
                seen[cid] = True
                unique_cells.append(cell)
        cell = unique_cells[unique_col_idx]
        para = cell.paragraphs[0]
        para.clear()
        para.add_run(str(text) if text is not None else '')
    except (IndexError, Exception):
        pass


def _set_cell(table, row, col, text, bold=False, align=None):
    """Write text into a table cell by absolute column index."""
    try:
        cell = table.rows[row].cells[col]
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(str(text) if text is not None else '')
        run.bold = bold
        if align:
            para.alignment = align
    except IndexError:
        pass


def _set_para(doc, para_index, text):
    """Replace the text of a paragraph by index."""
    try:
        p = doc.paragraphs[para_index]
        for run in p.runs:
            run.text = ''
        if p.runs:
            p.runs[0].text = text
        else:
            p.add_run(text)
    except IndexError:
        pass

def _f(v, d=3):
    """Format a number to d decimal places, empty string if None."""
    if v is None or v == '':
        return ''
    try:
        return f'{float(v):.{d}f}'
    except (TypeError, ValueError):
        return str(v)


def _f4(v):
    return _f(v, 4)


def _load(key):
    path = os.path.abspath(TEMPLATES[key])
    if not os.path.exists(path):
        raise FileNotFoundError(f'Template not found: {path}')
    return Document(path)


def _save(doc):
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Shore Tank Calculation ────────────────────────────────────────────────────

def generate_shore_tank_document(shore_calc):
    """
    Fill SHORE TANK CALCULATIONS.docx template.

    Template structure (verified from docx):
      Para [0]  THE UNITED REPUBLIC OF TANZANIA
      Para [1]  PETROLEUM BULK PROCUREMENT AGENCY
      Para [7]  TERMINAL REPRESENTATIVE  (signature block)

      Table 0 (4r x 4c) – header
        R0: VESSEL NAME | | Date |
        R1: PRODUCT     | | Vessel Density |
        R2: TERMINAL    | | Vessel Temperature |

      Table 1 (19r x 12c) – measurements (cols 1-2 = tank1 initial/final, 3-4 = tank2, etc.)
        R0:  Tank No. headers
        R1:  PARTICULARS | INITIAL | L.DISP/PROV/FINAL | ...
        R2:  Overall Dip(mm)
        R3:  Product Dip(mm)
        R4:  Water Dip(mm)
        R5:  Tank Temperature
        R6:  Specific Gravity(SG)
        R7:  Sample Temperature
        R8:  Density @20
        R9:  VCF
        R10: WCF
        R11: Gross Observed Volume
        R12: Roof Displacement Vol
        R13: Water Volume
        R14: Net Obs Volume
        R15: Standard Vol@20
        R16: Standard Vol@20 Received
        R17: Weight in Air (Mt)
        R18: Weight in Air Received (Mt)

      Table 2 (7r x 4c) – summary
        R0: STATUS | Obs Vol | M3 | M/Tons
        R1: TERMINAL
        R2: VESSEL
        R3: METER QUANTITY
        R4: DIFFERENCE(TERMINAL VS VESSEL)
        R6: INLET MANIFOLD SEAL

      Table 3 (4r x 6c) – Terminal Rep signatures
      Table 4 (4r x 6c) – PBPA Inspector signatures
    """
    doc = _load('shore_tank')
    items = list(shore_calc.tank_items.all())

    # ── Table 0: header ──
    t0 = doc.tables[0]
    _set_cell(t0, 0, 1, shore_calc.vessel_name)
    _set_cell(t0, 0, 3, str(shore_calc.calculation_date))
    _set_cell(t0, 1, 1, shore_calc.product_name)
    _set_cell(t0, 1, 3, _f4(shore_calc.vessel_density_kg_m3) if shore_calc.vessel_density_kg_m3 else '')
    _set_cell(t0, 2, 1, shore_calc.terminal)
    _set_cell(t0, 2, 3, _f(shore_calc.vessel_temperature_c, 1) if shore_calc.vessel_temperature_c else '')

    # ── Table 1: measurements ──
    # Unique cell layout per data row (rows 2-15, 17):
    #   unique[0]=label, unique[1]=tank1_initial, unique[2]=tank1_final,
    #   unique[3]=tank2_initial, unique[4]=tank2_final
    # Row 16 (Std Vol Received): unique[0]=label, unique[1]=tank1_recv, unique[3]=tank2_recv
    # Row 18 (Wt Received):      unique[0]=label, unique[1]=tank1_recv, unique[3]=tank2_recv
    t1 = doc.tables[1]

    # Row 0: single merged cell — write tank numbers as combined label
    tank_labels = '  |  '.join(item.tank_no or f'Tank {i+1}' for i, item in enumerate(items[:2]))
    _set_unique_cell(t1, 0, 0, tank_labels)

    # Data rows per tank pair (template supports 2 tanks in unique cells 1-4)
    STANDARD_ROWS = [
        # (row_idx, attr_initial, attr_final)
        (2,  'overall_dip_initial_mm',       'overall_dip_final_mm'),
        (3,  'product_dip_initial_mm',        'product_dip_final_mm'),
        (4,  'water_dip_initial_mm',          'water_dip_final_mm'),
        (5,  'tank_temperature_initial_c',    'tank_temperature_final_c'),
        (6,  'density_initial_kg_l',          'density_final_kg_l'),
        (7,  'sample_temperature_initial_c',  'sample_temperature_final_c'),
        (11, 'gross_observed_initial_m3',     'gross_observed_final_m3'),
        (12, 'roof_displacement_initial_m3',  'roof_displacement_final_m3'),
        (13, 'water_volume_initial_m3',       'water_volume_final_m3'),
        (14, 'net_observed_initial_m3',       'net_observed_final_m3'),
        (15, 'standard_volume_initial_m3',    'standard_volume_final_m3'),
        (17, 'weight_air_initial_mt',         'weight_air_final_mt'),
    ]

    for idx, item in enumerate(items[:2]):
        ui = 1 + idx * 2   # unique cell index for initial
        uf = 2 + idx * 2   # unique cell index for final

        for row_idx, attr_i, attr_f in STANDARD_ROWS:
            _set_unique_cell(t1, row_idx, ui, _f(getattr(item, attr_i, None)))
            _set_unique_cell(t1, row_idx, uf, _f(getattr(item, attr_f, None)))

        # ASTM rows (4 decimal places)
        vcf_i = item.effective_vcf_initial
        vcf_f = item.effective_vcf_final
        wcf_i = item.effective_wcf_initial
        wcf_f = item.effective_wcf_final
        _set_unique_cell(t1, 8,  ui, _f4(item.density_initial_kg_l))
        _set_unique_cell(t1, 8,  uf, _f4(item.density_final_kg_l))
        _set_unique_cell(t1, 9,  ui, _f4(vcf_i))
        _set_unique_cell(t1, 9,  uf, _f4(vcf_f))
        _set_unique_cell(t1, 10, ui, _f4(wcf_i))
        _set_unique_cell(t1, 10, uf, _f4(wcf_f))

        # Received rows have different merge patterns per row:
        # Row 16 unique: [0=label, 1=t1, 3=t2, 6=t3, 8=t4, 11=t5]
        # Row 18 unique: [0=label, 1=t1, 2=t2, 5=t3, 8=t4, 10=t5]
        recv16_idx = [1, 3, 6, 8, 11]
        recv18_idx = [1, 2, 5, 8, 10]
        if idx < len(recv16_idx):
            _set_unique_cell(t1, 16, recv16_idx[idx], _f(item.received_standard_volume_m3))
        if idx < len(recv18_idx):
            _set_unique_cell(t1, 18, recv18_idx[idx], _f(item.received_weight_air_mt))

    # ── Table 2: summary ──
    t2 = doc.tables[2]
    _set_cell(t2, 1, 1, _f(shore_calc.terminal_observed_volume_m3))
    _set_cell(t2, 1, 2, _f(shore_calc.terminal_standard_volume_m3))
    _set_cell(t2, 1, 3, _f(shore_calc.terminal_weight_air_mt))

    _set_cell(t2, 2, 1, _f(shore_calc.vessel_observed_volume_m3))
    _set_cell(t2, 2, 2, _f(shore_calc.vessel_standard_volume_m3))
    _set_cell(t2, 2, 3, _f(shore_calc.vessel_weight_air_mt))

    _set_cell(t2, 3, 2, _f(shore_calc.meter_quantity_m3))

    _set_cell(t2, 4, 1, _f(shore_calc.difference_observed_volume_m3))
    _set_cell(t2, 4, 2, _f(shore_calc.difference_standard_volume_m3))
    _set_cell(t2, 4, 3, _f(shore_calc.difference_weight_air_mt))

    # ── Table 3: Terminal Rep signatures ──
    t3 = doc.tables[3]
    _set_cell(t3, 1, 0, 'NAME')
    _set_cell(t3, 1, 1, shore_calc.terminal_representative_name or '')

    # ── Table 4: PBPA Inspector signatures ──
    t4 = doc.tables[4]
    _set_cell(t4, 1, 0, 'NAME')
    _set_cell(t4, 1, 1, shore_calc.pbpa_inspector_name or '')

    return _save(doc)


# ── Dip Ticket ────────────────────────────────────────────────────────────────

def generate_dip_ticket_document(inspection):
    """
    Fill DIP TICKET.docx template.

    Template structure:
      Para [0]  THE UNITED REPUBLIC OF TANZANIA
      Para [9]  Terminal representative: ... PBPA Inspector:
      Para [11] Name ... Name
      Para [13] Signature ... Signature

      Table 0 (3r x 2c) – vessel/product/terminal header
        R0: Vessel: | Product:
        R1: Terminal: | Tank No:
        R2: Date: | Time:

      Table 1 (10r x 5c) – measurements
        R0: PARTICULARS | MEASUREMENTS x4
        R1:  | 1st | 2nd | 3rd | Average
        R2:  Overall Dip (mm)
        R3:  Product Dip (mm)
        R4:  Product volume (L)
        R5:  Free water volume (L)
        R6:  Tank temperature
        R7:  Specific gravity (SG)
        R8:  Sample temperature
        R9:  (spare)

      Table 2 (4r x 4c) – seals & meter
        R0: PBPA SEAL POSITION | PBPA SEAL NUMBER | METER READINGS x2
        R1: Outlet valve seal  | | OBS |
        R2: Water valve seal   | | @20 |
        R3: Other branches seal| | MTS |
    """
    doc = _load('dip_ticket')

    # ── Table 0: header ──
    t0 = doc.tables[0]
    _set_cell(t0, 0, 0, f'Vessel: {inspection.vessel_name or ""}')
    _set_cell(t0, 0, 1, f'Product: {inspection.product_name or ""}')
    _set_cell(t0, 1, 0, f'Terminal: {inspection.terminal or ""}')
    _set_cell(t0, 1, 1, f'Tank No: {inspection.tank.tank_id if inspection.tank else ""}')
    date_str = inspection.inspection_date.strftime('%d-%m-%Y') if inspection.inspection_date else ''
    time_str = inspection.inspection_time.strftime('%H:%M') if inspection.inspection_time else ''
    _set_cell(t0, 2, 0, f'Date: {date_str}')
    _set_cell(t0, 2, 1, f'Time: {time_str}')

    # ── Table 1: measurements ──
    t1 = doc.tables[1]

    def _avg(vals):
        nums = [v for v in vals if v is not None]
        return round(sum(nums) / len(nums), 3) if nums else None

    rows = [
        # (row_idx, [v1, v2, v3])
        (2, [inspection.overall_dip_1_mm,       inspection.overall_dip_2_mm,       inspection.overall_dip_3_mm]),
        (3, [inspection.product_dip_1_mm,        inspection.product_dip_2_mm,       inspection.product_dip_3_mm]),
        (4, [inspection.product_volume_1_l,      inspection.product_volume_2_l,     inspection.product_volume_3_l]),
        (5, [inspection.free_water_volume_1_l,   inspection.free_water_volume_2_l,  inspection.free_water_volume_3_l]),
        (6, [inspection.tank_temperature_1_c,    inspection.tank_temperature_2_c,   inspection.tank_temperature_3_c]),
        (7, [inspection.specific_gravity_1,      inspection.specific_gravity_2,     inspection.specific_gravity_3]),
        (8, [inspection.sample_temperature_1_c,  inspection.sample_temperature_2_c, inspection.sample_temperature_3_c]),
    ]
    for row_idx, vals in rows:
        for col, v in enumerate(vals, start=1):
            _set_cell(t1, row_idx, col, _f(v))
        _set_cell(t1, row_idx, 4, _f(_avg(vals)))

    # ── Table 2: seals & meter ──
    t2 = doc.tables[2]
    _set_cell(t2, 1, 1, inspection.outlet_valve_seal_number or '')
    _set_cell(t2, 2, 1, inspection.water_valve_seal_number or '')
    _set_cell(t2, 3, 1, inspection.other_branches_seal_number or '')
    _set_cell(t2, 1, 3, _f(inspection.meter_reading_obs))
    _set_cell(t2, 2, 3, _f(inspection.meter_reading_at_20))
    _set_cell(t2, 3, 3, _f(inspection.meter_reading_mts))

    # ── Signature paragraphs ──
    # Para 11: Name line
    try:
        p11 = doc.paragraphs[11]
        p11.clear()
        r = p11.add_run(f'Name: {inspection.terminal_representative_name or ""}')
        r2_text = f'Name: {inspection.pbpa_inspector_name or ""}'
        # add spacing then second name
        p11.add_run('                                                    ' + r2_text)
    except Exception:
        pass

    return _save(doc)


# ── Product Receipt Certificate ───────────────────────────────────────────────

def generate_product_receipt_document(certificate):
    """
    Fill PRODUCT RECEIPT CERTIFICATE.docx template.

    Template structure:
      Para [6]  Vessel Name:
      Para [7]  Terminal:..
      Para [8]  Date:..
      Para [9]  Time:.
      Para [10] This is to certify...
      Para [12] Quantity received through inlet flowmeters:
      Para [14] Terminal representative: ... PBPA Inspector:
      Para [16] Name ... Name
      Para [18] Signature ... Signature

      Table 0 (13r x 4c) – items
        R0:  Tank No. | Product | Weight in Tonnage | Volume in Liters
        R1-R11: data rows
        R12: TOTAL | TOTAL | total_weight | total_volume

      Table 1 (2r x 2c) – totals summary
        R0: Weight in Tonnage | Volume in Litres
        R1: value | value
    """
    doc = _load('product_receipt')

    # ── Header paragraphs ──
    paras = doc.paragraphs
    _set_para_text(paras, 6, f'Vessel Name: {certificate.vessel_name}')
    _set_para_text(paras, 7, f'Terminal: {certificate.terminal}')
    date_str = certificate.receipt_date.strftime('%d-%m-%Y') if certificate.receipt_date else ''
    _set_para_text(paras, 8, f'Date: {date_str}')
    time_str = certificate.receipt_time.strftime('%H:%M') if certificate.receipt_time else ''
    _set_para_text(paras, 9, f'Time: {time_str}')
    _set_para_text(paras, 12,
        f'Quantity received through inlet flowmeters: '
        f'{_f(certificate.quantity_received_through_inlet_flowmeters)} Liters'
    )

    # ── Table 0: items ──
    t0 = doc.tables[0]
    items = list(certificate.items.all())
    data_rows = 11  # rows 1-11

    for i in range(data_rows):
        if i < len(items):
            item = items[i]
            _set_cell(t0, i + 1, 0, item.tank_no or '')
            _set_cell(t0, i + 1, 1, item.product_name or '')
            _set_cell(t0, i + 1, 2, _f(item.weight_tonnage))
            _set_cell(t0, i + 1, 3, _f(item.volume_liters))
        else:
            for col in range(4):
                _set_cell(t0, i + 1, col, '')

    # Total row (row 12)
    _set_cell(t0, 12, 2, _f(certificate.total_weight_tonnage))
    _set_cell(t0, 12, 3, _f(certificate.total_volume_liters))

    # ── Table 1: summary totals ──
    t1 = doc.tables[1]
    _set_cell(t1, 1, 0, _f(certificate.total_weight_tonnage))
    _set_cell(t1, 1, 1, _f(certificate.total_volume_liters))

    # ── Signature paragraphs ──
    _set_para_text(paras, 16,
        f'Name: {certificate.terminal_representative_name or ""}' +
        '                                                    ' +
        f'Name: {certificate.pbpa_inspector_name or ""}'
    )

    return _save(doc)


# ── Seal and Isolation ────────────────────────────────────────────────────────

def generate_seal_isolation_document(report):
    """
    Fill SEAL AND ISOLATION.docx template.

    Template structure:
      Para [4]  Vessel Name:.
      Para [5]  Product:.
      Para [6]  Terminal:...
      Para [7]  Date:..
      Para [14] Name ... Name
      Para [16] Signature ... Signature

      Table 0 (16r x 2c) – seal entries
        R0:  Location | Seal Number
        R1-R15: data rows
    """
    doc = _load('seal_isolation')
    paras = doc.paragraphs

    _set_para_text(paras, 4, f'Vessel Name: {report.vessel_name}')
    _set_para_text(paras, 5, f'Product: {report.product_name}')
    _set_para_text(paras, 6, f'Terminal: {report.terminal}')
    date_str = report.report_date.strftime('%d-%m-%Y') if report.report_date else ''
    _set_para_text(paras, 7, f'Date: {date_str}')

    # ── Table 0: entries ──
    t0 = doc.tables[0]
    entries = list(report.entries.all())
    data_rows = 15  # rows 1-15

    for i in range(data_rows):
        if i < len(entries):
            _set_cell(t0, i + 1, 0, entries[i].location or '')
            _set_cell(t0, i + 1, 1, entries[i].seal_number or '')
        else:
            _set_cell(t0, i + 1, 0, '')
            _set_cell(t0, i + 1, 1, '')

    # ── Signature paragraphs ──
    _set_para_text(paras, 14,
        f'Name: {report.terminal_representative_name or ""}' +
        '                                                    ' +
        f'Name: {report.pbpa_inspector_name or ""}'
    )

    return _save(doc)


# ── internal helper ───────────────────────────────────────────────────────────

def _set_para_text(paras, idx, text):
    """Replace all text in a paragraph while keeping the first run's formatting."""
    try:
        p = paras[idx]
        # clear all runs
        for run in p.runs:
            run.text = ''
        if p.runs:
            p.runs[0].text = text
        else:
            p.add_run(text)
    except IndexError:
        pass


# ── Compatibility shim (keeps existing views.py calls working) ────────────────

class ShoreTankDocumentGenerator:
    def fill_shore_tank_calculation_document(self, shore_calc, template_path=None):
        return generate_shore_tank_document(shore_calc)

    def fill_product_receipt_certificate_document(self, certificate, template_path=None):
        return generate_product_receipt_document(certificate)

    def fill_seal_isolation_document(self, report, template_path=None):
        return generate_seal_isolation_document(report)

    def fill_dip_ticket_document(self, inspection, template_path=None):
        return generate_dip_ticket_document(inspection)


# ── PDF Generation Functions ──────────────────────────────────────────────

from reportlab.lib.pagesizes import A4, landscape
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
import io


def generate_dip_ticket_pdf(inspection):
    """Generate Dip Ticket as PDF using ReportLab."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=10*mm, leftMargin=10*mm, topMargin=10*mm, bottomMargin=10*mm)
    
    elements = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#1a1a1a'),
        spaceAfter=6,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    elements.append(Paragraph("THE UNITED REPUBLIC OF TANZANIA", title_style))
    elements.append(Paragraph("PETROLEUM BULK PROCUREMENT AGENCY - DIP TICKET", title_style))
    elements.append(Spacer(1, 8*mm))
    
    # Header info table - use Paragraph for HTML support
    header_style = ParagraphStyle(
        'HeaderCell',
        parent=styles['Normal'],
        fontSize=9,
        alignment=TA_LEFT
    )
    
    header_data = [
        [
            Paragraph(f"<b>Vessel:</b> {inspection.vessel_name or ''}", header_style),
            Paragraph(f"<b>Product:</b> {inspection.product_name or ''}", header_style)
        ],
        [
            Paragraph(f"<b>Terminal:</b> {inspection.terminal or ''}", header_style),
            Paragraph(f"<b>Tank No:</b> {inspection.tank.tank_id if inspection.tank else ''}", header_style)
        ],
        [
            Paragraph(f"<b>Date:</b> {inspection.inspection_date.strftime('%d-%m-%Y') if inspection.inspection_date else ''}", header_style),
            Paragraph(f"<b>Time:</b> {inspection.inspection_time.strftime('%H:%M') if inspection.inspection_time else ''}", header_style)
        ]
    ]
    
    header_table = Table(header_data, colWidths=[doc.width/2-5*mm, doc.width/2-5*mm])
    header_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f5f5f5')),
    ]))
    elements.append(header_table)
    elements.append(Spacer(1, 8*mm))
    
    # Measurements table
    def _avg(vals):
        nums = [v for v in vals if v is not None]
        return round(sum(nums) / len(nums), 3) if nums else None
    
    normal_style = ParagraphStyle('Normal2', parent=styles['Normal'], fontSize=9, alignment=TA_CENTER)
    
    measurements_data = [
        ['PARTICULARS', '1st', '2nd', '3rd', 'Average'],
        ['Overall Dip (mm)', _f(inspection.overall_dip_1_mm), _f(inspection.overall_dip_2_mm), _f(inspection.overall_dip_3_mm), _f(_avg([inspection.overall_dip_1_mm, inspection.overall_dip_2_mm, inspection.overall_dip_3_mm]))],
        ['Product Dip (mm)', _f(inspection.product_dip_1_mm), _f(inspection.product_dip_2_mm), _f(inspection.product_dip_3_mm), _f(_avg([inspection.product_dip_1_mm, inspection.product_dip_2_mm, inspection.product_dip_3_mm]))],
        ['Product Volume (L)', _f(inspection.product_volume_1_l), _f(inspection.product_volume_2_l), _f(inspection.product_volume_3_l), _f(_avg([inspection.product_volume_1_l, inspection.product_volume_2_l, inspection.product_volume_3_l]))],
        ['Free Water Volume (L)', _f(inspection.free_water_volume_1_l), _f(inspection.free_water_volume_2_l), _f(inspection.free_water_volume_3_l), _f(_avg([inspection.free_water_volume_1_l, inspection.free_water_volume_2_l, inspection.free_water_volume_3_l]))],
        ['Tank Temperature (°C)', _f(inspection.tank_temperature_1_c, 1), _f(inspection.tank_temperature_2_c, 1), _f(inspection.tank_temperature_3_c, 1), _f(_avg([inspection.tank_temperature_1_c, inspection.tank_temperature_2_c, inspection.tank_temperature_3_c]), 1)],
        ['Specific Gravity (SG)', _f(inspection.specific_gravity_1), _f(inspection.specific_gravity_2), _f(inspection.specific_gravity_3), _f(_avg([inspection.specific_gravity_1, inspection.specific_gravity_2, inspection.specific_gravity_3]))],
        ['Sample Temperature (°C)', _f(inspection.sample_temperature_1_c, 1), _f(inspection.sample_temperature_2_c, 1), _f(inspection.sample_temperature_3_c, 1), _f(_avg([inspection.sample_temperature_1_c, inspection.sample_temperature_2_c, inspection.sample_temperature_3_c]), 1)],
    ]
    
    measurements_table = Table(measurements_data, colWidths=[doc.width*0.35, doc.width*0.13, doc.width*0.13, doc.width*0.13, doc.width*0.13])
    measurements_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#8B1A1A')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGNMENT', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f9f9f9')]),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
    ]))
    elements.append(measurements_table)
    elements.append(Spacer(1, 8*mm))
    
    # Seals & meter readings table
    seals_data = [
        ['PBPA SEAL POSITION', 'SEAL NUMBER', 'METER READINGS', ''],
        ['', '', 'OBS', '@20'],
        ['Outlet valve seal', inspection.outlet_valve_seal_number or '', _f(inspection.meter_reading_obs), ''],
        ['Water valve seal', inspection.water_valve_seal_number or '', _f(inspection.meter_reading_at_20), ''],
        ['Other branches seal', inspection.other_branches_seal_number or '', _f(inspection.meter_reading_mts), ''],
    ]
    
    seals_table = Table(seals_data, colWidths=[doc.width*0.35, doc.width*0.25, doc.width*0.2, doc.width*0.2])
    seals_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#8B1A1A')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGNMENT', (1, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ROWBACKGROUNDS', (0, 2), (-1, -1), [colors.white, colors.HexColor('#f9f9f9')]),
    ]))
    elements.append(seals_table)
    elements.append(Spacer(1, 12*mm))
    
    # Signatures
    sig_data = [
        ['Terminal Representative', '', 'PBPA Inspector'],
        ['Name: ' + (inspection.terminal_representative_name or ''), '', 'Name: ' + (inspection.pbpa_inspector_name or '')],
        ['Signature: _____________________', '', 'Signature: _____________________'],
    ]
    
    sig_table = Table(sig_data, colWidths=[doc.width/3-5*mm, doc.width/3-5*mm, doc.width/3-5*mm])
    sig_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGNMENT', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
    ]))
    elements.append(sig_table)
    
    doc.build(elements)
    buffer.seek(0)
    return buffer


def generate_seal_isolation_pdf(report):
    """Generate Seal & Isolation Report as PDF using ReportLab."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=10*mm, leftMargin=10*mm, topMargin=10*mm, bottomMargin=10*mm)
    
    elements = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#1a1a1a'),
        spaceAfter=6,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    elements.append(Paragraph("THE UNITED REPUBLIC OF TANZANIA", title_style))
    elements.append(Paragraph("PBPA SEALING AND ISOLATION REPORT", title_style))
    elements.append(Spacer(1, 8*mm))
    
    # Header info - use Paragraph for HTML support
    header_style = ParagraphStyle(
        'HeaderCell',
        parent=styles['Normal'],
        fontSize=9,
        alignment=TA_LEFT
    )
    
    header_data = [
        [
            Paragraph(f"<b>Vessel Name:</b> {report.vessel_name}", header_style),
            Paragraph(f"<b>Product:</b> {report.product_name}", header_style)
        ],
        [
            Paragraph(f"<b>Terminal:</b> {report.terminal}", header_style),
            Paragraph(f"<b>Date:</b> {report.report_date.strftime('%d-%m-%Y')}", header_style)
        ],
        [
            Paragraph(f"<b>Report No:</b> {report.report_number}", header_style),
            Paragraph(f"<b>Status:</b> {report.status.upper()}", header_style)
        ],
    ]
    
    header_table = Table(header_data, colWidths=[doc.width/2-5*mm, doc.width/2-5*mm])
    header_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f5f5f5')),
    ]))
    elements.append(header_table)
    elements.append(Spacer(1, 8*mm))
    
    # Seals entries table
    entries_data = [['Location', 'Seal Number']]
    for entry in report.entries.all():
        entries_data.append([entry.location or '', entry.seal_number or ''])
    
    # Pad with empty rows if fewer than 15
    while len(entries_data) < 16:
        entries_data.append(['', ''])
    
    entries_table = Table(entries_data, colWidths=[doc.width*0.6, doc.width*0.4])
    entries_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#8B1A1A')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f9f9f9')]),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
    ]))
    elements.append(entries_table)
    elements.append(Spacer(1, 12*mm))
    
    # Notes
    if report.notes:
        elements.append(Paragraph("<b>Notes:</b>", styles['Normal']))
        elements.append(Paragraph(report.notes, styles['Normal']))
        elements.append(Spacer(1, 6*mm))
    
    # Signatures
    sig_data = [
        ['Terminal Representative', '', 'PBPA Inspector'],
        ['Name: ' + (report.terminal_representative_name or ''), '', 'Name: ' + (report.pbpa_inspector_name or '')],
        ['Signature: _____________________', '', 'Signature: _____________________'],
    ]
    
    sig_table = Table(sig_data, colWidths=[doc.width/3-5*mm, doc.width/3-5*mm, doc.width/3-5*mm])
    sig_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGNMENT', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
    ]))
    elements.append(sig_table)
    
    doc.build(elements)
    buffer.seek(0)
    return buffer


def generate_shore_tank_pdf(shore_calc):
    """
    Generate Shore Tank Calculation PDF matching the official PBPA format:
    - Portrait A4
    - PBPA letterhead
    - Vessel/Product/Terminal header box
    - SHORE TANK CALCULATIONS title with document number
    - Full measurements table: PARTICULARS | INITIAL | L.DISP/PROV/FINAL (per tank)
    - Summary table: TERMINAL / VESSEL / METER QUANTITY / DIFFERENCE
    - PBPA Inspector & Terminal Representative signature blocks
    """
    buf = io.BytesIO()
    W, H = A4
    M = 12 * mm
    TW = W - 2 * M
    c = rl_canvas.Canvas(buf, pagesize=A4)

    def line(x1, y1, x2, y2, lw=0.5):
        c.setLineWidth(lw)
        c.line(x1, y1, x2, y2)

    def rect(x, y, w, h, fill=0, lw=0.5):
        c.setLineWidth(lw)
        c.rect(x, y, w, h, fill=fill)

    def cell(x, y, w, h, text, font='Helvetica', size=7, align='C', bold=False):
        fn = 'Helvetica-Bold' if bold else font
        c.setFont(fn, size)
        tx = x + w / 2 if align == 'C' else (x + w - 1*mm if align == 'R' else x + 1*mm)
        ty = y + h / 2 - size * 0.35
        c.drawString(tx, ty, str(text)) if align == 'L' else (
            c.drawRightString(tx, ty, str(text)) if align == 'R' else
            c.drawCentredString(tx, ty, str(text))
        )

    items = list(shore_calc.tank_items.all())
    n = min(len(items), 2)

    # Compute terminal totals from tank items (no stored summary fields on model)
    def _sum_attr(attr):
        vals = [getattr(it, attr, None) for it in items]
        nums = [float(v) for v in vals if v is not None]
        return round(sum(nums), 3) if nums else None

    terminal_obs = _sum_attr('gross_observed_final_m3')
    terminal_std = _sum_attr('standard_volume_final_m3')
    terminal_wt  = _sum_attr('weight_air_final_mt')
    vessel_obs   = shore_calc.vessel_observed_volume_m3
    vessel_std   = shore_calc.vessel_standard_volume_m3
    vessel_wt    = shore_calc.vessel_weight_air_mt
    meter_qty    = shore_calc.meter_quantity_m3

    def _diff(a, b):
        if a is not None and b is not None:
            return round(float(a) - float(b), 3)
        return None

    diff_obs = _diff(terminal_obs, vessel_obs)
    diff_std = _diff(terminal_std, vessel_std)
    diff_wt  = _diff(terminal_wt, vessel_wt)

    y = H - 8 * mm

    # ── Letterhead ────────────────────────────────────────────────────────────
    c.setFont('Helvetica', 7.5)
    c.drawCentredString(W / 2, y, 'THE UNITED REPUBLIC OF TANZANIA')
    y -= 5 * mm
    c.setFont('Helvetica-Bold', 10)
    c.drawCentredString(W / 2, y, 'PETROLEUM BULK PROCUREMENT AGENCY')
    y -= 4 * mm
    c.setFont('Helvetica', 6)
    c.drawCentredString(W / 2, y, 'TANZANIA PORTS AUTHORITY, ONE STOP CENTER BUILDING 11TH FLOOR, SOKOINE DRIVE, PLOT NO:1/2')
    y -= 3.5 * mm
    c.drawCentredString(W / 2, y,
        'Tel: +255222129009 / Fax: +255222129093 / info@pbpa.go.tz / WEBSITE: www.pbpa.go.tz / P.O. BOX 2634 Dar es Salaam, TANZANIA')
    y -= 4 * mm
    line(M, y, W - M, y)
    y -= 3 * mm

    # ── Vessel / Product / Terminal header box ────────────────────────────────
    rh = 6 * mm
    lw1 = TW * 0.28
    lw2 = TW * 0.32
    lw3 = TW * 0.40
    date_str = shore_calc.calculation_date.strftime('%d-%m-%Y') if shore_calc.calculation_date else ''

    rows_hdr = [
        ('VESSEL NAME', shore_calc.vessel_name or '', 'Date', date_str),
        ('PRODUCT',     shore_calc.product_name or '', 'Vessel Density', _f4(shore_calc.vessel_density_kg_m3) if shore_calc.vessel_density_kg_m3 else ''),
        ('TERMINAL',    shore_calc.terminal or '',    'Vessel Temperature', _f(shore_calc.vessel_temperature_c, 1) if shore_calc.vessel_temperature_c else ''),
    ]
    for label, val, label2, val2 in rows_hdr:
        rect(M, y - rh, lw1, rh)
        rect(M + lw1, y - rh, lw2, rh)
        rect(M + lw1 + lw2, y - rh, lw3 * 0.45, rh)
        rect(M + lw1 + lw2 + lw3 * 0.45, y - rh, lw3 * 0.55, rh)
        cell(M, y - rh, lw1, rh, label, bold=True, size=7)
        cell(M + lw1, y - rh, lw2, rh, val, size=7, align='L')
        cell(M + lw1 + lw2, y - rh, lw3 * 0.45, rh, label2, bold=True, size=7)
        cell(M + lw1 + lw2 + lw3 * 0.45, y - rh, lw3 * 0.55, rh, val2, size=7, align='L')
        y -= rh
    y -= 2 * mm

    # ── Title + Document Number ───────────────────────────────────────────────
    title_h = 8 * mm
    rect(M, y - title_h, TW * 0.65, title_h, lw=1.5)
    c.setFont('Helvetica-Bold', 11)
    c.drawCentredString(M + TW * 0.65 / 2, y - title_h + 2.5 * mm, 'SHORE TANK CALCULATIONS')
    # Doc number box top-right
    rect(M + TW * 0.65 + 2 * mm, y - title_h, TW * 0.35 - 2 * mm, title_h, lw=1.5)
    c.setFont('Helvetica-Bold', 9)
    c.drawCentredString(M + TW * 0.65 + 2 * mm + (TW * 0.35 - 2 * mm) / 2,
                        y - title_h + 2.5 * mm,
                        shore_calc.calculation_number or '')
    y -= title_h + 2 * mm

    # ── Measurements table ────────────────────────────────────────────────────
    # Columns: PARTICULARS | Tank1 INITIAL | Tank1 L.DISP/PROV/FINAL | Tank2 INITIAL | Tank2 L.DISP/PROV/FINAL
    # (if only 1 tank, last 2 cols are blank)
    col0 = TW * 0.30
    col_pair = (TW - col0) / 2  # width per tank pair
    col_val = col_pair / 2

    # Tank number header row
    rh2 = 5.5 * mm
    rect(M, y - rh2, col0, rh2)
    cell(M, y - rh2, col0, rh2, 'Tank No.', bold=True, size=7)
    for i in range(2):
        x0 = M + col0 + i * col_pair
        rect(x0, y - rh2, col_pair, rh2)
        tank_no = items[i].tank_no if i < n else ''
        cell(x0, y - rh2, col_pair, rh2, tank_no, bold=True, size=8)
    y -= rh2

    # Sub-header: INITIAL / L.DISP/PROV/FINAL
    rect(M, y - rh2, col0, rh2)
    cell(M, y - rh2, col0, rh2, 'PARTICULARS', bold=True, size=7)
    for i in range(2):
        x0 = M + col0 + i * col_pair
        rect(x0, y - rh2, col_val, rh2)
        rect(x0 + col_val, y - rh2, col_val, rh2)
        cell(x0, y - rh2, col_val, rh2, 'INITIAL', bold=True, size=6.5)
        cell(x0 + col_val, y - rh2, col_val, rh2, 'L.DISP/PROV/FINAL', bold=True, size=5.5)
    y -= rh2

    def mrow(label, attr_i, attr_f, dp=3):
        nonlocal y
        rect(M, y - rh2, col0, rh2)
        cell(M, y - rh2, col0, rh2, label, size=6.5, align='L')
        for i in range(2):
            x0 = M + col0 + i * col_pair
            rect(x0, y - rh2, col_val, rh2)
            rect(x0 + col_val, y - rh2, col_val, rh2)
            vi = _f(getattr(items[i], attr_i, None), dp) if i < n else ''
            vf = _f(getattr(items[i], attr_f, None), dp) if i < n else ''
            cell(x0, y - rh2, col_val, rh2, vi, size=6.5)
            cell(x0 + col_val, y - rh2, col_val, rh2, vf, size=6.5)
        y -= rh2

    def mrow4(label, attr_i, attr_f, dp=4):
        """4 decimal places for ASTM values."""
        nonlocal y
        rect(M, y - rh2, col0, rh2)
        cell(M, y - rh2, col0, rh2, label, size=6.5, align='L')
        for i in range(2):
            x0 = M + col0 + i * col_pair
            rect(x0, y - rh2, col_val, rh2)
            rect(x0 + col_val, y - rh2, col_val, rh2)
            vi = _f(getattr(items[i], attr_i, None), dp) if i < n else ''
            vf = _f(getattr(items[i], attr_f, None), dp) if i < n else ''
            cell(x0, y - rh2, col_val, rh2, vi, size=6.5)
            cell(x0 + col_val, y - rh2, col_val, rh2, vf, size=6.5)
        y -= rh2

    def mrow_recv(label, attr, dp=3):
        """Single value per tank (received row)."""
        nonlocal y
        rect(M, y - rh2, col0, rh2)
        cell(M, y - rh2, col0, rh2, label, size=6.5, align='L')
        for i in range(2):
            x0 = M + col0 + i * col_pair
            rect(x0, y - rh2, col_pair, rh2)
            v = _f(getattr(items[i], attr, None), dp) if i < n else ''
            cell(x0, y - rh2, col_pair, rh2, v, size=6.5)
        y -= rh2

    mrow('Overall Dip (Mm)',        'overall_dip_initial_mm',      'overall_dip_final_mm')
    mrow('Product Dip (Mm)',        'product_dip_initial_mm',      'product_dip_final_mm')
    mrow('Water Dip (Mm)',          'water_dip_initial_mm',        'water_dip_final_mm')
    mrow('Tank Temperature',        'tank_temperature_initial_c',  'tank_temperature_final_c', 1)
    mrow('Specific Gravity',        'density_initial_kg_l',        'density_final_kg_l', 4)
    mrow('Sample Temperature',      'sample_temperature_initial_c','sample_temperature_final_c', 1)
    mrow4('Density @20',            'density_initial_kg_l',        'density_final_kg_l')
    mrow4('VCF',                    'effective_vcf_initial',       'effective_vcf_final')
    mrow4('WCF',                    'effective_wcf_initial',       'effective_wcf_final')
    mrow('Gross Obs Volume',        'gross_observed_initial_m3',   'gross_observed_final_m3')
    mrow('Roof Displacement Vol.',  'roof_displacement_initial_m3','roof_displacement_final_m3')
    mrow('Water Volume',            'water_volume_initial_m3',     'water_volume_final_m3')
    mrow('Net Obs Volume',          'net_observed_initial_m3',     'net_observed_final_m3')
    mrow('Standard Vol@20',         'standard_volume_initial_m3',  'standard_volume_final_m3')
    mrow_recv('Standard Vol@20 Received', 'received_standard_volume_m3')
    mrow('Weight in Air (Mt)',      'weight_air_initial_mt',       'weight_air_final_mt')
    mrow_recv('Weight in Air Received (Mt)', 'received_weight_air_mt')

    y -= 2 * mm

    # ── Summary table ─────────────────────────────────────────────────────────
    sc0 = TW * 0.28
    sc1 = TW * 0.24
    sc2 = TW * 0.24
    sc3 = TW * 0.24
    srh = 6 * mm

    # Header
    for xi, (lbl, w) in enumerate([('STATUS', sc0), ('Obs Vol', sc1), ('M³', sc2), ('M/Tons', sc3)]):
        rx = M + sum([sc0, sc1, sc2, sc3][:xi])
        rect(rx, y - srh, w, srh)
        cell(rx, y - srh, w, srh, lbl, bold=True, size=7)
    y -= srh

    def srow(label, obs, std, wt):
        nonlocal y
        for xi, (val, w) in enumerate([(label, sc0), (_f(obs), sc1), (_f(std), sc2), (_f(wt), sc3)]):
            rx = M + sum([sc0, sc1, sc2, sc3][:xi])
            rect(rx, y - srh, w, srh)
            cell(rx, y - srh, w, srh, val, bold=(xi == 0), size=7, align='L' if xi == 0 else 'C')
        y -= srh

    srow('TERMINAL',  terminal_obs, terminal_std, terminal_wt)
    srow('VESSEL',    vessel_obs,   vessel_std,   vessel_wt)

    # Meter quantity row (spans obs+std columns)
    rect(M, y - srh, sc0, srh)
    rect(M + sc0, y - srh, sc1 + sc2, srh)
    rect(M + sc0 + sc1 + sc2, y - srh, sc3, srh)
    cell(M, y - srh, sc0, srh, 'METER QUANTITY', bold=True, size=7, align='L')
    cell(M + sc0, y - srh, sc1 + sc2, srh, _f(meter_qty), size=7)
    cell(M + sc0 + sc1 + sc2, y - srh, sc3, srh, '', size=7)
    y -= srh

    srow('DIFFERENCE (TERMINAL VS VESSEL)', diff_obs, diff_std, diff_wt)

    y -= 2 * mm

    # Inlet manifold seal
    rect(M, y - srh, TW, srh)
    cell(M, y - srh, TW * 0.35, srh, 'INLET MANIFOLD SEAL', bold=True, size=7, align='L')
    cell(M + TW * 0.35, y - srh, TW * 0.65, srh, '', size=7, align='L')
    y -= srh + 3 * mm

    # ── Signature blocks ──────────────────────────────────────────────────────
    # Two blocks side by side: PBPA Inspector (left) | Terminal Rep (right)
    gap = 6 * mm
    bw = (TW - gap) / 2          # each block width
    pbpa_x = M
    term_x = M + bw + gap

    # Block title headers
    c.setFont('Helvetica-Bold', 7)
    c.drawCentredString(pbpa_x + bw / 2, y, 'PETROLEUM BULK PROCUREMENT AGENCY INSPECTOR')
    c.drawCentredString(term_x + bw / 2, y, 'TERMINAL REPRESENTATIVE')
    y -= 5 * mm

    # Each block: label col + 5 data cols (INITIAL, 1ST DISPL, 2ND DISPL, PROVISIONAL, FINAL)
    lbl_w  = bw * 0.24
    data_cols = ['INITIAL', '1ST DISPL', '2ND DISPL', 'PROVISIONAL', 'FINAL']
    dcw = (bw - lbl_w) / len(data_cols)
    sig_rh = 5.5 * mm

    # Sub-column header row
    for bx in [pbpa_x, term_x]:
        rect(bx, y - sig_rh, lbl_w, sig_rh)
        cell(bx, y - sig_rh, lbl_w, sig_rh, 'STATUS', bold=True, size=6)
        for j, dc in enumerate(data_cols):
            rect(bx + lbl_w + j * dcw, y - sig_rh, dcw, sig_rh)
            cell(bx + lbl_w + j * dcw, y - sig_rh, dcw, sig_rh, dc, bold=True, size=5.5)
    y -= sig_rh

    # Data rows: STATUS / NAME / SIGNATURE / DATE
    for bi, bx in enumerate([pbpa_x, term_x]):
        name_val = shore_calc.pbpa_inspector_name if bi == 0 else shore_calc.terminal_representative_name
        row_data = [
            ('STATUS',    ['', '', '', '', '']),
            ('NAME',      [name_val or '', '', '', '', '']),
            ('SIGNATURE', ['', '', '', '', '']),
            ('DATE',      [date_str, '', '', '', '']),
        ]
        ry = y
        for row_lbl, vals in row_data:
            rect(bx, ry - sig_rh, lbl_w, sig_rh)
            cell(bx, ry - sig_rh, lbl_w, sig_rh, row_lbl, bold=True, size=6.5, align='L')
            for j, v in enumerate(vals):
                rect(bx + lbl_w + j * dcw, ry - sig_rh, dcw, sig_rh)
                if v:
                    cell(bx + lbl_w + j * dcw, ry - sig_rh, dcw, sig_rh, v, size=6)
            ry -= sig_rh

    c.showPage()
    c.save()
    buf.seek(0)
    return buf
